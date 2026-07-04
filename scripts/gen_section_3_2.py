#!/usr/bin/env python3
"""Generate Section 3.2 Data Model & Data Governance for Thai Zoo ARK proposal."""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.font_manager as fm
import numpy as np
from io import BytesIO

# Register ALL Thai font variants for matplotlib
_FONT_DIR = os.path.expanduser("~/Library/Fonts/")
for _fname in [
    "THSarabunNew.ttf", "THSarabunNew Bold.ttf",
    "THSarabunNew Italic.ttf", "THSarabunNew BoldItalic.ttf",
    "THSarabun.ttf", "THSarabun Bold.ttf",
]:
    _path = os.path.join(_FONT_DIR, _fname)
    if os.path.exists(_path):
        fm.fontManager.addfont(_path)

plt.rcParams['font.family'] = 'TH Sarabun New'
plt.rcParams['font.sans-serif'] = ['TH Sarabun New', 'TH SarabunPSK', 'DejaVu Sans']

# ── Constants ──
FONT_TH = "TH Sarabun New"
FONT_TH_PSK = "TH SarabunPSK"
BODY_SIZE = Pt(16)
HEADING_SIZE = Pt(16)
CAPTION_SIZE = Pt(14)
HEADER_BG = "D6E3F8"

OUT_DIR = os.path.expanduser(
    "~/Library/CloudStorage/OneDrive-Personal/Personal/ZTL"
)
OUT_FILE = os.path.join(OUT_DIR, "Section_3_2_Data_Model_Data_Governance.docx")


# ── Helpers ──

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name=FONT_TH, size=BODY_SIZE, bold=False, italic=False, color=None):
    """Configure run font properties for Thai text."""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    # Set complex script font (for Thai)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="{font_name}" w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
    # Set complex script size
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="{int(size.pt * 2)}"/>')
        rPr.append(szCs)
    else:
        szCs.set(qn('w:val'), str(int(size.pt * 2)))
    # bold cs
    if bold:
        bCs = rPr.find(qn('w:bCs'))
        if bCs is None:
            rPr.append(parse_xml(f'<w:bCs {nsdecls("w")}/>'))
    if color:
        run.font.color.rgb = color


def add_chapter_heading(doc, number, title):
    """Add X.X level heading (Chapter within chapter 3)."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    # Spacing
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="240" w:after="120"/>')
    pPr.append(spacing)
    # Indent
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="567" w:hanging="567"/>')
    pPr.append(ind)
    # Tab
    tabs = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="567"/></w:tabs>')
    pPr.append(tabs)

    run = p.add_run(f"{number}\t{title}")
    set_run_font(run, bold=True, size=HEADING_SIZE)
    return p


def add_section_heading(doc, number, title):
    """Add X.X.X level heading."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="200" w:after="80"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="1418" w:hanging="851"/>')
    pPr.append(ind)
    tabs = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="1418"/></w:tabs>')
    pPr.append(tabs)

    run = p.add_run(f"{number}\t{title}")
    set_run_font(run, bold=True, size=HEADING_SIZE)
    return p


def add_subsection_heading(doc, number, title):
    """Add X.X.X.X level heading."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="160" w:after="80"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="2268" w:hanging="850"/>')
    pPr.append(ind)
    tabs = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="2268"/></w:tabs>')
    pPr.append(tabs)

    run = p.add_run(f"{number}\t{title}")
    set_run_font(run, bold=True, size=HEADING_SIZE)
    return p


def add_body_after_chapter(doc, text):
    """Body paragraph after X.X heading."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="120" w:line="240" w:lineRule="auto"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="567"/>')
    pPr.append(ind)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_body_after_section(doc, text):
    """Body paragraph after X.X.X heading."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="120" w:line="240" w:lineRule="auto"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="1418"/>')
    pPr.append(ind)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_body_after_subsection(doc, text):
    """Body paragraph after X.X.X.X heading."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="120" w:line="240" w:lineRule="auto"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="2268"/>')
    pPr.append(ind)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_numbered_item_section(doc, number, text):
    """Numbered item after X.X.X heading."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="60" w:line="240" w:lineRule="auto"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="1778" w:hanging="360"/>')
    pPr.append(ind)
    run = p.add_run(f"{number}) {text}")
    set_run_font(run)
    return p


def add_numbered_item_chapter(doc, number, text):
    """Numbered item after X.X heading."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="60" w:line="240" w:lineRule="auto"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="927" w:hanging="360"/>')
    pPr.append(ind)
    run = p.add_run(f"{number}) {text}")
    set_run_font(run)
    return p


def add_table_caption(doc, caption_text):
    """Table caption above table, left-aligned, bold."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="240" w:after="120"/>')
    pPr.append(spacing)
    run = p.add_run(caption_text)
    set_run_font(run, bold=True)
    return p


def add_figure_caption(doc, caption_text):
    """Figure caption below figure, center, bold+italic."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="240" w:after="120"/>')
    pPr.append(spacing)
    run = p.add_run(caption_text)
    set_run_font(run, bold=True, italic=True, size=CAPTION_SIZE)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, bold=True)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_image_from_fig(doc, fig, width_inches=5.5):
    """Save matplotlib figure to buffer and add to doc."""
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))
    plt.close(fig)
    return p


# ── Diagram Generation Functions ──

def create_governance_framework_diagram():
    """Create Data Governance Framework diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')

    # Title
    ax.text(5, 6.6, 'กรอบธรรมาภิบาลข้อมูล (Data Governance Framework)',
            ha='center', va='center', fontsize=13, fontweight='bold',
            fontfamily='sans-serif')

    # Central box - Data Governance Council
    center = FancyBboxPatch((3.2, 4.5), 3.6, 1.2, boxstyle="round,pad=0.1",
                             facecolor='#1a5276', edgecolor='#1a5276', alpha=0.9)
    ax.add_patch(center)
    ax.text(5, 5.1, 'Data Governance\nCouncil', ha='center', va='center',
            fontsize=11, color='white', fontweight='bold')

    # Six pillars
    pillars = [
        (0.5, 2.5, '#2980b9', 'Master Data\nManagement'),
        (2.3, 2.5, '#27ae60', 'Data Quality\n& Validation'),
        (4.1, 2.5, '#8e44ad', 'Data Security\n& Access'),
        (5.9, 2.5, '#e67e22', 'Data Lineage\n& Audit'),
        (7.7, 2.5, '#c0392b', 'Data Lifecycle\n& Retention'),
        (0.5, 0.5, '#16a085', 'Data Dictionary\n& Standards'),
    ]
    # Add more pillars in bottom row
    pillars_bottom = [
        (2.3, 0.5, '#2c3e50', 'Integration\n& Interoperability'),
        (4.1, 0.5, '#d4ac0d', 'Data Privacy\n& Compliance'),
        (5.9, 0.5, '#7f8c8d', 'Monitoring\n& Reporting'),
        (7.7, 0.5, '#e74c3c', 'Change\nManagement'),
    ]

    for x, y, color, label in pillars:
        box = FancyBboxPatch((x, y), 1.6, 1.2, boxstyle="round,pad=0.08",
                              facecolor=color, edgecolor=color, alpha=0.85)
        ax.add_patch(box)
        ax.text(x + 0.8, y + 0.6, label, ha='center', va='center',
                fontsize=8, color='white', fontweight='bold')
        # Arrow from center to pillar
        ax.annotate('', xy=(x + 0.8, y + 1.2), xytext=(5, 4.5),
                    arrowprops=dict(arrowstyle='->', color='#555555', lw=1.2))

    for x, y, color, label in pillars_bottom:
        box = FancyBboxPatch((x, y), 1.6, 1.2, boxstyle="round,pad=0.08",
                              facecolor=color, edgecolor=color, alpha=0.85)
        ax.add_patch(box)
        ax.text(x + 0.8, y + 0.6, label, ha='center', va='center',
                fontsize=8, color='white', fontweight='bold')
        ax.annotate('', xy=(x + 0.8, y + 1.2), xytext=(5, 4.5),
                    arrowprops=dict(arrowstyle='->', color='#555555', lw=1.2))

    fig.tight_layout()
    return fig


def create_master_data_diagram():
    """Create Master Data relationship diagram centered on Animal Registry."""
    fig, ax = plt.subplots(1, 1, figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.axis('off')

    ax.text(5, 7.6, 'โครงสร้าง Master Data — Animal ID เป็น Single Source of Truth',
            ha='center', va='center', fontsize=12, fontweight='bold')

    # Central entity
    center = FancyBboxPatch((3.5, 4.8), 3, 1.4, boxstyle="round,pad=0.1",
                             facecolor='#c0392b', edgecolor='#922b21', linewidth=2)
    ax.add_patch(center)
    ax.text(5, 5.5, 'Animal Registry\n(Single Source of Truth)\nanimal_id | species_id | name\nsex | dob | status',
            ha='center', va='center', fontsize=9, color='white', fontweight='bold')

    # Reference data boxes (top)
    refs_top = [
        (0.3, 6.5, '#2980b9', 'Species Taxonomy\nspecies_id | sci_name\ncommon_name | class'),
        (7.2, 6.5, '#8e44ad', 'Zoo / Location\nzoo_id | zoo_name\nregion | province'),
    ]
    for x, y, color, label in refs_top:
        box = FancyBboxPatch((x, y), 2.5, 1.0, boxstyle="round,pad=0.08",
                              facecolor=color, alpha=0.85)
        ax.add_patch(box)
        ax.text(x + 1.25, y + 0.5, label, ha='center', va='center',
                fontsize=8, color='white', fontweight='bold')

    # Dependent modules
    deps = [
        (0.2, 3.2, '#27ae60', 'Studbook\nstudbook_no\nsire_id | dam_id'),
        (0.2, 1.2, '#16a085', 'Genetic Analysis\nf_coefficient\ncompleteness_idx'),
        (3.5, 1.2, '#e67e22', 'Billing\ninvoice_id\nprocedure | medicine'),
        (3.5, 3.2, '#2c3e50', 'Medical Records\nvisit_id | diagnosis\nvet_id | treatment'),
        (7.0, 3.2, '#d4ac0d', 'Nutrition\ndiet_card_id\nfood_item | qty'),
        (7.0, 1.2, '#7f8c8d', 'Keeper Care\ncare_log_id\nactivity | notes'),
    ]
    for x, y, color, label in deps:
        box = FancyBboxPatch((x, y), 2.5, 1.2, boxstyle="round,pad=0.08",
                              facecolor=color, alpha=0.85)
        ax.add_patch(box)
        ax.text(x + 1.25, y + 0.6, label, ha='center', va='center',
                fontsize=8, color='white', fontweight='bold')
        # Arrow from center to dependent
        ax.annotate('', xy=(x + 1.25, min(y + 1.2, 4.8)),
                    xytext=(5, 4.8),
                    arrowprops=dict(arrowstyle='->', color='#333', lw=1.5,
                                    connectionstyle='arc3,rad=0.1'))

    # Arrows from ref to center
    ax.annotate('', xy=(3.5, 5.8), xytext=(2.8, 7.0),
                arrowprops=dict(arrowstyle='->', color='#333', lw=1.5))
    ax.annotate('', xy=(6.5, 5.8), xytext=(7.2, 7.0),
                arrowprops=dict(arrowstyle='->', color='#333', lw=1.5))

    # Legend
    ax.text(0.3, 0.3, 'FK = Foreign Key อ้างอิงไปยัง Animal Registry  |  '
            'ข้อมูลเข้าถึงผ่าน Read-Only API  |  ไม่เขียนตรงไปที่ตารางเดิม',
            fontsize=8, style='italic', color='#555')

    fig.tight_layout()
    return fig


def create_er_diagram():
    """Create conceptual ER diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(11, 8))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 9)
    ax.axis('off')

    ax.text(5.5, 8.6, 'Entity-Relationship Diagram — ระบบ Thai Zoo ARK (Conceptual)',
            ha='center', va='center', fontsize=12, fontweight='bold')

    # Entities
    entities = {
        'animal':    (4.0, 6.2, 3.0, 1.5, '#c0392b', 'ANIMAL\n(ทะเบียนสัตว์)\nPK: animal_id\nspecies_id | sex | dob\nstatus | zoo_id'),
        'species':   (0.3, 7.0, 2.5, 1.0, '#2980b9', 'SPECIES\n(อนุกรมวิธาน)\nPK: species_id\nsci_name | class'),
        'zoo':       (8.2, 7.0, 2.5, 1.0, '#8e44ad', 'ZOO\n(สวนสัตว์)\nPK: zoo_id\nname | region'),
        'studbook':  (0.3, 4.3, 2.8, 1.2, '#27ae60', 'STUDBOOK\n(ทะเบียนสายพันธุ์)\nPK: studbook_no\nFK: animal_id\nFK: sire_id | dam_id'),
        'genetic':   (0.3, 2.3, 2.8, 1.2, '#16a085', 'GENETIC_ANALYSIS\n(วิเคราะห์พันธุกรรม)\nPK: id\nFK: studbook_id\nf_coefficient'),
        'pedigree':  (0.3, 0.3, 2.8, 1.0, '#1abc9c', 'PEDIGREE_CACHE\nPK: id\nFK: animal_id\nancestor_path | depth'),
        'medical':   (4.0, 4.0, 3.0, 1.2, '#2c3e50', 'MEDICAL_RECORD\n(เวชระเบียน)\nPK: visit_id\nFK: animal_id\ndiagnosis | vet_id'),
        'nutrition': (8.0, 4.5, 2.8, 1.2, '#d4ac0d', 'NUTRITION\n(โภชนาการ)\nPK: diet_card_id\nFK: animal_id\nfood_item | qty'),
        'billing':   (4.0, 1.5, 3.0, 1.2, '#e67e22', 'BILLING_INVOICE\n(ค่ารักษา)\nPK: invoice_id\nFK: animal_id\ntotal | status'),
        'pharmacy':  (8.0, 2.5, 2.8, 1.2, '#e74c3c', 'PHARMACY\n(ยา/เวชภัณฑ์)\nPK: drug_id\nbatch | source\nexpiry | barcode'),
        'syslog':    (8.0, 0.5, 2.8, 1.0, '#7f8c8d', 'SYSTEM_LOG\nPK: id\nuser_id | action\ntimestamp | module'),
        'user':      (4.0, 0.0, 3.0, 1.0, '#34495e', 'USER\n(เจ้าหน้าที่)\nPK: user_id\nrole_id | zoo_id'),
    }

    for key, (x, y, w, h, color, label) in entities.items():
        box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.06",
                              facecolor=color, alpha=0.88, edgecolor='#222', linewidth=1.2)
        ax.add_patch(box)
        ax.text(x + w/2, y + h/2, label, ha='center', va='center',
                fontsize=7, color='white', fontweight='bold')

    # Relationships (simplified arrows)
    rels = [
        ('species', 'animal', '1:N'),
        ('zoo', 'animal', '1:N'),
        ('animal', 'studbook', '1:1'),
        ('studbook', 'genetic', '1:N'),
        ('animal', 'pedigree', '1:N'),
        ('animal', 'medical', '1:N'),
        ('animal', 'nutrition', '1:N'),
        ('animal', 'billing', '1:N'),
        ('medical', 'pharmacy', 'N:M'),
        ('user', 'syslog', '1:N'),
        ('medical', 'billing', '1:1'),
    ]

    for src, dst, card in rels:
        sx, sy, sw, sh = entities[src][0], entities[src][1], entities[src][2], entities[src][3]
        dx, dy, dw, dh = entities[dst][0], entities[dst][1], entities[dst][2], entities[dst][3]
        ax.annotate('', xy=(dx + dw/2, dy + dh/2), xytext=(sx + sw/2, sy + sh/2),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=0.8,
                                    connectionstyle='arc3,rad=0.15'))

    # Self-referencing for animal (sire/dam)
    ax.annotate('', xy=(7.0, 7.2), xytext=(7.0, 6.5),
                arrowprops=dict(arrowstyle='->', color='#c0392b', lw=1.5,
                                connectionstyle='arc3,rad=-0.5'))
    ax.text(7.8, 7.2, 'Self-ref\n(sire/dam)', fontsize=7, color='#c0392b', style='italic')

    fig.tight_layout()
    return fig


def create_integration_strategy_diagram():
    """Create data integration strategy diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')

    ax.text(5, 5.7, 'กลยุทธ์การบูรณาการข้อมูลระหว่างโมดูลใหม่และระบบเดิม',
            ha='center', va='center', fontsize=12, fontweight='bold')

    # Existing system
    existing = FancyBboxPatch((0.3, 2.5), 3.5, 2.5, boxstyle="round,pad=0.1",
                               facecolor='#ecf0f1', edgecolor='#2c3e50', linewidth=2)
    ax.add_patch(existing)
    ax.text(2.05, 4.7, 'ระบบเดิม (Existing)', ha='center', fontsize=10,
            fontweight='bold', color='#2c3e50')
    ax.text(2.05, 4.1, 'Animal Registry DB\nMedical Records\nNutrition\nUser Management',
            ha='center', fontsize=8, color='#2c3e50')
    ax.text(2.05, 2.8, 'Schema ไม่ถูกแก้ไข', ha='center', fontsize=8,
            color='#c0392b', fontweight='bold', style='italic')

    # API Gateway
    gw = FancyBboxPatch((4.2, 2.8), 1.6, 2.0, boxstyle="round,pad=0.08",
                          facecolor='#f39c12', edgecolor='#d68910', linewidth=2)
    ax.add_patch(gw)
    ax.text(5.0, 3.8, 'API\nGateway\n(Read-Only\nAccess)', ha='center',
            fontsize=9, color='white', fontweight='bold')

    # New modules
    new = FancyBboxPatch((6.2, 2.5), 3.5, 2.5, boxstyle="round,pad=0.1",
                          facecolor='#eafaf1', edgecolor='#27ae60', linewidth=2)
    ax.add_patch(new)
    ax.text(7.95, 4.7, 'โมดูลใหม่ (New Modules)', ha='center', fontsize=10,
            fontweight='bold', color='#27ae60')
    ax.text(7.95, 4.1, 'Studbook | Genetic | Pedigree\nBilling | Dashboard\nChange Request | System Log',
            ha='center', fontsize=8, color='#27ae60')
    ax.text(7.95, 2.8, 'Additive Schema Only', ha='center', fontsize=8,
            color='#27ae60', fontweight='bold', style='italic')

    # Arrows
    ax.annotate('', xy=(4.2, 3.8), xytext=(3.8, 3.8),
                arrowprops=dict(arrowstyle='->', color='#2c3e50', lw=2))
    ax.annotate('', xy=(6.2, 3.8), xytext=(5.8, 3.8),
                arrowprops=dict(arrowstyle='->', color='#27ae60', lw=2))

    # Labels
    ax.text(4.0, 4.2, 'Read', fontsize=8, color='#2c3e50', ha='center')
    ax.text(6.0, 4.2, 'Read', fontsize=8, color='#27ae60', ha='center')

    # Bottom: validation layer
    val = FancyBboxPatch((1.5, 0.5), 7.0, 1.5, boxstyle="round,pad=0.1",
                          facecolor='#fdebd0', edgecolor='#e67e22', linewidth=1.5)
    ax.add_patch(val)
    ax.text(5.0, 1.6, 'Data Validation & Quality Layer', ha='center',
            fontsize=10, fontweight='bold', color='#e67e22')
    ax.text(5.0, 0.9, 'Unique Constraints  |  FK Integrity Check  |  Duplicate Detection  |  '
            'Business Rules  |  Audit Trail',
            ha='center', fontsize=8, color='#e67e22')

    ax.annotate('', xy=(5.0, 2.0), xytext=(5.0, 2.5),
                arrowprops=dict(arrowstyle='<->', color='#e67e22', lw=1.5))

    fig.tight_layout()
    return fig


def create_data_lineage_diagram():
    """Create data lineage and audit trail flow diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5.5)
    ax.axis('off')

    ax.text(5, 5.2, 'Data Lineage และ Audit Trail Flow',
            ha='center', va='center', fontsize=12, fontweight='bold')

    # Flow steps
    steps = [
        (0.3, 3.0, 1.8, 1.5, '#3498db', 'User Action\n(Create/Update/\nDelete)'),
        (2.5, 3.0, 1.8, 1.5, '#2ecc71', 'Validation\nLayer\n(Business Rules\n+ Data Quality)'),
        (4.7, 3.0, 1.8, 1.5, '#9b59b6', 'Data Store\n(Write to\nTarget Table)'),
        (6.9, 3.0, 1.8, 1.5, '#e67e22', 'Audit Log\n(who/when/what\n/before/after)'),
        (6.9, 0.8, 1.8, 1.5, '#e74c3c', 'System Log\n(timestamp\nuser_id | IP\nmodule | action)'),
    ]

    for x, y, w, h, color, label in steps:
        box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08",
                              facecolor=color, alpha=0.85)
        ax.add_patch(box)
        ax.text(x + w/2, y + h/2, label, ha='center', va='center',
                fontsize=8, color='white', fontweight='bold')

    # Arrows between steps
    for i in range(len(steps) - 2):
        x1 = steps[i][0] + steps[i][2]
        x2 = steps[i+1][0]
        y_mid = steps[i][1] + steps[i][3]/2
        ax.annotate('', xy=(x2, y_mid), xytext=(x1, y_mid),
                    arrowprops=dict(arrowstyle='->', color='#333', lw=1.5))

    # Arrow from Data Store to Audit Log
    ax.annotate('', xy=(6.9, 3.75), xytext=(6.5, 3.75),
                arrowprops=dict(arrowstyle='->', color='#333', lw=1.5))
    # Arrow from Audit Log to System Log
    ax.annotate('', xy=(7.8, 2.3), xytext=(7.8, 3.0),
                arrowprops=dict(arrowstyle='->', color='#333', lw=1.5))

    # Metadata fields
    meta_box = FancyBboxPatch((0.3, 0.5), 5.5, 1.8, boxstyle="round,pad=0.08",
                               facecolor='#f8f9fa', edgecolor='#6c757d', linewidth=1.2)
    ax.add_patch(meta_box)
    ax.text(3.05, 2.0, 'Audit Metadata Fields (ทุกตาราง)', ha='center',
            fontsize=9, fontweight='bold', color='#2c3e50')
    ax.text(3.05, 1.3, 'created_by | created_at | updated_by | updated_at\n'
            'deleted_by | deleted_at | deletion_reason | cr_reference\n'
            'version_no | source_module | ip_address',
            ha='center', fontsize=8, color='#2c3e50')

    fig.tight_layout()
    return fig


def create_data_quality_diagram():
    """Create data quality process diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5.5)
    ax.axis('off')

    ax.text(5, 5.2, 'กระบวนการควบคุมคุณภาพข้อมูลและป้องกันความซ้ำซ้อน',
            ha='center', va='center', fontsize=12, fontweight='bold')

    # Input
    inp = FancyBboxPatch((0.2, 3.2), 1.5, 1.3, boxstyle="round,pad=0.08",
                          facecolor='#3498db', alpha=0.85)
    ax.add_patch(inp)
    ax.text(0.95, 3.85, 'Data Input\n(User/API/\nImport)', ha='center',
            fontsize=8, color='white', fontweight='bold')

    # Validation steps
    checks = [
        (2.2, 3.2, '#e74c3c', 'Schema\nValidation\n(Type/Format/\nRequired)'),
        (3.8, 3.2, '#e67e22', 'Business\nRule Check\n(Range/Logic/\nDiet Card)'),
        (5.4, 3.2, '#f1c40f', 'Duplicate\nDetection\n(Animal ID/\nStudbook No)'),
        (7.0, 3.2, '#27ae60', 'Referential\nIntegrity\n(FK Exists/\nStatus Valid)'),
    ]
    for x, y, color, label in checks:
        box = FancyBboxPatch((x, y), 1.3, 1.3, boxstyle="round,pad=0.06",
                              facecolor=color, alpha=0.85)
        ax.add_patch(box)
        ax.text(x + 0.65, y + 0.65, label, ha='center', va='center',
                fontsize=7, color='white', fontweight='bold')

    # Output
    ok = FancyBboxPatch((8.7, 3.5), 1.1, 0.8, boxstyle="round,pad=0.06",
                         facecolor='#2ecc71', alpha=0.9)
    ax.add_patch(ok)
    ax.text(9.25, 3.9, 'PASS\nSave', ha='center', fontsize=9,
            color='white', fontweight='bold')

    reject = FancyBboxPatch((4.0, 0.8), 2.0, 1.0, boxstyle="round,pad=0.06",
                             facecolor='#c0392b', alpha=0.85)
    ax.add_patch(reject)
    ax.text(5.0, 1.3, 'REJECT\nError Message\nto User', ha='center',
            fontsize=9, color='white', fontweight='bold')

    # Arrows
    ax.annotate('', xy=(2.2, 3.85), xytext=(1.7, 3.85),
                arrowprops=dict(arrowstyle='->', color='#333', lw=1.3))
    for i in range(len(checks)-1):
        ax.annotate('', xy=(checks[i+1][0], 3.85), xytext=(checks[i][0]+1.3, 3.85),
                    arrowprops=dict(arrowstyle='->', color='#333', lw=1.3))
    ax.annotate('', xy=(8.7, 3.9), xytext=(8.3, 3.85),
                arrowprops=dict(arrowstyle='->', color='#333', lw=1.3))
    # Reject arrows down
    for x, y, _, _ in checks:
        ax.annotate('', xy=(5.0, 1.8), xytext=(x+0.65, 3.2),
                    arrowprops=dict(arrowstyle='->', color='#c0392b', lw=0.8,
                                    linestyle='dashed', connectionstyle='arc3,rad=0.2'))

    ax.text(5, 2.3, 'Fail', fontsize=8, color='#c0392b', ha='center', style='italic')

    fig.tight_layout()
    return fig


# ──────────────────────────────────────────
# MAIN DOCUMENT CONSTRUCTION
# ──────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Twips(11906 * 20)  # A4
    section.page_height = Twips(16838 * 20)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_TH
    font.size = BODY_SIZE
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="{FONT_TH}" w:ascii="{FONT_TH}" w:hAnsi="{FONT_TH}"/>')
        rPr.insert(0, rFonts)

    # ══════════════════════════════════════════════════════════════
    # 3.2  Data Model & Data Governance
    # ══════════════════════════════════════════════════════════════

    add_chapter_heading(doc, "3.2", "แบบจำลองข้อมูลและธรรมาภิบาลข้อมูล (Data Model & Data Governance)")

    add_body_after_chapter(doc,
        "การออกแบบแบบจำลองข้อมูลและการวางกรอบธรรมาภิบาลข้อมูลเป็นรากฐานสำคัญที่สุดของการพัฒนา"
        "ปรับปรุงระบบ Thai Zoo ARK ให้สามารถทำงานร่วมกับระบบเดิมได้อย่างราบรื่น โดยไม่กระทบ"
        "โครงสร้างฐานข้อมูลและการทำงานที่มีอยู่ ทั้งยังสามารถใช้ข้อมูลจากระบบเดิมได้อย่างเหมาะสม "
        "ไม่เกิดความซ้ำซ้อนหรือขัดแย้ง ตรงตามเกณฑ์การให้คะแนนข้อ 3 ที่กำหนดไว้ใน TOR "
        "หัวข้อนี้ประกอบด้วยเนื้อหา 8 ส่วน ส่วนแรกคือหลักการธรรมาภิบาลข้อมูล 6 ประการ "
        "(หัวข้อ 3.2.1) ซึ่งวางกรอบการกำกับดูแลข้อมูลทั้งระบบตั้งแต่การกำหนดเจ้าของข้อมูล "
        "การควบคุมคุณภาพ จนถึงนโยบายการจัดเก็บ ส่วนที่สองคือโครงสร้าง Master Data (หัวข้อ 3.2.2) "
        "ที่กำหนดให้ทะเบียนสัตว์ (Animal Registry) เป็น Single Source of Truth "
        "ส่วนที่สามคือแบบจำลองข้อมูลเชิงแนวคิด (หัวข้อ 3.2.3) ที่แสดง Entity-Relationship Diagram "
        "ครอบคลุมทุกเอนทิตีทั้งระบบเดิมและระบบใหม่ ส่วนที่สี่คือโครงสร้างพจนานุกรมข้อมูล "
        "(หัวข้อ 3.2.4) ตาม TOR ข้อ 4.26 ส่วนที่ห้าคือกลยุทธ์การบูรณาการข้อมูลกับระบบเดิม "
        "(หัวข้อ 3.2.5) ที่ออกแบบให้โมดูลใหม่เข้าถึงข้อมูลเดิมผ่าน Read-Only API เท่านั้น "
        "ส่วนที่หกคือกลไกป้องกันความซ้ำซ้อนและควบคุมคุณภาพข้อมูล (หัวข้อ 3.2.6) "
        "ส่วนที่เจ็ดคือ Data Lineage และ Audit Trail (หัวข้อ 3.2.7) ตาม TOR ข้อ 4.21 "
        "และส่วนที่แปดคือนโยบายวงจรชีวิตข้อมูลและการจัดเก็บ (หัวข้อ 3.2.8) "
        "ตารางที่ 3.2-1 ถึง 3.2-9 แสดงรายละเอียดแต่ละประเด็น "
        "และแผนภาพที่ 3.2-1 ถึง 3.2-6 แสดงภาพรวมสถาปัตยกรรมข้อมูลและกระบวนการที่เกี่ยวข้อง"
    )

    # ─────────────────────────────────────────────
    # 3.2.1  Data Governance Framework
    # ─────────────────────────────────────────────
    add_section_heading(doc, "3.2.1", "หลักการธรรมาภิบาลข้อมูล (Data Governance Framework)")

    add_body_after_section(doc,
        "กรอบธรรมาภิบาลข้อมูล (Data Governance Framework) เป็นหัวใจสำคัญในการกำกับดูแลข้อมูล"
        "ทั้งระบบให้มีคุณภาพ ความถูกต้อง ความปลอดภัย และความสอดคล้องกันตลอดวงจรชีวิตของข้อมูล "
        "ผู้เสนอได้ออกแบบกรอบธรรมาภิบาลข้อมูลสำหรับระบบ Thai Zoo ARK โดยยึดหลักการ 6 ประการ"
        "ที่ครอบคลุมทุกมิติของการจัดการข้อมูล ตั้งแต่การกำหนดเจ้าของข้อมูลที่ชัดเจน "
        "การรักษาคุณภาพผ่านกระบวนการตรวจสอบอัตโนมัติ การควบคุมการเข้าถึงตามบทบาทหน้าที่ "
        "การติดตามที่มาของข้อมูลผ่าน Audit Trail การกำหนดนโยบายวงจรชีวิตข้อมูล "
        "และการจัดทำมาตรฐานข้อมูลกลางผ่านพจนานุกรมข้อมูล (Data Dictionary) ตาม TOR ข้อ 4.26 "
        "กรอบนี้ออกแบบให้สอดคล้องกับข้อกำหนด TOR ที่ระบุว่าระบบต้องทำงานร่วมกับระบบเดิม"
        "โดยไม่กระทบโครงสร้าง และใช้ข้อมูลจากระบบเดิมอย่างเหมาะสมไม่ซ้ำซ้อน "
        "ตารางที่ 3.2-1 แสดงรายละเอียดหลักการทั้ง 6 ประการ "
        "และแผนภาพที่ 3.2-1 แสดงภาพรวมกรอบธรรมาภิบาลข้อมูลทั้งระบบ"
    )

    # Table 3.2-1
    add_table_caption(doc, "ตารางที่ 3.2-1 หลักการธรรมาภิบาลข้อมูล 6 ประการของระบบ Thai Zoo ARK")
    add_table(doc,
        ["ลำดับ", "หลักการ", "คำอธิบาย", "ตัวอย่างการประยุกต์ใช้", "TOR ที่เกี่ยวข้อง"],
        [
            ["1", "Data Ownership\n(ความเป็นเจ้าของข้อมูล)",
             "กำหนดเจ้าของข้อมูลแต่ละ Domain อย่างชัดเจน เช่น ทะเบียนสัตว์มีเจ้าของคือ Zoo Admin "
             "เวชระเบียนมีเจ้าของคือสัตวแพทย์ เพื่อให้มีผู้รับผิดชอบคุณภาพและความถูกต้อง",
             "Animal Registry → Zoo Admin\nMedical Records → สัตวแพทย์\nStudbook → เจ้าหน้าที่ Studbook",
             "4.5, 4.6"],
            ["2", "Data Quality\n(คุณภาพข้อมูล)",
             "บังคับใช้กฎการตรวจสอบข้อมูล (Validation Rules) ทุกจุดนำเข้า ตั้งแต่ระดับ Schema "
             "(ชนิดข้อมูล รูปแบบ) ระดับ Business Rule (ช่วงค่า ตรรกะ) และระดับ Cross-reference "
             "(ความสมบูรณ์ของ Foreign Key)",
             "วันเกิดสัตว์ไม่เป็นอนาคต\nอาหารต้องอยู่ใน Diet Card\nAnimal ID ต้องไม่ซ้ำ",
             "4.6, 4.7, 4.14"],
            ["3", "Data Security & Access Control\n(ความปลอดภัยและสิทธิ์)",
             "ควบคุมการเข้าถึงข้อมูลตามบทบาทหน้าที่ (RBAC) ร่วมกับสิทธิ์รายบุคคล "
             "การลบข้อมูลต้องมีสิทธิ์ Super Admin และอ้างอิงใบคำร้อง ทุกการกระทำถูกบันทึก",
             "Super Admin เท่านั้นลบทะเบียนสัตว์\nสัตวแพทย์เข้าถึงเฉพาะเวชระเบียน\nKeeper เห็นเฉพาะสัตว์ที่ตนดูแล",
             "4.5, 4.6, 4.21"],
            ["4", "Data Lineage & Audit Trail\n(การติดตามที่มาข้อมูล)",
             "ทุก Record มีข้อมูลติดตามที่มาครบถ้วน ได้แก่ ผู้สร้าง ผู้แก้ไข วันเวลาสร้าง/แก้ไข "
             "และ Version History เพื่อให้สามารถตรวจสอบย้อนหลังได้ตลอด",
             "created_by, updated_by\ncreated_at, updated_at\nversion_no, change_log",
             "4.21, 4.18.7.1"],
            ["5", "Data Lifecycle Management\n(วงจรชีวิตข้อมูล)",
             "กำหนดนโยบายระยะเวลาจัดเก็บข้อมูลแต่ละประเภท ตั้งแต่ข้อมูลปฏิบัติการ "
             "ข้อมูล Archive จนถึง System Log ที่ต้องเก็บอย่างน้อย 90 วัน ตาม TOR",
             "System Log ≥ 90 วัน\nเวชระเบียน: ตลอดชีวิตสัตว์\nBilling: ตามกฎหมาย",
             "4.21.5"],
            ["6", "Data Dictionary & Standards\n(พจนานุกรมข้อมูลและมาตรฐาน)",
             "จัดทำพจนานุกรมข้อมูล (Data Dictionary) สำหรับทุกตารางและฟิลด์ "
             "กำหนด Code Lists มาตรฐาน (Species Taxonomy, Procedure Codes, Zoo Codes) "
             "และส่งมอบเป็นเอกสารแยกตาม TOR ข้อ 4.26",
             "Data Dictionary Document\nSpecies Code List\nProcedure Code List",
             "4.26"],
        ],
    )

    # Diagram 3.2-1
    fig = create_governance_framework_diagram()
    add_image_from_fig(doc, fig, width_inches=5.8)
    add_figure_caption(doc, "แผนภาพที่ 3.2-1 กรอบธรรมาภิบาลข้อมูลของระบบ Thai Zoo ARK")

    # ─────────────────────────────────────────────
    # 3.2.2  Master Data & Single Source of Truth
    # ─────────────────────────────────────────────
    add_section_heading(doc, "3.2.2", "โครงสร้าง Master Data และหลัก Single Source of Truth")

    add_body_after_section(doc,
        "หลัก Single Source of Truth (SSOT) เป็นแนวคิดสำคัญในการออกแบบระบบ Thai Zoo ARK "
        "โดยกำหนดให้ข้อมูลแต่ละประเภทมีแหล่งเก็บหลักเพียงแห่งเดียว และโมดูลอื่นทั้งหมด"
        "ต้องอ้างอิงข้อมูลจากแหล่งหลักผ่าน Foreign Key หรือ API เท่านั้น "
        "ไม่อนุญาตให้คัดลอกข้อมูลซ้ำไปเก็บในตารางอื่น (No Duplication Rule) "
        "แนวทางนี้ตอบโจทย์ TOR โดยตรงที่ระบุว่าระบบต้อง \"ใช้ข้อมูลจากระบบเดิมอย่างเหมาะสม "
        "ไม่ซ้ำซ้อน/ขัดแย้ง\" ซึ่งเป็นเกณฑ์ที่มีน้ำหนัก 30 คะแนน "
        "โครงสร้าง Master Data ของระบบ Thai Zoo ARK มีทะเบียนสัตว์ (Animal Registry) "
        "เป็นศูนย์กลาง โดยทุกโมดูลทั้งระบบเดิม (เวชระเบียน โภชนาการ ผู้ดูแลสัตว์) "
        "และระบบใหม่ (Studbook, Genetic Analysis, Pedigree, Billing, Dashboard) "
        "ล้วนอ้างอิงกลับมาที่ animal_id เป็น Primary Identifier "
        "ข้อมูลอ้างอิง (Reference Data) เช่น รหัสชนิดสัตว์ รหัสสวนสัตว์ รหัสหัตถการ "
        "จะถูกจัดเก็บเป็น Code Lists ที่ใช้ร่วมกันทั้งระบบ "
        "ตารางที่ 3.2-2 แสดงรายการ Master Data หลักและเจ้าของข้อมูล "
        "และแผนภาพที่ 3.2-2 แสดงโครงสร้างความสัมพันธ์ระหว่าง Master Data กับโมดูลต่างๆ"
    )

    # Table 3.2-2
    add_table_caption(doc, "ตารางที่ 3.2-2 รายการ Master Data หลักของระบบ Thai Zoo ARK")
    add_table(doc,
        ["ลำดับ", "Master Data", "คำอธิบาย", "Primary Key", "Owner Module",
         "สถานะ", "ผู้บริโภคข้อมูล (Data Consumers)"],
        [
            ["1", "Animal Registry\n(ทะเบียนสัตว์)", "ข้อมูลอัตลักษณ์ สถานะ และลำดับวงศ์ตระกูล"
             "ของสัตว์ทุกตัว เป็น SSOT ของข้อมูลสัตว์",
             "animal_id", "ทะเบียนสัตว์", "ตารางเดิม\n(ไม่แก้ไข)",
             "ทุกโมดูล"],
            ["2", "Species Taxonomy\n(อนุกรมวิธาน)", "ข้อมูลชนิดสัตว์ ชื่อวิทยาศาสตร์ "
             "ชื่อสามัญ Class Order Family",
             "species_id", "ทะเบียนสัตว์", "ตารางเดิม\n(ไม่แก้ไข)",
             "Studbook, Pedigree, Dashboard"],
            ["3", "Zoo / Location\n(สวนสัตว์/สถานที่)", "ข้อมูลสวนสัตว์ในเครือ อสส. "
             "7 แห่ง พร้อมพิกัดและข้อมูลติดต่อ",
             "zoo_id", "ระบบบริหาร", "ตารางเดิม\n(ไม่แก้ไข)",
             "ทุกโมดูล"],
            ["4", "User / Staff\n(เจ้าหน้าที่)", "ข้อมูลผู้ใช้งาน บทบาท สิทธิ์ "
             "และโครงสร้างองค์กร",
             "user_id", "จัดการเจ้าหน้าที่", "ตารางเดิม\n(เพิ่ม Permission)",
             "System Log, Audit Trail"],
            ["5", "Procedure Code\n(รหัสหัตถการ)", "รหัสมาตรฐานหัตถการทางสัตวแพทย์ "
             "พร้อมอัตราค่าบริการ",
             "procedure_code", "Billing (ใหม่)", "ตารางใหม่",
             "Medical Record, Invoice"],
            ["6", "Drug / Medicine\n(ยา/เวชภัณฑ์)", "รายการยาและเวชภัณฑ์ทั้งหมด "
             "พร้อมแหล่งที่มา ล็อต วันหมดอายุ",
             "drug_id", "ยา/เวชภัณฑ์", "ตารางเดิม\n(เพิ่ม Source)",
             "Billing, Medical Record"],
            ["7", "Food Item\n(วัตถุดิบอาหาร)", "รายการวัตถุดิบอาหารสัตว์ "
             "หน่วยนับ ประเภท",
             "food_item_id", "คลังอาหาร", "ตารางเดิม\n(ไม่แก้ไข)",
             "Nutrition, Keeper"],
        ],
    )

    # Diagram 3.2-2
    fig = create_master_data_diagram()
    add_image_from_fig(doc, fig, width_inches=5.8)
    add_figure_caption(doc, "แผนภาพที่ 3.2-2 โครงสร้าง Master Data และความสัมพันธ์ — Animal ID เป็น Single Source of Truth")

    # ─────────────────────────────────────────────
    # 3.2.3  Conceptual Data Model (ER Diagram)
    # ─────────────────────────────────────────────
    add_section_heading(doc, "3.2.3", "แบบจำลองข้อมูลเชิงแนวคิด (Conceptual Data Model)")

    add_body_after_section(doc,
        "แบบจำลองข้อมูลเชิงแนวคิด (Conceptual Data Model) แสดงภาพรวมเอนทิตีหลักทั้งหมด"
        "ของระบบ Thai Zoo ARK พร้อมความสัมพันธ์ระหว่างเอนทิตี ซึ่งออกแบบตามหลัก "
        "Additive Schema Changes โดยตารางเดิมไม่ถูกแก้ไขโครงสร้าง (No Breaking Changes) "
        "ในขณะที่ตารางใหม่สำหรับโมดูล Studbook, Genetic Analysis, Pedigree, Billing, "
        "Change Request, Manual Portal และ System Log จะถูกเพิ่มเข้ามาและเชื่อมโยง"
        "กับตารางเดิมผ่าน Foreign Key ในทิศทาง Read-Only เท่านั้น "
        "การออกแบบนี้มีจุดเด่นสำคัญ 4 ประการ ประการแรกคือเอนทิตี Animal เป็นศูนย์กลาง"
        "ของทุกความสัมพันธ์ ประการที่สองคือความสัมพันธ์แบบ Self-referencing "
        "สำหรับข้อมูลสายพันธุ์ (Sire/Dam) ซึ่งจำเป็นสำหรับ Studbook และ Pedigree "
        "ประการที่สามคือการแยก System Log ออกจาก Business Data อย่างชัดเจน "
        "และประการที่สี่คือการออกแบบ Genetic Analysis ให้เป็น Calculated Entity "
        "ที่คำนวณค่า F-coefficient จากข้อมูล Studbook โดยอัตโนมัติ "
        "แผนภาพที่ 3.2-3 แสดง Entity-Relationship Diagram ภาพรวมทั้งระบบ "
        "และตารางที่ 3.2-3 สรุปเอนทิตีหลักทั้งหมดพร้อมประเภทและ Cardinality"
    )

    # Diagram 3.2-3
    fig = create_er_diagram()
    add_image_from_fig(doc, fig, width_inches=6.0)
    add_figure_caption(doc, "แผนภาพที่ 3.2-3 Entity-Relationship Diagram ภาพรวมระบบ Thai Zoo ARK")

    # Table 3.2-3
    add_table_caption(doc, "ตารางที่ 3.2-3 สรุปเอนทิตีหลักและความสัมพันธ์ของระบบ Thai Zoo ARK")
    add_table(doc,
        ["เอนทิตี", "ประเภท", "Primary Key", "ความสัมพันธ์หลัก", "Cardinality", "หมายเหตุ"],
        [
            ["Animal\n(ทะเบียนสัตว์)", "เดิม", "animal_id",
             "→ Species (FK)\n→ Zoo (FK)\n→ Self (sire/dam)", "1:N\n1:N\nSelf-ref",
             "SSOT ไม่แก้ไข Schema"],
            ["Species\n(อนุกรมวิธาน)", "เดิม", "species_id",
             "← Animal\n← Studbook", "1:N\n1:N", "Code List กลาง"],
            ["Zoo\n(สวนสัตว์)", "เดิม", "zoo_id",
             "← Animal\n← User", "1:N\n1:N", "7 สวนสัตว์ + HQ"],
            ["Studbook\n(ทะเบียนสายพันธุ์)", "ใหม่", "studbook_no",
             "→ Animal (FK)\n→ Sire/Dam (FK)", "1:1\nN:1",
             "Auto-number ตาม DOB"],
            ["Genetic Analysis\n(วิเคราะห์พันธุกรรม)", "ใหม่", "id",
             "→ Studbook (FK)", "N:1", "Calculated Entity\nF-coefficient"],
            ["Pedigree Cache\n(แคชสายโลหิต)", "ใหม่", "id",
             "→ Animal (FK)", "N:1", "Auto-generated\nAncestor Path"],
            ["Medical Record\n(เวชระเบียน)", "เดิม", "visit_id",
             "→ Animal (FK)\n→ Drug (FK)", "N:1\nN:M", "ไม่แก้ไข Schema"],
            ["Billing Invoice\n(ค่ารักษา)", "ใหม่", "invoice_id",
             "→ Animal (FK)\n→ Medical (FK)\n→ Rate Card (FK)", "N:1\n1:1\nN:1",
             "Sum Check\nAudit Trail"],
            ["Rate Card\n(อัตราค่าบริการ)", "ใหม่", "procedure_code",
             "← Billing", "1:N", "Effective Date\nEnd Date"],
            ["System Log\n(บันทึกเหตุการณ์)", "ใหม่", "id",
             "→ User (FK)", "N:1", "Partitioned\n≥ 90 วัน"],
            ["User\n(เจ้าหน้าที่)", "เดิม\n(เพิ่ม Permission)", "user_id",
             "→ Zoo (FK)\n→ Role (FK)", "N:1\nN:1", "เพิ่ม Individual\nPermission"],
        ],
    )

    # ─────────────────────────────────────────────
    # 3.2.4  Data Dictionary Structure
    # ─────────────────────────────────────────────
    add_section_heading(doc, "3.2.4", "โครงสร้างพจนานุกรมข้อมูล (Data Dictionary Structure)")

    add_body_after_section(doc,
        "พจนานุกรมข้อมูล (Data Dictionary) เป็นเอกสารสำคัญที่ TOR ข้อ 4.26 กำหนดให้ต้องส่งมอบ "
        "โดยผู้เสนอจะจัดทำพจนานุกรมข้อมูลสำหรับทุกตารางในระบบ ทั้งตารางเดิมที่มีอยู่แล้ว"
        "และตารางใหม่ที่พัฒนาเพิ่ม โดยพจนานุกรมข้อมูลจะระบุรายละเอียดในระดับฟิลด์ "
        "ครอบคลุมชื่อฟิลด์ ชนิดข้อมูล ขนาด ข้อจำกัด (Constraints) ค่าเริ่มต้น "
        "คำอธิบาย กฎการตรวจสอบ (Validation Rules) และหมายเหตุการ Migration "
        "ตารางที่ 3.2-4 แสดงรูปแบบมาตรฐานของพจนานุกรมข้อมูลที่จะใช้ "
        "และตารางที่ 3.2-5 แสดงตัวอย่างพจนานุกรมข้อมูลสำหรับตาราง studbook "
        "ซึ่งเป็นตารางใหม่ที่สำคัญที่สุดของโครงการ "
        "พจนานุกรมข้อมูลฉบับสมบูรณ์จะถูกส่งมอบเป็นเอกสารแยกตาม TOR ข้อ 4.26"
    )

    # Table 3.2-4
    add_table_caption(doc, "ตารางที่ 3.2-4 โครงสร้างพจนานุกรมข้อมูลของตารางหลักในระบบ")
    add_table(doc,
        ["ตาราง (Table)", "ฟิลด์หลัก (Key Fields)", "Owner Module",
         "กฎการตรวจสอบ (Validation Rules)", "หมายเหตุ Migration"],
        [
            ["animal\n(ตารางเดิม)", "animal_id, species_id,\nname, sex, dob, status",
             "ทะเบียนสัตว์", "ไม่แก้ไข Schema เดิม\nRead-only จากโมดูลใหม่",
             "ไม่มีการเปลี่ยนแปลง"],
            ["studbook\n(ตารางใหม่)", "studbook_no, animal_id,\nspecies_id, sire_id,\ndam_id, dob, dod",
             "Studbook", "FK → animal (must exist)\nAuto-number ตาม DOB\nสถานะครอบคลุม TOR 4.14.2",
             "ตารางใหม่\nสร้างพร้อม Index"],
            ["genetic_analysis\n(ตารางใหม่)", "id, studbook_id,\nf_coefficient,\ncompleteness_index",
             "Genetic", "Calculated field ≥ 0\nFK → studbook\nRecalculate on parent update",
             "ตารางใหม่\nTrigger-based"],
            ["pedigree_cache\n(ตารางใหม่)", "id, animal_id,\nancestor_path, depth",
             "Pedigree", "Auto-generated\nFK → animal\nRefresh on data change",
             "ตารางใหม่\nMaterialized View"],
            ["billing_invoice\n(ตารางใหม่)", "invoice_id, animal_id,\ntotal_procedure,\ntotal_medicine, status",
             "Billing", "Sum check\nFK → animal, medical\nStatus workflow",
             "ตารางใหม่\nFull Audit Trail"],
            ["rate_card\n(ตารางใหม่)", "procedure_code, rate,\nunit, effective_date,\nend_date",
             "Billing", "Date range validation\nNo overlap\nChange history (TOR 4.18.4.2)",
             "ตารางใหม่\nVersioned"],
            ["change_request\n(ตารางใหม่)", "cr_id, category,\nrequester_id,\nstatus, attachment",
             "คำร้อง", "หมวดหมู่ตาม TOR 4.19.1\nPDF attachment\nStatus workflow",
             "ตารางใหม่"],
            ["system_log\n(ตารางใหม่)", "id, timestamp,\nuser_id, action,\nmodule, detail, ip",
             "System Log", "Retention ≥ 90 วัน (TOR 4.21.5)\nPartitioned by month\nImmutable (append-only)",
             "ตารางใหม่\nPartitioned"],
        ],
    )

    # Table 3.2-5 — Detailed example for studbook
    add_table_caption(doc, "ตารางที่ 3.2-5 ตัวอย่างพจนานุกรมข้อมูลตาราง studbook (รายละเอียดระดับฟิลด์)")
    add_table(doc,
        ["ฟิลด์", "ชนิดข้อมูล", "ขนาด", "Nullable", "Constraint", "คำอธิบาย", "Validation Rule"],
        [
            ["studbook_no", "BIGINT", "-", "NOT NULL", "PK, AUTO_INCREMENT",
             "ลำดับในทะเบียนสายพันธุ์\nเรียงตาม DOB (TOR 4.14.1)", "Auto-generate ตาม DOB"],
            ["animal_id", "INT", "-", "NOT NULL", "FK → animal.animal_id\nUNIQUE",
             "รหัสประจำตัวสัตว์\nอ้างอิง SSOT", "ต้องมีอยู่ในตาราง animal"],
            ["species_id", "INT", "-", "NOT NULL", "FK → species.species_id",
             "รหัสชนิดสัตว์", "ต้องตรงกับ animal.species_id"],
            ["sex", "ENUM", "M/F/U", "NOT NULL", "-",
             "เพศ: ตัวผู้/ตัวเมีย/ไม่ทราบ", "ต้องตรงกับ animal.sex"],
            ["dob", "DATE", "-", "NULL", "-",
             "วันเกิด", "ไม่เป็นอนาคต\nถ้าไม่ทราบใช้วันรับเข้า"],
            ["dod", "DATE", "-", "NULL", "-",
             "วันตาย", "ต้องหลัง dob\nNULL = ยังมีชีวิต"],
            ["sire_id", "INT", "-", "NULL", "FK → animal.animal_id",
             "รหัสพ่อ", "ต้องมีในตาราง animal\nเพศ = M"],
            ["dam_id", "INT", "-", "NULL", "FK → animal.animal_id",
             "รหัสแม่", "ต้องมีในตาราง animal\nเพศ = F"],
            ["origin", "VARCHAR", "100", "NULL", "-",
             "แหล่งที่มา", "เช่น เกิดในสวนสัตว์ ย้ายเข้า รับบริจาค"],
            ["current_location", "VARCHAR", "100", "NULL", "-",
             "สถานที่ปัจจุบัน", "ชื่อสวนสัตว์/ส่วนจัดแสดง"],
            ["status", "ENUM", "-", "NOT NULL", "-",
             "สถานะตาม TOR 4.14.2:\nเกิด/ตาย/ย้ายเข้า/ย้ายออก/\nสูญหาย/รับบริจาค/จัดซื้อ",
             "ต้องตรงกับ TOR 4.14.2"],
            ["created_by", "INT", "-", "NOT NULL", "FK → user.user_id",
             "ผู้สร้างรายการ", "Auto-fill จาก Session"],
            ["created_at", "DATETIME", "-", "NOT NULL", "DEFAULT NOW()",
             "วันเวลาสร้าง", "Auto-fill"],
            ["updated_by", "INT", "-", "NULL", "FK → user.user_id",
             "ผู้แก้ไขล่าสุด", "Auto-fill จาก Session"],
            ["updated_at", "DATETIME", "-", "NULL", "ON UPDATE NOW()",
             "วันเวลาแก้ไขล่าสุด", "Auto-fill"],
        ],
    )

    # ─────────────────────────────────────────────
    # 3.2.5  Data Integration Strategy
    # ─────────────────────────────────────────────
    add_section_heading(doc, "3.2.5", "กลยุทธ์การบูรณาการข้อมูลกับระบบเดิม (Data Integration Strategy)")

    add_body_after_section(doc,
        "กลยุทธ์การบูรณาการข้อมูลเป็นหัวใจของการทำให้ระบบใหม่และระบบเดิมทำงานร่วมกัน"
        "ได้อย่างราบรื่น โดยไม่กระทบต่อโครงสร้างและการทำงานเดิม "
        "ผู้เสนอได้ออกแบบกลยุทธ์การบูรณาการตามหลัก \"Read-Only Access to Existing Data\" "
        "กล่าวคือ โมดูลใหม่ทั้งหมด (Studbook, Genetic Analysis, Pedigree, Billing, "
        "Dashboard, Change Request, System Log) จะเข้าถึงข้อมูลจากตารางเดิม "
        "ผ่าน API Gateway ในโหมด Read-Only เท่านั้น โดยไม่เขียนข้อมูลกลับไปที่ตารางเดิมโดยตรง "
        "การเขียนข้อมูลจะเกิดขึ้นเฉพาะในตารางใหม่ที่ถูกออกแบบมาสำหรับโมดูลนั้นๆ "
        "โดยใช้ Foreign Key อ้างอิงไปยังตารางเดิม "
        "กลยุทธ์นี้มีข้อดี 3 ประการ ประการแรกคือไม่มีความเสี่ยงต่อข้อมูลเดิม "
        "ประการที่สองคือสามารถ Rollback ได้โดยลบตารางใหม่ออกโดยไม่กระทบตารางเดิม "
        "และประการที่สามคือสามารถพัฒนาและทดสอบแต่ละโมดูลได้อิสระโดยไม่กระทบโมดูลอื่น "
        "แผนภาพที่ 3.2-4 แสดงภาพรวมกลยุทธ์การบูรณาการ "
        "และตารางที่ 3.2-6 แสดงแนวทางการเข้าถึงข้อมูลแยกตามโมดูล"
    )

    # Diagram 3.2-4
    fig = create_integration_strategy_diagram()
    add_image_from_fig(doc, fig, width_inches=5.8)
    add_figure_caption(doc, "แผนภาพที่ 3.2-4 กลยุทธ์การบูรณาการข้อมูลระหว่างโมดูลใหม่และระบบเดิม")

    # Table 3.2-6
    add_table_caption(doc, "ตารางที่ 3.2-6 แนวทางการเข้าถึงข้อมูลแยกตามโมดูล")
    add_table(doc,
        ["โมดูลใหม่", "ตารางเดิมที่อ้างอิง", "วิธีเข้าถึง",
         "สิทธิ์เข้าถึง", "ตารางใหม่ที่เขียน", "TOR Ref"],
        [
            ["Studbook", "animal, species", "API (Read-Only)",
             "SELECT only", "studbook", "4.14"],
            ["Genetic Analysis", "animal (via studbook)", "API (Read-Only)",
             "SELECT only", "genetic_analysis,\npedigree_cache", "4.15"],
            ["Pedigree Diagram", "animal, species,\nstudbook, genetic", "API (Read-Only)",
             "SELECT only", "pedigree_cache", "4.16"],
            ["Billing", "animal,\nmedical_record, drug", "API (Read-Only)",
             "SELECT only", "billing_invoice,\nrate_card,\nbilling_payment", "4.18"],
            ["Dashboard", "animal, studbook,\ngenetic, nutrition,\nmedical", "API (Read-Only)",
             "SELECT only", "dashboard_cache\n(Materialized View)", "4.17"],
            ["Change Request", "animal, user", "API (Read-Only)",
             "SELECT only", "change_request,\ncr_attachment", "4.19"],
            ["System Log", "user", "API (Read-Only)",
             "SELECT only", "system_log", "4.21"],
            ["Manual Portal", "user", "API (Read-Only)",
             "SELECT only", "knowledge_base,\nkb_attachment", "4.20"],
        ],
    )

    # ─────────────────────────────────────────────
    # 3.2.6  Data Quality & Deduplication
    # ─────────────────────────────────────────────
    add_section_heading(doc, "3.2.6", "กลไกป้องกันความซ้ำซ้อนและการควบคุมคุณภาพข้อมูล")

    add_body_after_section(doc,
        "การป้องกันความซ้ำซ้อนของข้อมูล (Data Deduplication) และการควบคุมคุณภาพข้อมูล "
        "(Data Quality Control) เป็นกลไกสำคัญที่ตอบโจทย์เกณฑ์การให้คะแนนข้อ 3 "
        "ในประเด็น \"ใช้ข้อมูลจากระบบเดิมอย่างเหมาะสม ไม่ซ้ำซ้อน/ขัดแย้ง\" (30 คะแนน) "
        "ผู้เสนอได้ออกแบบกลไกการควบคุมคุณภาพข้อมูล 4 ระดับ ได้แก่ "
        "ระดับที่ 1 คือ Schema Validation ซึ่งตรวจสอบชนิดข้อมูล รูปแบบ และฟิลด์บังคับ "
        "ระดับที่ 2 คือ Business Rule Check ซึ่งตรวจสอบช่วงค่า ตรรกะทางธุรกิจ "
        "เช่น อาหารต้องอยู่ใน Diet Card ยาบริจาคต้องไม่คิดค่ายา "
        "ระดับที่ 3 คือ Duplicate Detection ซึ่งตรวจสอบข้อมูลซ้ำ เช่น Animal ID "
        "และ Studbook Number ต้องไม่ซ้ำ "
        "และระดับที่ 4 คือ Referential Integrity ซึ่งตรวจสอบว่า Foreign Key "
        "ที่อ้างอิงมีอยู่จริงในตารางเดิมและสถานะยังใช้งานได้ "
        "ตารางที่ 3.2-7 แสดงรายละเอียดกฎการป้องกันความซ้ำซ้อนและการควบคุมคุณภาพ "
        "และแผนภาพที่ 3.2-5 แสดงกระบวนการตรวจสอบคุณภาพข้อมูลแบบ 4 ระดับ"
    )

    # Table 3.2-7
    add_table_caption(doc, "ตารางที่ 3.2-7 กฎการป้องกันความซ้ำซ้อนและการควบคุมคุณภาพข้อมูล")
    add_table(doc,
        ["ลำดับ", "ระดับ", "กฎ (Rule)", "ตัวอย่าง", "โมดูลที่ใช้", "การดำเนินการเมื่อผิดกฎ"],
        [
            ["1", "Schema\nValidation",
             "ตรวจสอบชนิดข้อมูล รูปแบบ\nและฟิลด์บังคับ (NOT NULL)",
             "วันเกิดต้องเป็น DATE\nเพศต้องเป็น M/F/U\nanimal_id ต้องไม่ว่าง",
             "ทุกโมดูล", "Reject + แจ้ง Error\nพร้อมระบุฟิลด์ที่ผิด"],
            ["2", "Business\nRule Check",
             "ตรวจสอบช่วงค่าและตรรกะธุรกิจ",
             "วันเกิดไม่เป็นอนาคต\nอาหารต้องอยู่ใน Diet Card\nยาบริจาค → ไม่คิดค่ายา (4.18.3.1)\nSire เพศ = M, Dam เพศ = F",
             "Nutrition, Billing,\nStudbook, Registry",
             "Reject + แจ้ง Warning\nพร้อมคำแนะนำ"],
            ["3", "Duplicate\nDetection",
             "ตรวจสอบข้อมูลซ้ำด้วย\nUnique Constraints",
             "Animal ID ไม่ซ้ำ\nStudbook No ไม่ซ้ำ\nInvoice No ไม่ซ้ำ\nDrug Barcode ไม่ซ้ำ",
             "Registry, Studbook,\nBilling, Pharmacy",
             "Reject + แสดง\nรายการที่ซ้ำ"],
            ["4", "Referential\nIntegrity",
             "ตรวจสอบ FK ว่ามีอยู่จริง\nในตารางเดิมและยัง Active",
             "animal_id ต้องมีในตาราง animal\nspecies_id ต้องมีในตาราง species\nuser_id ต้องมีในตาราง user\nstatus ต้อง Active",
             "ทุกโมดูลใหม่",
             "Reject + แจ้งว่า\nข้อมูลอ้างอิงไม่ถูกต้อง"],
            ["5", "No Duplication\nRule",
             "ข้อมูลสัตว์เก็บที่ทะเบียน\nเดียว โมดูลอื่นอ้างอิง\nผ่าน FK ไม่คัดลอก",
             "Studbook ไม่คัดลอกชื่อสัตว์\nBilling ไม่คัดลอก diagnosis\nDashboard ใช้ View/API",
             "ทุกโมดูลใหม่",
             "Design enforcement\n(ระดับ Architecture)"],
            ["6", "Cross-Module\nConsistency",
             "ข้อมูลข้ามโมดูลต้องสอดคล้อง",
             "สถานะสัตว์ใน Studbook\nตรงกับ Registry\nค่ายาใน Billing ตรงกับ Pharmacy",
             "Studbook, Billing,\nDashboard",
             "Periodic consistency\ncheck + Alert"],
        ],
    )

    # Diagram 3.2-5
    fig = create_data_quality_diagram()
    add_image_from_fig(doc, fig, width_inches=5.8)
    add_figure_caption(doc, "แผนภาพที่ 3.2-5 กระบวนการควบคุมคุณภาพข้อมูลแบบ 4 ระดับ")

    # ─────────────────────────────────────────────
    # 3.2.7  Data Lineage & Audit Trail
    # ─────────────────────────────────────────────
    add_section_heading(doc, "3.2.7", "Data Lineage และ Audit Trail")

    add_body_after_section(doc,
        "การติดตามที่มาของข้อมูล (Data Lineage) และร่องรอยการตรวจสอบ (Audit Trail) "
        "เป็นกลไกสำคัญที่ TOR กำหนดไว้ในหลายข้อ ทั้ง TOR ข้อ 4.21 (System Log) "
        "ข้อ 4.6 (ควบคุมการลบข้อมูลทะเบียนสัตว์) และข้อ 4.18.7.1 (Audit Trail ของ Billing) "
        "ผู้เสนอได้ออกแบบระบบ Audit Trail แบบครอบคลุม 2 ระดับ "
        "ระดับแรกคือ Record-level Audit ซึ่งทุกตารางในระบบทั้งเดิมและใหม่จะมีฟิลด์ Audit "
        "ประกอบด้วย created_by, created_at, updated_by, updated_at, version_no "
        "เพื่อให้สามารถตรวจสอบได้ว่าใครสร้าง ใครแก้ไข และเมื่อไหร่ "
        "ระดับที่สองคือ System Log ซึ่งเป็นตาราง Centralized Log ที่บันทึกทุกเหตุการณ์สำคัญ "
        "ครอบคลุม Login/Logout, CRUD Actions, Timestamp, User ID, IP Address, Module "
        "และรายละเอียดการเปลี่ยนแปลง (before/after values) "
        "ระบบ System Log ออกแบบเป็น Append-only (ไม่สามารถแก้ไขหรือลบได้) "
        "และมี Retention Policy อย่างน้อย 90 วันตาม TOR ข้อ 4.21.5 โดยตั้งค่าได้ "
        "ตารางที่ 3.2-8 แสดงฟิลด์ Audit ที่บังคับในทุกตาราง "
        "และแผนภาพที่ 3.2-6 แสดง Data Lineage Flow ทั้งระบบ"
    )

    # Table 3.2-8
    add_table_caption(doc, "ตารางที่ 3.2-8 ฟิลด์ Audit ที่บังคับในทุกตารางของระบบ")
    add_table(doc,
        ["ฟิลด์", "ชนิดข้อมูล", "คำอธิบาย", "ตัวอย่างค่า",
         "วิธีการบันทึก", "ระดับการใช้งาน"],
        [
            ["created_by", "INT (FK → user)", "ผู้สร้างรายการ",
             "user_id = 1042", "Auto-fill จาก Session", "ทุกตาราง"],
            ["created_at", "DATETIME", "วันเวลาที่สร้าง",
             "2569-03-15 10:30:00", "DEFAULT CURRENT_TIMESTAMP", "ทุกตาราง"],
            ["updated_by", "INT (FK → user)", "ผู้แก้ไขล่าสุด",
             "user_id = 1055", "Auto-fill จาก Session", "ทุกตาราง"],
            ["updated_at", "DATETIME", "วันเวลาที่แก้ไขล่าสุด",
             "2569-03-16 14:20:00", "ON UPDATE CURRENT_TIMESTAMP", "ทุกตาราง"],
            ["deleted_by", "INT (FK → user)", "ผู้ลบรายการ (Soft Delete)",
             "user_id = 1001\n(Super Admin)", "Manual (Super Admin only)", "ตารางที่ลบได้"],
            ["deleted_at", "DATETIME", "วันเวลาที่ลบ",
             "2569-04-01 09:00:00", "Manual entry", "ตารางที่ลบได้"],
            ["deletion_reason", "TEXT", "เหตุผลการลบ",
             "\"สัตว์ตายแล้ว ตาม CR#001\"", "ต้องกรอก เมื่อลบ", "ทะเบียนสัตว์"],
            ["cr_reference", "VARCHAR(50)", "เลขที่ใบคำร้องอ้างอิง",
             "CR-2569-0042", "ต้องกรอก เมื่อลบ\n(TOR 4.6)", "ทะเบียนสัตว์"],
            ["version_no", "INT", "ลำดับเวอร์ชันของ Record",
             "3", "Auto-increment เมื่อแก้ไข", "ตารางสำคัญ"],
            ["source_module", "VARCHAR(50)", "โมดูลที่เป็นต้นทางข้อมูล",
             "\"BILLING\"", "Auto-fill ตามโมดูล", "System Log"],
            ["ip_address", "VARCHAR(45)", "IP Address ผู้ใช้งาน",
             "10.20.30.40", "Auto-capture", "System Log"],
        ],
    )

    # Diagram 3.2-6
    fig = create_data_lineage_diagram()
    add_image_from_fig(doc, fig, width_inches=5.8)
    add_figure_caption(doc, "แผนภาพที่ 3.2-6 Data Lineage และ Audit Trail Flow ของระบบ Thai Zoo ARK")

    # ─────────────────────────────────────────────
    # 3.2.8  Data Retention & Lifecycle Policy
    # ─────────────────────────────────────────────
    add_section_heading(doc, "3.2.8", "นโยบายวงจรชีวิตข้อมูลและการจัดเก็บ (Data Retention & Lifecycle Policy)")

    add_body_after_section(doc,
        "นโยบายวงจรชีวิตข้อมูล (Data Lifecycle Policy) กำหนดแนวทางการจัดการข้อมูล"
        "ตั้งแต่การสร้าง การใช้งาน การเก็บรักษา จนถึงการทำลาย "
        "โดยคำนึงถึงข้อกำหนดทาง TOR ข้อ 4.21.5 ที่ระบุว่า System Log ต้องเก็บ"
        "อย่างน้อย 90 วัน และตั้งค่าระยะเวลาจัดเก็บได้ "
        "ผู้เสนอได้ออกแบบนโยบายการจัดเก็บข้อมูลแยกตามประเภท โดยข้อมูลหลัก (Master Data) "
        "เช่น ทะเบียนสัตว์ อนุกรมวิธาน จะถูกเก็บรักษาถาวร เนื่องจากเป็นข้อมูลวิทยาศาสตร์"
        "ที่มีคุณค่าต่อการอนุรักษ์และการวิจัยระยะยาว "
        "ข้อมูลปฏิบัติการ (Operational Data) เช่น เวชระเบียน โภชนาการ จะถูกเก็บรักษา"
        "ตลอดชีวิตของสัตว์ และ Archive หลังจากสัตว์เสียชีวิตหรือย้ายออก "
        "ข้อมูลทรานแซกชัน (Transaction Data) เช่น Billing จะถูกเก็บตามระเบียบการเงิน "
        "และ System Log จะถูกเก็บอย่างน้อย 90 วัน โดยสามารถตั้งค่าได้ถึง 365 วัน "
        "ตารางที่ 3.2-9 แสดงรายละเอียดนโยบายระยะเวลาจัดเก็บข้อมูลแยกตามประเภท"
    )

    # Table 3.2-9
    add_table_caption(doc, "ตารางที่ 3.2-9 นโยบายระยะเวลาจัดเก็บข้อมูลแยกตามประเภท")
    add_table(doc,
        ["ประเภทข้อมูล", "ตัวอย่าง", "ระยะเวลาจัดเก็บ",
         "วิธีจัดเก็บ", "การทำลาย/Archive", "TOR Ref"],
        [
            ["Master Data\n(ข้อมูลหลัก)", "ทะเบียนสัตว์, อนุกรมวิธาน,\nสวนสัตว์, เจ้าหน้าที่",
             "ถาวร (Permanent)", "Database + Backup",
             "ไม่ทำลาย\nSoft Delete เท่านั้น", "4.6"],
            ["Conservation Data\n(ข้อมูลอนุรักษ์)", "Studbook, Genetic Analysis,\nPedigree",
             "ถาวร (Permanent)", "Database + Backup\n+ Export to ZIMS",
             "ไม่ทำลาย\nExport สำรอง", "4.14-4.16"],
            ["Operational Data\n(ข้อมูลปฏิบัติการ)", "เวชระเบียน, โภชนาการ,\nผู้ดูแลสัตว์",
             "ตลอดชีวิตสัตว์\n+ 5 ปีหลังตาย/ย้าย", "Database (Active)\n→ Archive (Cold)",
             "Archive หลัง 5 ปี\nไม่ทำลาย", "-"],
            ["Transaction Data\n(ข้อมูลทรานแซกชัน)", "Billing Invoice,\nPayment, Rate Card",
             "7 ปีตามระเบียบการเงิน", "Database (Active)\n→ Archive (Cold)",
             "Archive หลัง 7 ปี", "4.18"],
            ["System Log\n(บันทึกเหตุการณ์)", "Login/Logout, CRUD,\nUser Actions",
             "ขั้นต่ำ 90 วัน\nตั้งค่าได้ถึง 365 วัน", "Partitioned Table\n(Monthly)",
             "Auto-purge ตาม\nRetention Setting", "4.21.5"],
            ["Change Request\n(คำร้อง)", "CR Records,\nPDF Attachments",
             "5 ปีหลังปิดคำร้อง", "Database + File Storage",
             "Archive หลัง 5 ปี", "4.19"],
            ["Knowledge Base\n(คู่มือ)", "PDF, วิดีโอ, เนื้อหา",
             "ถาวร (จนกว่าจะถูกแทนที่)", "Object Storage\n(File Server)",
             "Replace เมื่อมี\nเวอร์ชันใหม่", "4.20"],
            ["Backup\n(สำรองข้อมูล)", "Full DB Backup,\nIncremental Backup",
             "Daily: 30 วัน\nWeekly: 90 วัน\nMonthly: 1 ปี",
             "Encrypted Storage\n(Off-site)", "Rotate ตามรอบ", "-"],
        ],
    )

    # ── Save Document ──
    os.makedirs(OUT_DIR, exist_ok=True)
    doc.save(OUT_FILE)
    print(f"Document saved to: {OUT_FILE}")
    return OUT_FILE


if __name__ == "__main__":
    build_document()
