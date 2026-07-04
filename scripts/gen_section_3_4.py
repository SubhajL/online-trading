#!/usr/bin/env python3
"""Generate Section 3.4 Logging & Audit for Thai Zoo ARK proposal.

Reuses helper infrastructure from gen_section_3_2.py.
"""

import os, sys
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from matplotlib.patches import FancyBboxPatch
from io import BytesIO

# ── Thai Font Registration ──
_FONT_DIR = os.path.expanduser("~/Library/Fonts/")
for _fname in [
    "THSarabunNew.ttf", "THSarabunNew Bold.ttf",
    "THSarabunNew Italic.ttf", "THSarabunNew BoldItalic.ttf",
    "THSarabun.ttf", "THSarabun Bold.ttf",
]:
    _path = os.path.join(_FONT_DIR, _fname)
    if os.path.exists(_path):
        fm.fontManager.addfont(_path)
plt.rcParams['font.family'] = 'TH Sarabun New'
plt.rcParams['font.sans-serif'] = ['TH Sarabun New', 'TH SarabunPSK', 'DejaVu Sans']

# ── Constants ──
FONT_TH = "TH Sarabun New"
BODY_SIZE = Pt(16)
HEADING_SIZE = Pt(16)
CAPTION_SIZE = Pt(14)
HEADER_BG = "D6E3F8"
OUT_DIR = os.path.expanduser("~/Library/CloudStorage/OneDrive-Personal/Personal/ZTL")
OUT_FILE = os.path.join(OUT_DIR, "Section_3_4_Logging_and_Audit.docx")

# ═══════════════════════════════════════════════════════════
# HELPERS  (identical to gen_section_3_2.py)
# ═══════════════════════════════════════════════════════════

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _ensure_cs(run, font_name, size, bold):
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="{font_name}" w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
        rPr.insert(0, rF)
    else:
        for a in ('cs','ascii','hAnsi'):
            rF.set(qn(f'w:{a}'), font_name)
    sz = rPr.find(qn('w:szCs'))
    v = str(int(size.pt * 2))
    if sz is None:
        rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="{v}"/>'))
    else:
        sz.set(qn('w:val'), v)
    if bold and rPr.find(qn('w:bCs')) is None:
        rPr.append(parse_xml(f'<w:bCs {nsdecls("w")}/>'))

def sfont(run, font_name=FONT_TH, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = font_name; run.font.size = size
    run.bold = bold; run.italic = italic
    _ensure_cs(run, font_name, size, bold)
    if color: run.font.color.rgb = color

def _heading(doc, num, title, tab, left, hang, bef, aft):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="{bef}" w:after="{aft}"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="{left}" w:hanging="{hang}"/>'))
    pPr.append(parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="{tab}"/></w:tabs>'))
    run = p.add_run(f"{num}\t{title}")
    sfont(run, bold=True, size=HEADING_SIZE)
    return p

def h_chapter(doc, n, t):   return _heading(doc, n, t, 567,  567,  567, 240, 120)
def h_section(doc, n, t):   return _heading(doc, n, t, 1418, 1418, 851, 200, 80)
def h_subsection(doc, n, t): return _heading(doc, n, t, 2268, 2268, 850, 160, 80)

def _body(doc, text, fi):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="120" w:line="240" w:lineRule="auto"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="{fi}"/>'))
    run = p.add_run(text); sfont(run)
    return p

def body_ch(doc, t):  return _body(doc, t, 567)
def body_sec(doc, t): return _body(doc, t, 1418)
def body_sub(doc, t): return _body(doc, t, 2268)

def nitem_sec(doc, n, t):
    p = doc.add_paragraph(); pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="60" w:line="240" w:lineRule="auto"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="1778" w:hanging="360"/>'))
    run = p.add_run(f"{n}) {t}"); sfont(run)
    return p

def nitem_sub(doc, n, t):
    p = doc.add_paragraph(); pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="60" w:line="240" w:lineRule="auto"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="2628" w:hanging="360"/>'))
    run = p.add_run(f"{n}) {t}"); sfont(run)
    return p

def tbl_cap(doc, t):
    p = doc.add_paragraph(); pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="240" w:after="120"/>'))
    run = p.add_run(t); sfont(run, bold=True)

def fig_cap(doc, t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="240" w:after="120"/>'))
    run = p.add_run(t); sfont(run, bold=True, italic=True, size=CAPTION_SIZE)

def tbl(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; set_cell_shading(c, HEADER_BG)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.paragraphs[0].add_run(h); sfont(run, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            run = table.rows[ri+1].cells[ci].paragraphs[0].add_run(str(val)); sfont(run)
    return table

def add_img(doc, fig, w=5.5):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight', facecolor='white', edgecolor='none')
    buf.seek(0)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Inches(w))
    plt.close(fig)

# ═══════════════════════════════════════════════════════════
# DIAGRAM GENERATORS
# ═══════════════════════════════════════════════════════════

def _box(ax, x, y, w, h, color, label, fs=9):
    ax.add_patch(FancyBboxPatch((x,y), w, h, boxstyle="round,pad=0.08",
                 facecolor=color, alpha=0.88, edgecolor='#222', linewidth=1.1))
    ax.text(x+w/2, y+h/2, label, ha='center', va='center',
            fontsize=fs, color='white', fontweight='bold')

def _arrow(ax, x1, y1, x2, y2, **kw):
    ax.annotate('', xy=(x2,y2), xytext=(x1,y1),
                arrowprops=dict(arrowstyle='->', color=kw.get('c','#333'), lw=kw.get('lw',1.5),
                                connectionstyle=kw.get('cs','arc3,rad=0')))

def diagram_system_log_architecture():
    """3-tier System Log architecture."""
    fig, ax = plt.subplots(figsize=(10,6)); ax.set_xlim(0,10); ax.set_ylim(0,7); ax.axis('off')
    ax.text(5,6.7,'สถาปัตยกรรมระบบบันทึกเหตุการณ์ (System Log Architecture)',
            ha='center', fontsize=14, fontweight='bold')

    # Source layer
    srcs = [
        (0.2,5.0,1.4,1.0,'#3498db','Auth\nModule'),
        (1.9,5.0,1.4,1.0,'#2ecc71','Animal\nRegistry'),
        (3.6,5.0,1.4,1.0,'#9b59b6','Medical\nRecords'),
        (5.3,5.0,1.4,1.0,'#e67e22','Studbook\n& Genetic'),
        (7.0,5.0,1.4,1.0,'#e74c3c','Billing\nModule'),
        (8.7,5.0,1.2,1.0,'#16a085','อื่นๆ\n(All)'),
    ]
    for s in srcs: _box(ax, *s, fs=8)
    ax.text(5,6.2,'Event Sources (ทุกโมดูล)', ha='center', fontsize=10, color='#555')

    # Collector
    _box(ax, 2.5, 3.2, 5.0, 1.2, '#2c3e50', 'Event Collector & Normalizer\nStandard Schema  |  Async Queue  |  Buffered Write')
    for s in srcs:
        _arrow(ax, s[0]+s[2]/2, s[1], 5.0, 4.4)

    # Storage
    _box(ax, 0.3, 1.0, 3.0, 1.5, '#1a5276',
         'System Log Store\n(Partitioned Table)\nImmutable / Append-only\nRetention ≥ 90 วัน (ตั้งค่าได้)')
    _box(ax, 3.8, 1.0, 3.0, 1.5, '#7d3c98',
         'Audit Trail Store\n(Per-Record Metadata)\ncreated/updated/deleted_by\nversion_no  |  change_log')
    _box(ax, 7.3, 1.0, 2.5, 1.5, '#117a65',
         'Dashboard & Reports\nSearch / Filter\nExport CSV/PDF\nAlert & Monitoring')

    _arrow(ax, 5.0, 3.2, 1.8, 2.5)
    _arrow(ax, 5.0, 3.2, 5.3, 2.5)
    _arrow(ax, 5.0, 3.2, 8.55, 2.5)

    fig.tight_layout()
    return fig


def diagram_event_flow():
    """Event lifecycle from user action → log entry."""
    fig, ax = plt.subplots(figsize=(10,4.5)); ax.set_xlim(0,10); ax.set_ylim(0,5); ax.axis('off')
    ax.text(5,4.7,'วงจรชีวิตเหตุการณ์ (Event Lifecycle) — จากการกระทำของผู้ใช้สู่บันทึกเหตุการณ์',
            ha='center', fontsize=12, fontweight='bold')

    steps = [
        (0.1,2.5,1.5,1.5,'#3498db','1. User Action\nLogin/Logout\nCRUD\nApproval'),
        (2.0,2.5,1.5,1.5,'#27ae60','2. Event\nCapture\nTimestamp\nIP / UserAgent'),
        (3.9,2.5,1.5,1.5,'#8e44ad','3. Normalize\n& Classify\nSeverity Level\nModule Tag'),
        (5.8,2.5,1.5,1.5,'#e67e22','4. Async\nQueue\nBuffered\nNon-blocking'),
        (7.7,2.5,1.5,1.5,'#c0392b','5. Persist\nPartitioned\nTable\nImmutable'),
    ]
    for s in steps: _box(ax, *s, fs=8)
    for i in range(len(steps)-1):
        _arrow(ax, steps[i][0]+steps[i][2], steps[i][1]+steps[i][3]/2,
               steps[i+1][0], steps[i+1][1]+steps[i+1][3]/2)

    # Bottom labels
    _box(ax, 0.5, 0.5, 3.5, 1.2, '#566573',
         'Before/After Values (สำหรับ Update/Delete)\nSQL Diff  |  JSON Patch  |  Field-level Change Log')
    _box(ax, 5.0, 0.5, 4.5, 1.2, '#1a5276',
         'Retention & Purge Policy\nConfigurable: 90 / 180 / 365 วัน (TOR 4.21.5)\nAuto-purge by Partition Drop  |  Archive to Cold Storage')
    _arrow(ax, 8.45, 2.5, 7.25, 1.7, c='#888')
    _arrow(ax, 2.25, 2.5, 2.25, 1.7, c='#888')

    fig.tight_layout()
    return fig


def diagram_deletion_control():
    """Deletion governance workflow."""
    fig, ax = plt.subplots(figsize=(10,6)); ax.set_xlim(0,10); ax.set_ylim(0,7); ax.axis('off')
    ax.text(5,6.7,'กระบวนการควบคุมการลบข้อมูล (Deletion Governance Workflow)',
            ha='center', fontsize=13, fontweight='bold')
    ax.text(5,6.3,'ตาม TOR ข้อ 4.6 — เฉพาะ Super Admin + อ้างอิงใบคำร้อง',
            ha='center', fontsize=10, color='#c0392b', style='italic')

    # Row 1 – request
    _box(ax,0.2,4.5,2.0,1.3,'#3498db','ผู้ใช้ทั่วไป\n(General User)\nไม่สามารถลบได้\nปุ่มลบถูกซ่อน',fs=8)
    _box(ax,2.8,4.5,2.2,1.3,'#e67e22','ยื่นคำร้อง\n(Change Request)\nหมวดหมู่: ลบข้อมูล\nแนบเอกสาร PDF',fs=8)
    _arrow(ax, 2.2, 5.15, 2.8, 5.15)

    # Row 2 – review
    _box(ax,5.6,4.5,2.0,1.3,'#8e44ad','Super Admin\nตรวจสอบคำร้อง\nพิจารณาเหตุผล',fs=8)
    _arrow(ax, 5.0, 5.15, 5.6, 5.15)

    # Decision
    ax.add_patch(plt.Polygon([[8.2,5.15],[8.9,5.8],[9.6,5.15],[8.9,4.5]],
                              facecolor='#f1c40f', edgecolor='#d4ac0d', linewidth=1.5))
    ax.text(8.9,5.15,'อนุมัติ?', ha='center', va='center', fontsize=9, fontweight='bold')
    _arrow(ax, 7.6, 5.15, 8.2, 5.15)

    # Reject
    _box(ax,7.8,2.8,2.0,1.0,'#95a5a6','ปฏิเสธ\nแจ้งเหตุผล\nกลับผู้ร้อง',fs=8)
    ax.annotate('ไม่อนุมัติ', xy=(8.9,2.8+1.0), xytext=(8.9,4.5),
                fontsize=8, color='#c0392b', ha='center',
                arrowprops=dict(arrowstyle='->', color='#c0392b', lw=1.3))

    # Approve path
    _box(ax,0.2,2.0,2.5,1.5,'#27ae60','Soft Delete\n(เปลี่ยน status)\nไม่ลบจริงจาก DB\nข้อมูลยังอยู่',fs=8)
    _box(ax,3.2,2.0,2.5,1.5,'#2c3e50','บันทึก Audit\ndeleted_by\ndeleted_at\ndeletion_reason\ncr_reference',fs=8)
    _box(ax,3.2,0.1,2.5,1.2,'#1a5276','System Log\nAction: DELETE\nAnimal ID\nTimestamp\nIP Address',fs=8)

    ax.annotate('อนุมัติ', xy=(1.45,3.5), xytext=(8.2,4.5),
                fontsize=8, color='#27ae60', ha='center',
                arrowprops=dict(arrowstyle='->', color='#27ae60', lw=1.3,
                                connectionstyle='arc3,rad=0.3'))
    _arrow(ax, 2.7, 2.75, 3.2, 2.75)
    _arrow(ax, 4.45, 2.0, 4.45, 1.3)

    fig.tight_layout()
    return fig


def diagram_audit_coverage():
    """Audit trail coverage across modules."""
    fig, ax = plt.subplots(figsize=(10,5.5)); ax.set_xlim(0,10); ax.set_ylim(0,6); ax.axis('off')
    ax.text(5,5.7,'ขอบเขตการครอบคลุม Audit Trail ตามโมดูล',
            ha='center', fontsize=13, fontweight='bold')

    # Central audit
    _box(ax, 3.0, 3.0, 4.0, 1.5, '#2c3e50',
         'Centralized Audit Layer\ncreated_by | updated_by | deleted_by\ncreated_at | updated_at | version_no\nchange_log (before/after JSON)')

    modules = [
        (0.2,4.5,1.8,1.0,'#3498db','ทะเบียนสัตว์\n(TOR 4.6)\nSoft Delete\n+ CR Ref'),
        (2.3,4.5,1.8,1.0,'#27ae60','ยา/เวชภัณฑ์\n(TOR 4.8)\nSource Tracking\nBarcode Audit'),
        (4.4,4.5,1.8,1.0,'#9b59b6','Billing\n(TOR 4.18.7.1)\nFull Audit Trail\nทุกขั้นตอน'),
        (6.5,4.5,1.8,1.0,'#e67e22','Rate Card\n(TOR 4.18.4.2)\nChange History\nVersion Control'),
        (8.6,4.5,1.2,1.0,'#e74c3c','ทุกโมดูล\nCRUD\nAudit',fs:=7),
    ]
    # fix last one
    modules[-1] = (8.5,4.5,1.3,1.0,'#e74c3c','ทุกโมดูล\nCRUD\nAudit')
    for m in modules:
        _box(ax, *m, fs=7)
        _arrow(ax, m[0]+m[2]/2, m[1], 5.0, 4.5, c='#555')

    # Bottom – System Log
    _box(ax, 0.5, 0.8, 4.0, 1.5, '#1a5276',
         'System Log (TOR 4.21)\nLogin/Logout (4.21.1)\nUser Actions (4.21.2)\nTimestamp ISO 8601 (4.21.3)\nUser ID (4.21.4)\nRetention ≥ 90 วัน (4.21.5)')
    _box(ax, 5.5, 0.8, 4.0, 1.5, '#7d3c98',
         'Deletion Audit (TOR 4.6)\nSuper Admin Only\nCR Reference Required\nSoft Delete + Reason\nBefore/After Snapshot\nImmutable Record')

    _arrow(ax, 5.0, 3.0, 2.5, 2.3, c='#555')
    _arrow(ax, 5.0, 3.0, 7.5, 2.3, c='#555')

    fig.tight_layout()
    return fig


def diagram_log_search_ui():
    """Conceptual UI for log search/filter."""
    fig, ax = plt.subplots(figsize=(10,5)); ax.set_xlim(0,10); ax.set_ylim(0,5.5); ax.axis('off')
    ax.text(5,5.2,'แนวคิดหน้าจอค้นหาและกรองบันทึกเหตุการณ์ (Log Search & Filter UI)',
            ha='center', fontsize=12, fontweight='bold')

    # Filter bar
    _box(ax, 0.3, 3.8, 9.4, 1.1, '#ecf0f1', '', fs=1)
    ax.text(0.5, 4.6, 'Filter Bar', fontsize=10, fontweight='bold', color='#2c3e50')
    filters = [
        (0.5,4.0,'Date Range\n[From] - [To]'),
        (2.5,4.0,'Module\n[Dropdown]'),
        (4.2,4.0,'Action Type\n[Login/CRUD/Delete]'),
        (6.0,4.0,'User\n[Search]'),
        (7.8,4.0,'Severity\n[Info/Warn/Error]'),
    ]
    for x,y,label in filters:
        ax.add_patch(FancyBboxPatch((x,y),1.5,0.7, boxstyle="round,pad=0.04",
                     facecolor='white', edgecolor='#bdc3c7', linewidth=1))
        ax.text(x+0.75, y+0.35, label, ha='center', va='center', fontsize=7, color='#2c3e50')

    # Results table header
    _box(ax, 0.3, 2.8, 9.4, 0.7, '#2c3e50',
         'Timestamp            |  User ID  |  Username  |  Module  |  Action  |  Detail  |  IP Address  |  Severity',
         fs=8)

    # Sample rows
    rows_data = [
        ('2569-03-15 10:30:15  |  1042  |  Dr.Somchai  |  Medical  |  UPDATE  |  วินิจฉัย: ปรับปรุง...  |  10.20.30.40  |  INFO', '#eafaf1'),
        ('2569-03-15 10:28:03  |  1001  |  Admin01  |  Registry  |  DELETE  |  animal_id=5023 CR#042  |  10.20.30.1  |  WARN', '#fdedec'),
        ('2569-03-15 10:25:44  |  2088  |  Keeper_KK  |  Keeper  |  CREATE  |  แจ้งป่วย animal_id=3011  |  10.30.40.5  |  INFO', '#eafaf1'),
        ('2569-03-15 10:20:00  |  1042  |  Dr.Somchai  |  Auth  |  LOGIN  |  Session start  |  10.20.30.40  |  INFO', '#f4f6f7'),
    ]
    for i, (txt, bg) in enumerate(rows_data):
        y = 2.1 - i*0.5
        ax.add_patch(FancyBboxPatch((0.3,y), 9.4, 0.45, boxstyle="square,pad=0",
                     facecolor=bg, edgecolor='#d5d8dc', linewidth=0.5))
        ax.text(0.5, y+0.22, txt, fontsize=7, color='#2c3e50', va='center')

    # Bottom
    ax.text(5, 0.2, 'Export CSV  |  Export PDF  |  Page 1 of 245  |  Total: 12,234 records',
            ha='center', fontsize=8, color='#7f8c8d', style='italic')
    fig.tight_layout()
    return fig


# ═══════════════════════════════════════════════════════════
# DOCUMENT CONSTRUCTION
# ═══════════════════════════════════════════════════════════

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)
    style = doc.styles['Normal']; style.font.name = FONT_TH; style.font.size = BODY_SIZE
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>'); style.element.append(rPr)
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rPr.insert(0, parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="{FONT_TH}" w:ascii="{FONT_TH}" w:hAnsi="{FONT_TH}"/>'))

    # ══════════════════════════════════════════
    # 3.4  Logging & Audit
    # ══════════════════════════════════════════

    h_chapter(doc, "3.4", "ระบบบันทึกเหตุการณ์และร่องรอยการตรวจสอบ (Logging & Audit)")

    body_ch(doc,
        "ระบบบันทึกเหตุการณ์ (System Log) และร่องรอยการตรวจสอบ (Audit Trail) เป็นองค์ประกอบ"
        "พื้นฐานของธรรมาภิบาลข้อมูลที่ช่วยให้องค์การสวนสัตว์แห่งประเทศไทย (อสส.) สามารถ"
        "ตรวจสอบย้อนหลังได้ว่าใครทำอะไร เมื่อไหร่ จากที่ใด และมีผลกระทบอย่างไรต่อข้อมูลในระบบ "
        "Thai Zoo ARK หัวข้อนี้ตอบสนองข้อกำหนด TOR หลายข้อโดยตรง ได้แก่ "
        "TOR ข้อ 4.21 (ระบบบันทึกเหตุการณ์ 5 ข้อย่อย) TOR ข้อ 4.6 (การควบคุมการลบข้อมูลทะเบียนสัตว์) "
        "และ TOR ข้อ 4.18.7.1 (Audit Trail ของระบบคิดค่าบริการ) ซึ่งทั้งหมดนี้มีส่วนสนับสนุน"
        "เกณฑ์การให้คะแนนข้อ 3 ในประเด็น \"ใช้ข้อมูลจากระบบเดิมอย่างเหมาะสม ไม่ซ้ำซ้อน/ขัดแย้ง\" "
        "(30 คะแนน) โดยเนื้อหาประกอบด้วย 6 ส่วน "
        "ส่วนแรก (หัวข้อ 3.4.1) คือสถาปัตยกรรมระบบบันทึกเหตุการณ์ที่ออกแบบเป็น Centralized "
        "Event Collection ครอบคลุมทุกโมดูลทั้งระบบเดิมและระบบใหม่ "
        "ส่วนที่สอง (หัวข้อ 3.4.2) คือรายละเอียดข้อมูลที่บันทึกตาม TOR ข้อ 4.21.1 ถึง 4.21.4 "
        "ส่วนที่สาม (หัวข้อ 3.4.3) คือขอบเขต Audit Trail ที่ครอบคลุมทุกโมดูลวิกฤต "
        "โดยเฉพาะ Billing (TOR 4.18.7.1) และ Rate Card (TOR 4.18.4.2) "
        "ส่วนที่สี่ (หัวข้อ 3.4.4) คือกระบวนการควบคุมการลบข้อมูลตาม TOR ข้อ 4.6 "
        "ส่วนที่ห้า (หัวข้อ 3.4.5) คือนโยบาย Retention และการตั้งค่าระยะเวลาจัดเก็บ "
        "ตาม TOR ข้อ 4.21.5 และส่วนที่หก (หัวข้อ 3.4.6) คือแนวคิดหน้าจอค้นหาและ"
        "กรองบันทึกเหตุการณ์สำหรับผู้ดูแลระบบ "
        "ตารางที่ 3.4-1 ถึง 3.4-7 แสดงรายละเอียดแต่ละประเด็น "
        "และแผนภาพที่ 3.4-1 ถึง 3.4-5 แสดงภาพรวมสถาปัตยกรรมและกระบวนการที่เกี่ยวข้อง"
    )

    # ═══════════════════════════════════════════
    # 3.4.1  System Log Architecture
    # ═══════════════════════════════════════════
    h_section(doc, "3.4.1", "สถาปัตยกรรมระบบบันทึกเหตุการณ์ (System Log Architecture)")

    body_sec(doc,
        "ระบบบันทึกเหตุการณ์ (System Log) ของ Thai Zoo ARK ถูกออกแบบเป็นสถาปัตยกรรม"
        "แบบ Centralized Event Collection ที่รวบรวมเหตุการณ์จากทุกโมดูลไปยังที่จัดเก็บกลาง"
        "เพียงแห่งเดียว โดยไม่กระจายบันทึกเหตุการณ์ไว้ในแต่ละโมดูลแยกกัน "
        "สถาปัตยกรรมนี้ประกอบด้วย 3 ชั้น ชั้นแรกคือ Event Sources ซึ่งครอบคลุมทุกโมดูล"
        "ตั้งแต่ Auth Module (Login/Logout), Animal Registry, Medical Records, Studbook & Genetic, "
        "Billing ไปจนถึงโมดูลอื่นทั้งหมด โดยทุกโมดูลจะส่ง Event ในรูปแบบมาตรฐานเดียวกัน "
        "ชั้นที่สองคือ Event Collector & Normalizer ที่ทำหน้าที่แปลง Event จากทุกแหล่ง"
        "ให้อยู่ใน Schema มาตรฐาน จัดลำดับความสำคัญ (Severity Level) และส่งเข้า Queue "
        "แบบ Asynchronous เพื่อไม่ให้กระทบประสิทธิภาพของระบบหลัก "
        "ชั้นที่สามคือ Persistent Storage ที่แบ่งออกเป็น 3 ส่วน ได้แก่ System Log Store "
        "(ตาราง Partitioned แบบ Append-only), Audit Trail Store (Metadata ระดับ Record) "
        "และ Dashboard & Reports สำหรับการค้นหา กรอง และส่งออกรายงาน "
        "แผนภาพที่ 3.4-1 แสดงภาพรวมสถาปัตยกรรมทั้ง 3 ชั้น "
        "และตารางที่ 3.4-1 แสดงรายละเอียดองค์ประกอบแต่ละชั้น"
    )

    fig = diagram_system_log_architecture()
    add_img(doc, fig, 5.8)
    fig_cap(doc, "แผนภาพที่ 3.4-1 สถาปัตยกรรมระบบบันทึกเหตุการณ์แบบ Centralized Event Collection")

    tbl_cap(doc, "ตารางที่ 3.4-1 องค์ประกอบสถาปัตยกรรมระบบบันทึกเหตุการณ์")
    tbl(doc,
        ["ชั้น (Layer)", "องค์ประกอบ", "หน้าที่", "เทคโนโลยี/แนวทาง", "TOR Ref"],
        [
            ["Event Sources\n(แหล่งเหตุการณ์)", "ทุกโมดูลในระบบ\nทั้งเดิมและใหม่",
             "ส่ง Event เมื่อมี User Action\n(Login/Logout/CRUD/Delete/Approve)",
             "Event Emitter Pattern\nMiddleware Interceptor", "4.21.1-4.21.4"],
            ["Event Collector\n(ตัวรวบรวม)", "Normalizer + Queue",
             "แปลง Event ให้อยู่ใน Schema เดียว\nจัดลำดับ Severity\nส่งเข้า Queue แบบ Async",
             "Message Queue\n(In-memory / Redis)\nAsync Processing", "-"],
            ["System Log Store\n(คลังบันทึก)", "Partitioned Table\n(Monthly Partition)",
             "จัดเก็บบันทึกเหตุการณ์\nแบบ Immutable (Append-only)\nPartitioned เพื่อประสิทธิภาพ",
             "MySQL/PostgreSQL\nTable Partitioning\nIndex on timestamp, user_id", "4.21.5"],
            ["Audit Trail Store\n(คลัง Audit)", "Per-Record Metadata\n+ Change Log Table",
             "เก็บ who/when/what ระดับ Record\nรวมถึง before/after values\nสำหรับ Update และ Delete",
             "Trigger-based\nApplication-level\nJSON diff storage", "4.6, 4.18.7.1"],
            ["Dashboard & Reports\n(รายงาน)", "Search/Filter UI\nExport Functions",
             "ค้นหาและกรองบันทึก\nส่งออก CSV/PDF\nAlert เมื่อมีเหตุการณ์วิกฤต",
             "Web UI\nElastic Search (Optional)\nScheduled Reports", "-"],
        ])

    # ═══════════════════════════════════════════
    # 3.4.2  System Log Fields (TOR 4.21)
    # ═══════════════════════════════════════════
    h_section(doc, "3.4.2", "รายละเอียดข้อมูลที่บันทึกในระบบ System Log (TOR 4.21)")

    body_sec(doc,
        "TOR ข้อ 4.21 กำหนดข้อมูลที่ระบบ System Log ต้องบันทึกไว้ 5 ข้อย่อย "
        "ครอบคลุมเหตุการณ์เข้า-ออกระบบ (TOR 4.21.1) การกระทำของผู้ใช้ (TOR 4.21.2) "
        "วันเวลา (TOR 4.21.3) ผู้ใช้งาน (TOR 4.21.4) และระยะเวลาจัดเก็บ (TOR 4.21.5) "
        "ผู้เสนอได้ออกแบบโครงสร้างข้อมูลของ System Log ให้ครอบคลุมทุกข้อกำหนดดังกล่าว "
        "โดยเพิ่มฟิลด์เสริมที่จำเป็นสำหรับการวิเคราะห์ความปลอดภัยและการตรวจสอบย้อนหลัง "
        "ได้แก่ IP Address, User Agent, Module Name, Severity Level และ Detail "
        "ซึ่งเป็นมาตรฐานสากลตามแนวปฏิบัติ OWASP Logging Cheat Sheet "
        "ตารางที่ 3.4-2 แสดงการเชื่อมโยงระหว่างข้อกำหนด TOR กับฟิลด์ในระบบ "
        "และตารางที่ 3.4-3 แสดงโครงสร้างตาราง system_log โดยละเอียด "
        "แผนภาพที่ 3.4-2 แสดงวงจรชีวิตเหตุการณ์ตั้งแต่การกระทำของผู้ใช้จนถึงการจัดเก็บ"
    )

    tbl_cap(doc, "ตารางที่ 3.4-2 การเชื่อมโยงข้อกำหนด TOR 4.21 กับฟิลด์ในระบบ System Log")
    tbl(doc,
        ["TOR ข้อ", "ข้อกำหนด", "ฟิลด์ในระบบ", "ตัวอย่างค่า", "หมายเหตุ"],
        [
            ["4.21.1", "บันทึกการเข้าสู่ระบบ/ออกจากระบบ\n(Login/Logout)",
             "action = 'LOGIN' หรือ 'LOGOUT'\ntimestamp, user_id, ip_address",
             "action='LOGIN'\ntimestamp='2569-03-15 08:30:00'\nuser_id=1042\nip='10.20.30.40'",
             "บันทึกทุกครั้ง\nรวม Failed Login"],
            ["4.21.2", "บันทึกการดำเนินการของผู้ใช้\nอย่างน้อย Create, Update",
             "action = 'CREATE'/'UPDATE'/'DELETE'\n/'APPROVE'/'EXPORT'\nmodule, detail",
             "action='UPDATE'\nmodule='MEDICAL'\ndetail='Updated diagnosis\nfor animal_id=5023'",
             "รวม Delete และ\nApprove ด้วย\n(เกินข้อกำหนดขั้นต่ำ)"],
            ["4.21.3", "บันทึกวันเวลา\n(Timestamp)",
             "timestamp\n(ISO 8601 with timezone)",
             "2569-03-15T10:30:15+07:00",
             "ใช้ Server time\nTimezone: Asia/Bangkok"],
            ["4.21.4", "บันทึกชื่อผู้ใช้งาน/รหัสผู้ใช้งาน\n(Username/User ID)",
             "user_id (FK → user)\nสามารถ JOIN ดู username",
             "user_id=1042\n→ username='Dr.Somchai'",
             "เก็บ user_id เป็นหลัก\nJOIN ดู username เมื่อแสดง"],
            ["4.21.5", "ตั้งค่าระยะเวลาจัดเก็บได้\nอย่างน้อย 90 วัน",
             "retention_days (System Config)\nPartitioned Table (Monthly)",
             "retention_days = 90\n(ค่าเริ่มต้น, ตั้งได้ถึง 365)",
             "Auto-purge โดย\nDrop Partition\nตาม Config"],
        ])

    tbl_cap(doc, "ตารางที่ 3.4-3 โครงสร้างตาราง system_log (Data Dictionary)")
    tbl(doc,
        ["ฟิลด์", "ชนิดข้อมูล", "Nullable", "Constraint", "คำอธิบาย", "ตัวอย่าง"],
        [
            ["id", "BIGINT", "NOT NULL", "PK, AUTO_INCREMENT", "ลำดับเหตุการณ์", "1, 2, 3, ..."],
            ["timestamp", "DATETIME(3)", "NOT NULL", "INDEX\nPARTITION KEY",
             "วันเวลาเหตุการณ์\n(ISO 8601 + milliseconds)", "2569-03-15 10:30:15.123"],
            ["user_id", "INT", "NOT NULL", "FK → user.user_id\nINDEX",
             "รหัสผู้ใช้งาน (TOR 4.21.4)", "1042"],
            ["action", "ENUM", "NOT NULL", "INDEX",
             "ประเภทเหตุการณ์\n(TOR 4.21.1 + 4.21.2)",
             "LOGIN, LOGOUT,\nCREATE, READ, UPDATE,\nDELETE, APPROVE, EXPORT"],
            ["module", "VARCHAR(50)", "NOT NULL", "INDEX",
             "โมดูลที่เกิดเหตุการณ์", "AUTH, REGISTRY,\nMEDICAL, STUDBOOK,\nBILLING, KEEPER, ..."],
            ["severity", "ENUM", "NOT NULL", "-",
             "ระดับความสำคัญ", "INFO, WARNING,\nERROR, CRITICAL"],
            ["detail", "TEXT", "NULL", "-",
             "รายละเอียดเหตุการณ์\n(free text / JSON)",
             "\"Updated diagnosis\nfor animal_id=5023\""],
            ["entity_type", "VARCHAR(50)", "NULL", "-",
             "ชนิดข้อมูลที่เกี่ยวข้อง", "animal, studbook,\ninvoice, drug, ..."],
            ["entity_id", "VARCHAR(50)", "NULL", "INDEX",
             "รหัสข้อมูลที่เกี่ยวข้อง", "5023, INV-001, ..."],
            ["ip_address", "VARCHAR(45)", "NOT NULL", "-",
             "IP Address ผู้ใช้งาน", "10.20.30.40"],
            ["user_agent", "VARCHAR(500)", "NULL", "-",
             "Browser / Device info", "Chrome/120.0\niPhone Safari/17"],
            ["session_id", "VARCHAR(100)", "NULL", "-",
             "รหัส Session", "sess_abc123..."],
            ["before_value", "JSON", "NULL", "-",
             "ค่าก่อนเปลี่ยนแปลง\n(สำหรับ UPDATE/DELETE)", "{\"status\":\"alive\"}"],
            ["after_value", "JSON", "NULL", "-",
             "ค่าหลังเปลี่ยนแปลง\n(สำหรับ UPDATE)", "{\"status\":\"deceased\"}"],
        ])

    fig = diagram_event_flow()
    add_img(doc, fig, 5.8)
    fig_cap(doc, "แผนภาพที่ 3.4-2 วงจรชีวิตเหตุการณ์จากการกระทำของผู้ใช้สู่การจัดเก็บถาวร")

    # ═══════════════════════════════════════════
    # 3.4.3  Audit Trail Coverage
    # ═══════════════════════════════════════════
    h_section(doc, "3.4.3", "ขอบเขตการครอบคลุม Audit Trail (Audit Trail Coverage)")

    body_sec(doc,
        "Audit Trail ในระบบ Thai Zoo ARK ถูกออกแบบให้ครอบคลุม 2 ระดับ "
        "ระดับแรกคือ Record-level Audit ซึ่งทุกตารางในระบบจะมีฟิลด์ Metadata "
        "ได้แก่ created_by, created_at, updated_by, updated_at และ version_no "
        "เพื่อให้ทราบว่าข้อมูลแต่ละ Record ถูกสร้างและแก้ไขโดยใครเมื่อไหร่ "
        "ระดับที่สองคือ Module-specific Audit ซึ่ง TOR กำหนดให้โมดูลบางโมดูล"
        "ต้องมี Audit Trail เฉพาะทางที่เข้มงวดมากกว่าปกติ โดยเฉพาะ 4 โมดูลวิกฤต ได้แก่ "
        "โมดูลทะเบียนสัตว์ (TOR 4.6) ที่ต้องควบคุมการลบอย่างเข้มงวด "
        "โมดูล Billing (TOR 4.18.7.1) ที่ต้องมี Audit Trail ทุกขั้นตอนของการคิดค่าบริการ "
        "โมดูล Rate Card (TOR 4.18.4.2) ที่ต้องเก็บประวัติการเปลี่ยนแปลงอัตราค่าบริการ "
        "และโมดูลยา/เวชภัณฑ์ (TOR 4.8) ที่ต้องติดตามแหล่งที่มาและ Lot/Batch ของยา "
        "แผนภาพที่ 3.4-3 แสดงขอบเขตการครอบคลุม Audit Trail ตามโมดูล "
        "และตารางที่ 3.4-4 แสดงรายละเอียดข้อกำหนด Audit ของแต่ละโมดูลวิกฤต"
    )

    fig = diagram_audit_coverage()
    add_img(doc, fig, 5.8)
    fig_cap(doc, "แผนภาพที่ 3.4-3 ขอบเขตการครอบคลุม Audit Trail ตามโมดูลทั้งระบบ")

    tbl_cap(doc, "ตารางที่ 3.4-4 ข้อกำหนด Audit Trail เฉพาะทางของโมดูลวิกฤต")
    tbl(doc,
        ["โมดูล", "TOR Ref", "ข้อกำหนด Audit", "ฟิลด์/กลไกที่ใช้",
         "ตัวอย่างเหตุการณ์ที่บันทึก", "ระดับ Severity"],
        [
            ["ทะเบียนสัตว์\n(Animal Registry)", "4.6",
             "ควบคุมการลบข้อมูล\nเฉพาะ Super Admin\nอ้างอิงใบคำร้อง\nบันทึกเหตุผล",
             "deleted_by, deleted_at\ndeletion_reason\ncr_reference\nSoft Delete flag",
             "Super Admin ลบ animal_id=5023\nCR#042 เหตุผล: สัตว์ตาย\nจาก sire_id=3001",
             "WARNING"],
            ["ค่าบริการรักษา\n(Billing)", "4.18.7.1",
             "Audit Trail ทุกขั้นตอน\nตั้งแต่สร้าง Invoice\nจนถึงชำระเงิน",
             "billing_audit table\nstep, actor, timestamp\nbefore/after values\nstatus transitions",
             "สร้าง Invoice INV-001\nเพิ่มหัตถการ P-005\nอนุมัติยอด 5,500 บาท\nชำระเงินสำเร็จ",
             "INFO"],
            ["อัตราค่าบริการ\n(Rate Card)", "4.18.4.2",
             "ประวัติการเปลี่ยนแปลง\nอัตราค่าบริการ\nEffective Date tracking",
             "rate_card_history table\nold_rate, new_rate\neffective_date\nchanged_by, changed_at",
             "เปลี่ยนค่าหัตถการ P-005\nจาก 500 เป็น 600 บาท\nมีผลตั้งแต่ 1 เม.ย. 2569",
             "INFO"],
            ["ยา/เวชภัณฑ์\n(Pharmacy)", "4.8",
             "ติดตามแหล่งที่มา\nLot/Batch tracking\nBarcode audit",
             "drug_source (donate/buy)\nlot_no, batch_no\nbarcode_scanned_at\nexpiry_date tracking",
             "รับยา Amoxicillin Lot#A23\nSource: บริจาค\nBarcode: 8859001234567\nExpiry: 15 ก.ย. 2570",
             "INFO"],
            ["เจ้าหน้าที่\n(Staff Mgmt)", "4.5",
             "ติดตามการเปลี่ยนแปลง\nสิทธิ์และบทบาท",
             "permission_audit table\nold_role, new_role\npermission changes\ngranted_by, granted_at",
             "เปลี่ยน Role user_id=2088\nจาก Keeper เป็น Zoo Admin\nโดย Super Admin",
             "WARNING"],
            ["ทุกโมดูล\n(Global)", "4.21",
             "System Log กลาง\nCRUD + Login/Logout\nTimestamp + User ID",
             "system_log table\n(ตามตารางที่ 3.4-3)",
             "ทุกเหตุการณ์ CRUD\nLogin/Logout\nExport/Approve",
             "INFO / WARNING"],
        ])

    # ═══════════════════════════════════════════
    # 3.4.4  Deletion Control (TOR 4.6)
    # ═══════════════════════════════════════════
    h_section(doc, "3.4.4", "กระบวนการควบคุมการลบข้อมูล (Deletion Governance — TOR 4.6)")

    body_sec(doc,
        "TOR ข้อ 4.6 กำหนดอย่างชัดเจนว่าการลบข้อมูลทะเบียนสัตว์ต้องถูกควบคุมอย่างเข้มงวด "
        "โดยเฉพาะอย่างยิ่งผู้ใช้ทั่วไปต้องไม่สามารถลบข้อมูลได้ และการลบต้องผ่านกระบวนการ"
        "อนุมัติที่มีหลักฐาน ผู้เสนอได้ออกแบบกระบวนการควบคุมการลบข้อมูล (Deletion Governance) "
        "ที่ประกอบด้วย 5 ขั้นตอน ขั้นตอนแรกคือผู้ใช้ทั่วไปจะไม่เห็นปุ่มลบในหน้าจอ "
        "(ปุ่มถูกซ่อนตาม Role-based UI Control) ขั้นตอนที่สองคือเมื่อต้องการลบ "
        "ต้องยื่นคำร้อง (Change Request) ผ่านระบบคำร้อง (TOR 4.19) โดยเลือกหมวดหมู่ "
        "\"ขอลบข้อมูล\" และแนบเอกสารประกอบ ขั้นตอนที่สามคือ Super Admin ตรวจสอบ"
        "และพิจารณาอนุมัติ ขั้นตอนที่สี่คือเมื่ออนุมัติ ระบบจะทำ Soft Delete "
        "(เปลี่ยนสถานะเป็น deleted แต่ไม่ลบข้อมูลจริงจากฐานข้อมูล) พร้อมบันทึก "
        "deleted_by, deleted_at, deletion_reason และ cr_reference "
        "และขั้นตอนที่ห้าคือระบบบันทึก System Log สำหรับเหตุการณ์ลบพร้อม before-snapshot "
        "เพื่อให้สามารถกู้คืนข้อมูลได้หากจำเป็น "
        "แผนภาพที่ 3.4-4 แสดงกระบวนการควบคุมการลบข้อมูลทั้งหมด "
        "และตารางที่ 3.4-5 แสดงเปรียบเทียบสิทธิ์ในการลบข้อมูลแยกตามบทบาท"
    )

    fig = diagram_deletion_control()
    add_img(doc, fig, 5.8)
    fig_cap(doc, "แผนภาพที่ 3.4-4 กระบวนการควบคุมการลบข้อมูลตาม TOR ข้อ 4.6")

    tbl_cap(doc, "ตารางที่ 3.4-5 เปรียบเทียบสิทธิ์ในการลบข้อมูลแยกตามบทบาท")
    tbl(doc,
        ["บทบาท", "สิทธิ์ลบ", "เงื่อนไข", "สิ่งที่บันทึก", "ผลที่เกิดขึ้น"],
        [
            ["ผู้ใช้ทั่วไป\n(General User)", "ไม่มีสิทธิ์ลบ\n(ปุ่มลบถูกซ่อน)",
             "ไม่มี — ไม่สามารถลบได้\nไม่ว่ากรณีใด",
             "ไม่มีการบันทึก\n(เพราะไม่สามารถทำได้)",
             "ข้อมูลถูกปกป้อง\nจากการลบโดยไม่ตั้งใจ"],
            ["Zoo Admin", "ไม่มีสิทธิ์ลบ\nข้อมูลทะเบียนสัตว์",
             "สามารถยื่นคำร้อง\nขอลบได้เท่านั้น",
             "บันทึกการยื่นคำร้อง\nในระบบ Change Request",
             "คำร้องรอ Super Admin\nพิจารณา"],
            ["Super Admin\n(HQ)", "มีสิทธิ์อนุมัติการลบ\n(เฉพาะผ่านคำร้อง)",
             "ต้องอ้างอิงเลขที่คำร้อง (CR#)\nต้องระบุเหตุผลการลบ\nต้องยืนยันก่อนลบ",
             "deleted_by = user_id\ndeleted_at = timestamp\ndeletion_reason = text\ncr_reference = CR#\nbefore_snapshot = JSON",
             "Soft Delete\n(เปลี่ยน status ไม่ลบจริง)\nข้อมูลสามารถกู้คืนได้"],
            ["ระบบ\n(System)", "ไม่สามารถลบข้อมูล\nโดยอัตโนมัติ",
             "ไม่มี Cascade Delete\nไม่มี Auto-purge\nสำหรับข้อมูลธุรกิจ",
             "System Log บันทึก\nทุก Error ที่เกี่ยวกับการลบ",
             "ป้องกัน Orphan Data\nและ Referential Integrity"],
        ])

    # ═══════════════════════════════════════════
    # 3.4.5  Retention & Configuration
    # ═══════════════════════════════════════════
    h_section(doc, "3.4.5", "นโยบาย Retention และการตั้งค่าระยะเวลาจัดเก็บ (TOR 4.21.5)")

    body_sec(doc,
        "TOR ข้อ 4.21.5 กำหนดให้ระบบ System Log สามารถตั้งค่าระยะเวลาจัดเก็บได้ "
        "โดยมีระยะเวลาขั้นต่ำอย่างน้อย 90 วัน ผู้เสนอได้ออกแบบระบบ Retention Management "
        "ที่ยืดหยุ่นและมีประสิทธิภาพ โดยใช้เทคนิค Table Partitioning แบ่งตาราง system_log "
        "เป็น Monthly Partitions ทำให้การลบข้อมูลเก่าเป็นเพียงการ Drop Partition "
        "ซึ่งมีประสิทธิภาพสูงกว่าการ DELETE ทีละ Row อย่างมาก "
        "ผู้ดูแลระบบสามารถตั้งค่าระยะเวลาจัดเก็บผ่านหน้า Admin Config "
        "โดยมีตัวเลือก 90 วัน (ค่าเริ่มต้นตาม TOR), 180 วัน และ 365 วัน "
        "ระบบจะมี Scheduled Job ทำการตรวจสอบและ Purge ข้อมูลที่เกินกำหนดโดยอัตโนมัติ "
        "สำหรับข้อมูลที่ถูก Purge จาก Active Storage สามารถ Archive ไปยัง Cold Storage "
        "ได้หากต้องการเก็บรักษาระยะยาว "
        "ตารางที่ 3.4-6 แสดงรายละเอียดนโยบาย Retention แยกตามประเภทข้อมูล"
    )

    tbl_cap(doc, "ตารางที่ 3.4-6 นโยบาย Retention แยกตามประเภทข้อมูล Log")
    tbl(doc,
        ["ประเภท Log", "ระยะเวลา Active", "ระยะเวลา Archive",
         "วิธี Purge", "การตั้งค่า", "TOR Ref"],
        [
            ["System Log\n(Login/Logout/Actions)", "90-365 วัน\n(ตั้งค่าได้)",
             "ไม่บังคับ\n(Optional Archive)",
             "Drop Partition\n(Monthly Partition)", "Admin Config UI\nค่าเริ่มต้น = 90 วัน", "4.21.5"],
            ["Audit Trail\n(Record-level)", "ตลอดอายุ Record\n(ไม่มี Purge)",
             "ไม่มี\n(เก็บถาวร)",
             "ไม่มี — เก็บตลอด\nเป็นส่วนหนึ่งของ Record", "ไม่สามารถตั้งค่า\n(Permanent)", "4.6, 4.18.7.1"],
            ["Deletion Audit\n(Before-snapshot)", "ตลอด\n(ไม่มี Purge)",
             "ไม่มี\n(เก็บถาวร)",
             "ไม่มี — เก็บตลอด\nเพื่อการกู้คืนข้อมูล", "ไม่สามารถตั้งค่า\n(Permanent)", "4.6"],
            ["Billing Audit\n(Step-by-step)", "7 ปี\n(ตามระเบียบการเงิน)",
             "Archive หลัง 7 ปี",
             "Move to Archive Table", "System Config\nค่าเริ่มต้น = 7 ปี", "4.18.7.1"],
            ["Rate Card History", "ถาวร\n(เก็บทุกเวอร์ชัน)",
             "ไม่มี",
             "ไม่มี — ประวัติราคา\nต้องเก็บถาวร", "ไม่สามารถตั้งค่า\n(Permanent)", "4.18.4.2"],
            ["Permission Change\nAudit", "3 ปี\n(ค่าเริ่มต้น)",
             "Archive หลัง 3 ปี",
             "Move to Archive Table", "Admin Config UI", "4.5"],
        ])

    # ═══════════════════════════════════════════
    # 3.4.6  Log Search & Filter UI
    # ═══════════════════════════════════════════
    h_section(doc, "3.4.6", "แนวคิดหน้าจอค้นหาและกรองบันทึกเหตุการณ์ (Log Search & Filter UI)")

    body_sec(doc,
        "เพื่อให้ผู้ดูแลระบบสามารถใช้ประโยชน์จากข้อมูล System Log ได้อย่างเต็มที่ "
        "ผู้เสนอได้ออกแบบหน้าจอค้นหาและกรองบันทึกเหตุการณ์ที่รองรับเงื่อนไขการค้นหา"
        "หลากหลาย ครอบคลุมการกรองตามช่วงวันที่ (Date Range) โมดูลที่เกิดเหตุการณ์ "
        "ประเภทเหตุการณ์ (Action Type) ผู้ใช้งาน (User) และระดับความสำคัญ (Severity) "
        "ผลลัพธ์จะแสดงเป็นตารางที่เรียงตามเวลาล่าสุดก่อน โดยแต่ละรายการจะแสดง Timestamp, "
        "User ID, Username, Module, Action, Detail, IP Address และ Severity "
        "ระบบรองรับการส่งออกผลลัพธ์เป็น CSV และ PDF สำหรับการตรวจสอบภายนอกระบบ "
        "สำหรับเหตุการณ์ที่มี Severity ระดับ WARNING หรือ CRITICAL ระบบจะใช้สีพื้นหลัง"
        "แตกต่างเพื่อให้สังเกตได้ง่าย (เช่น สีแดงอ่อนสำหรับ DELETE) "
        "นอกจากนี้ ยังมีระบบ Alert อัตโนมัติที่จะแจ้งเตือน Super Admin "
        "เมื่อมีเหตุการณ์วิกฤต เช่น Failed Login มากกว่า 5 ครั้งติดต่อกัน "
        "หรือมีการ DELETE ข้อมูลที่มีผลกระทบสูง "
        "แผนภาพที่ 3.4-5 แสดงแนวคิดหน้าจอค้นหาบันทึกเหตุการณ์ "
        "และตารางที่ 3.4-7 แสดงรายละเอียดเงื่อนไขการค้นหาที่รองรับ"
    )

    fig = diagram_log_search_ui()
    add_img(doc, fig, 5.8)
    fig_cap(doc, "แผนภาพที่ 3.4-5 แนวคิดหน้าจอค้นหาและกรองบันทึกเหตุการณ์ (Log Search & Filter)")

    tbl_cap(doc, "ตารางที่ 3.4-7 เงื่อนไขการค้นหาและกรองบันทึกเหตุการณ์ที่รองรับ")
    tbl(doc,
        ["เงื่อนไข", "ประเภท Input", "ค่าที่รองรับ", "ค่าเริ่มต้น", "หมายเหตุ"],
        [
            ["ช่วงวันที่\n(Date Range)", "Date Picker\n(From - To)",
             "ช่วงวันที่ใดก็ได้\nภายใน Retention Period", "7 วันล่าสุด",
             "ไม่สามารถค้นหา\nเกิน Retention Period"],
            ["โมดูล\n(Module)", "Dropdown\n(Multi-select)",
             "AUTH, REGISTRY,\nMEDICAL, NUTRITION,\nPHARMACY, WAREHOUSE,\nKEEPER, STUDBOOK,\nGENETIC, BILLING,\nCR, CMS, SYSLOG",
             "ทั้งหมด (All)", "สามารถเลือกได้\nหลายโมดูลพร้อมกัน"],
            ["ประเภทเหตุการณ์\n(Action)", "Dropdown\n(Multi-select)",
             "LOGIN, LOGOUT,\nCREATE, READ,\nUPDATE, DELETE,\nAPPROVE, EXPORT",
             "ทั้งหมด (All)", "สามารถเลือกได้\nหลายประเภทพร้อมกัน"],
            ["ผู้ใช้งาน\n(User)", "Search / Autocomplete",
             "ค้นหาด้วย User ID\nหรือ Username", "ทั้งหมด (All)",
             "รองรับ Autocomplete\nจากตาราง User"],
            ["ระดับความสำคัญ\n(Severity)", "Dropdown\n(Multi-select)",
             "INFO, WARNING,\nERROR, CRITICAL", "ทั้งหมด (All)",
             "สีพื้นหลังแตกต่าง\nตาม Severity"],
            ["ข้อมูลที่เกี่ยวข้อง\n(Entity)", "Search",
             "Entity Type + Entity ID\nเช่น animal:5023", "ว่าง (ไม่กรอง)",
             "สำหรับค้นหา\nประวัติข้อมูลเฉพาะตัว"],
            ["คำค้นหา\n(Full Text)", "Text Input",
             "ค้นหาใน Detail field", "ว่าง (ไม่กรอง)",
             "Full-text search\nใน detail column"],
            ["การส่งออก\n(Export)", "Button",
             "CSV, PDF", "-", "ส่งออกตาม Filter\nที่ตั้งค่าไว้"],
        ])

    # ── Save ──
    os.makedirs(OUT_DIR, exist_ok=True)
    doc.save(OUT_FILE)
    print(f"Document saved to: {OUT_FILE}")


if __name__ == "__main__":
    build()
