#!/usr/bin/env python3
"""Generate Section 3.6 Module-by-Module Design for Thai Zoo ARK proposal."""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import matplotlib.font_manager as fm
import numpy as np
from io import BytesIO

# ── Font Setup ──
_FONT_DIR = os.path.expanduser("~/Library/Fonts/")
for _fname in [
    "THSarabunNew.ttf", "THSarabunNew Bold.ttf",
    "THSarabunNew Italic.ttf", "THSarabunNew BoldItalic.ttf",
]:
    _path = os.path.join(_FONT_DIR, _fname)
    if os.path.exists(_path):
        fm.fontManager.addfont(_path)
plt.rcParams['font.family'] = 'TH Sarabun New'
plt.rcParams['font.sans-serif'] = ['TH Sarabun New', 'TH SarabunPSK', 'DejaVu Sans']

# ── Constants ──
FONT_TH = "TH Sarabun New"
BODY_SIZE = Pt(16)
HEADING_SIZE = Pt(16)
CAPTION_SIZE = Pt(14)
HEADER_BG = "D6E3F8"
OUT_DIR = os.path.expanduser("~/Library/CloudStorage/OneDrive-Personal/Personal/ZTL")
OUT_FILE = os.path.join(OUT_DIR, "Section_3_6_Module_by_Module_Design.docx")

# ── Counters for continuous numbering ──
table_counter = 0
diagram_counter = 0

def next_table_num():
    global table_counter
    table_counter += 1
    return f"3.6-{table_counter}"

def next_diagram_num():
    global diagram_counter
    diagram_counter += 1
    return f"3.6-{diagram_counter}"


# ── Helpers ──
def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, font_name=FONT_TH, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="{font_name}" w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="{int(size.pt * 2)}"/>')
        rPr.append(szCs)
    else:
        szCs.set(qn('w:val'), str(int(size.pt * 2)))
    if bold:
        if rPr.find(qn('w:bCs')) is None:
            rPr.append(parse_xml(f'<w:bCs {nsdecls("w")}/>'))
    if color:
        run.font.color.rgb = color

def add_chapter_heading(doc, number, title):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="240" w:after="120"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="567" w:hanging="567"/>'))
    pPr.append(parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="567"/></w:tabs>'))
    run = p.add_run(f"{number}\t{title}")
    set_run_font(run, bold=True, size=HEADING_SIZE)
    return p

def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="200" w:after="80"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="1418" w:hanging="851"/>'))
    pPr.append(parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="1418"/></w:tabs>'))
    run = p.add_run(f"{number}\t{title}")
    set_run_font(run, bold=True, size=HEADING_SIZE)
    return p

def add_subsection_heading(doc, number, title):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="160" w:after="80"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="2268" w:hanging="850"/>'))
    pPr.append(parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="2268"/></w:tabs>'))
    run = p.add_run(f"{number}\t{title}")
    set_run_font(run, bold=True, size=HEADING_SIZE)
    return p

def add_body_after_chapter(doc, text):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="120" w:line="240" w:lineRule="auto"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="567"/>'))
    run = p.add_run(text)
    set_run_font(run)
    return p

def add_body_after_section(doc, text):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="120" w:line="240" w:lineRule="auto"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="1418"/>'))
    run = p.add_run(text)
    set_run_font(run)
    return p

def add_body_after_subsection(doc, text):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="120" w:line="240" w:lineRule="auto"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="2268"/>'))
    run = p.add_run(text)
    set_run_font(run)
    return p

def add_numbered_item_section(doc, number, text):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="60" w:line="240" w:lineRule="auto"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="1778" w:hanging="360"/>'))
    run = p.add_run(f"{number}) {text}")
    set_run_font(run)
    return p

def add_numbered_item_subsection(doc, number, text):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="60" w:line="240" w:lineRule="auto"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="2628" w:hanging="360"/>'))
    run = p.add_run(f"{number}) {text}")
    set_run_font(run)
    return p

def add_table_caption(doc, num, title):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="240" w:after="120"/>'))
    run = p.add_run(f"ตารางที่ {num} {title}")
    set_run_font(run, bold=True)
    return p

def add_figure_caption(doc, num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="240" w:after="120"/>'))
    run = p.add_run(f"แผนภาพที่ {num} {title}")
    set_run_font(run, bold=True, italic=True, size=CAPTION_SIZE)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_image_from_fig(doc, fig, width_inches=5.5):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight', facecolor='white', edgecolor='none')
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))
    plt.close(fig)
    return p


# ── Diagram Generators ──

def create_flow_diagram(title, boxes, arrows, figsize=(10, 5)):
    """Generic flow diagram generator. boxes = [(x, y, w, h, color, label), ...], arrows = [(from_idx, to_idx), ...]"""
    fig, ax = plt.subplots(1, 1, figsize=figsize)
    ax.set_xlim(-0.5, 11)
    ax.set_ylim(-0.5, figsize[1] + 0.5)
    ax.axis('off')
    ax.text(5.25, figsize[1] + 0.1, title, ha='center', va='center', fontsize=12, fontweight='bold')

    centers = []
    for x, y, w, h, color, label in boxes:
        box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08", facecolor=color, alpha=0.9)
        ax.add_patch(box)
        ax.text(x + w/2, y + h/2, label, ha='center', va='center', fontsize=8, color='white', fontweight='bold')
        centers.append((x + w/2, y + h/2, x, y, w, h))

    for fi, ti in arrows:
        fx, fy = centers[fi][0], centers[fi][1]
        tx, ty = centers[ti][0], centers[ti][1]
        # Determine connection points
        if abs(fx - tx) > abs(fy - ty):
            if fx < tx:
                start = (centers[fi][2] + centers[fi][4], fy)
                end = (centers[ti][2], ty)
            else:
                start = (centers[fi][2], fy)
                end = (centers[ti][2] + centers[ti][4], ty)
        else:
            if fy > ty:
                start = (fx, centers[fi][3])
                end = (tx, centers[ti][3] + centers[ti][5])
            else:
                start = (fx, centers[fi][3] + centers[fi][5])
                end = (tx, centers[ti][3])
        ax.annotate('', xy=end, xytext=start, arrowprops=dict(arrowstyle='->', color='#444', lw=1.5))

    fig.tight_layout()
    return fig


# ══════════════════════════════════════════════
# MAIN DOCUMENT GENERATION
# ══════════════════════════════════════════════

def build_document():
    global table_counter, diagram_counter
    table_counter = 0
    diagram_counter = 0

    doc = Document()

    # Page setup A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ══════════════════════════════════════════
    # 3.6 Chapter Heading + Narrative
    # ══════════════════════════════════════════
    add_chapter_heading(doc, "3.6", "การออกแบบระบบรายโมดูล (Module-by-Module Design)")

    add_body_after_chapter(doc,
        "การออกแบบระบบรายโมดูลเป็นหัวใจสำคัญของข้อเสนอด้านเทคนิคที่สอดคล้องกับข้อกำหนด TOR ข้อ 4.5 ถึง 4.21 "
        "ครอบคลุมทั้งการปรับปรุงระบบเดิม 6 โมดูล และการพัฒนาระบบใหม่ 8 โมดูล รวมทั้งสิ้น 14 โมดูล "
        "หัวข้อนี้แบ่งเนื้อหาออกเป็น 14 ส่วน เริ่มจากระบบจัดการเจ้าหน้าที่ (หัวข้อ 3.6.1) ระบบทะเบียนสัตว์ (หัวข้อ 3.6.2) "
        "ระบบจัดการโภชนาการสัตว์ (หัวข้อ 3.6.3) ระบบบริหารการเก็บและใช้ยา/เวชภัณฑ์ (หัวข้อ 3.6.4) "
        "ระบบคลังอาหารสัตว์ (หัวข้อ 3.6.5) ระบบผู้ดูแลสัตว์ (หัวข้อ 3.6.6) ระบบ Studbook Management (หัวข้อ 3.6.7) "
        "โมดูล Genetic Analysis (หัวข้อ 3.6.8) ระบบ Pedigree Diagram (หัวข้อ 3.6.9) ระบบ Reports & Dashboard (หัวข้อ 3.6.10) "
        "ระบบคิดค่าบริการดูแลรักษาสัตว์ (หัวข้อ 3.6.11) ระบบคำร้องขอเปลี่ยนแปลงแก้ไข (หัวข้อ 3.6.12) "
        "ระบบจัดเก็บเนื้อหาคู่มือการใช้งาน (หัวข้อ 3.6.13) และระบบบันทึกเหตุการณ์ System Log (หัวข้อ 3.6.14) "
        "แต่ละโมดูลนำเสนอ User Stories ที่แสดงมุมมองของผู้ใช้จริง การออกแบบหน้าจอ จุดเชื่อมต่อข้อมูลกับโมดูลอื่น "
        "กฎการตรวจสอบข้อมูล เกณฑ์การทดสอบและยอมรับ รวมถึงแผนความเข้ากันได้กับระบบเดิม "
        "สำหรับโมดูลที่ปรับปรุง 6 โมดูล (หัวข้อ 3.6.1-3.6.6) แต่ละโมดูลยังมี Gap Analysis "
        "เปรียบเทียบสภาพระบบเดิม (As-Is) กับเป้าหมายหลังปรับปรุง (To-Be) พร้อมแผนภาพประกอบ "
        "นอกจากนี้ หัวข้อ 3.6.15 นำเสนอ Benchmarks และแนวปฏิบัติที่ดีระดับสากล "
        "สำหรับโมดูลใหม่ 8 โมดูล โดยอ้างอิงมาตรฐาน Species360/ZIMS, EAZA Population Management, "
        "OWASP Security Logging และแนวปฏิบัติ Veterinary Practice Management สมัยใหม่ "
        "โดยมีตารางสรุปข้อมูลสำคัญและแผนภาพกระบวนการทำงานประกอบทุกโมดูล"
    )

    # ══════════════════════════════════════════
    # 3.6.1 Staff Management (TOR 4.5)
    # ══════════════════════════════════════════
    build_module_staff(doc)
    build_module_animal_registry(doc)
    build_module_nutrition(doc)
    build_module_pharma(doc)
    build_module_food_warehouse(doc)
    build_module_keeper(doc)
    build_module_studbook(doc)
    build_module_genetic(doc)
    build_module_pedigree(doc)
    build_module_reports_dashboard(doc)
    build_module_billing(doc)
    build_module_change_request(doc)
    build_module_manual_cms(doc)
    build_module_system_log(doc)
    build_module_benchmarks(doc)

    # Save
    os.makedirs(OUT_DIR, exist_ok=True)
    doc.save(OUT_FILE)
    print(f"Saved: {OUT_FILE}")


# ══════════════════════════════════════════════
# MODULE 1: Staff Management
# ══════════════════════════════════════════════

def build_module_staff(doc):
    add_section_heading(doc, "3.6.1", "ระบบจัดการเจ้าหน้าที่ (Staff & Permission Management) — TOR 4.5")

    add_body_after_section(doc,
        "ระบบจัดการเจ้าหน้าที่เป็นโมดูลพื้นฐานที่กำหนดโครงสร้างการเข้าถึงระบบ Thai Zoo ARK ทั้งหมด "
        "โดยในระบบเดิมยังไม่มีการ Validate โครงสร้างองค์กร (ฝ่าย สำนัก ตำแหน่ง) อย่างครบถ้วน "
        "และยังไม่รองรับการกำหนดสิทธิ์แบบรายบุคคล (Per-user Permission) "
        "การปรับปรุงครั้งนี้จึงมุ่งเน้นการยกระดับเป็น Identity and Access Administration "
        "ที่ตรวจสอบโครงสร้างองค์กรตั้งแต่ขั้นตอนสร้างผู้ใช้ รองรับ Role-Based Access Control (RBAC) "
        "ร่วมกับสิทธิ์รายบุคคล และมี Audit Trail สำหรับทุกการเปลี่ยนแปลงสิทธิ์ "
        "หัวข้อนี้ประกอบด้วย User Stories (หัวข้อ 3.6.1.1) การออกแบบหน้าจอ (หัวข้อ 3.6.1.2) "
        "กฎการตรวจสอบข้อมูล (หัวข้อ 3.6.1.3) เกณฑ์การทดสอบ (หัวข้อ 3.6.1.4) "
        "และแผนความเข้ากันได้กับระบบเดิม (หัวข้อ 3.6.1.5) "
        "โดยมีตารางที่ 3.6-1 แสดง User Stories และตารางที่ 3.6-2 แสดงเกณฑ์การทดสอบ "
        "และแผนภาพที่ 3.6-1 แสดงกระบวนการจัดการเจ้าหน้าที่และสิทธิ์การเข้าถึง"
    )

    # 3.6.1.1 User Stories
    add_subsection_heading(doc, "3.6.1.1", "User Stories")
    add_body_after_subsection(doc,
        "User Stories ต่อไปนี้กำหนดความต้องการจากมุมมองผู้ใช้จริงสำหรับระบบจัดการเจ้าหน้าที่ "
        "ครอบคลุมทั้งผู้ดูแลระบบระดับกลาง (Super Admin) และผู้ดูแลระบบระดับสวนสัตว์ (Zoo Admin) "
        "ดังแสดงในตารางที่ 3.6-1"
    )
    tn = next_table_num()
    add_table_caption(doc, tn, "User Stories ระบบจัดการเจ้าหน้าที่")
    add_table(doc,
        ["รหัส", "บทบาท", "ความต้องการ", "เป้าหมาย", "TOR"],
        [
            ["US-5.1", "Super Admin", "ต้องการ Validate โครงสร้างองค์กร (ฝ่าย/สำนัก/ตำแหน่ง) เมื่อสร้างผู้ใช้", "ข้อมูลผู้ใช้ถูกต้องตามโครงสร้าง", "4.5"],
            ["US-5.2", "Super Admin", "ต้องการกำหนดสิทธิ์แบบรายบุคคลนอกเหนือจาก Role", "ผู้ใช้แต่ละคนเห็นเฉพาะข้อมูลที่จำเป็น", "4.5"],
            ["US-5.3", "Zoo Admin", "ต้องการจัดการผู้ใช้ภายในสวนสัตว์ของตนเอง", "กระจายอำนาจการบริหารจัดการ", "4.5"],
            ["US-5.4", "Super Admin", "ต้องการดู Audit Trail การเปลี่ยนแปลงสิทธิ์ทั้งหมด", "ตรวจสอบย้อนหลังได้ทุกการเปลี่ยนแปลง", "4.5"],
        ],
        col_widths=[1.8, 2.5, 5.0, 4.0, 1.2]
    )

    # 3.6.1.2 Screen Design
    add_subsection_heading(doc, "3.6.1.2", "การออกแบบหน้าจอและกระบวนการทำงาน")
    add_body_after_subsection(doc,
        "การออกแบบหน้าจอระบบจัดการเจ้าหน้าที่ประกอบด้วยหน้าจอหลัก 3 ส่วน ได้แก่ "
        "หน้าสร้าง/แก้ไขผู้ใช้พร้อม Dropdown Validate โครงสร้างองค์กร "
        "หน้า Permission Matrix ที่แสดง Role และสิทธิ์รายบุคคลแบบ Override ได้ "
        "และหน้า Audit Trail ที่แสดงประวัติการเปลี่ยนแปลงทุกรายการ "
        "แผนภาพที่ 3.6-1 แสดงกระบวนการทำงานตั้งแต่การสร้างผู้ใช้ การตรวจสอบโครงสร้าง "
        "จนถึงการกำหนดสิทธิ์และบันทึก Audit Log"
    )

    # Diagram
    fig = create_flow_diagram(
        "กระบวนการจัดการเจ้าหน้าที่และสิทธิ์การเข้าถึง (TOR 4.5)",
        [
            (0.0, 3.0, 2.0, 1.0, '#2c3e50', 'Super Admin\n/ Zoo Admin'),
            (2.8, 3.0, 2.2, 1.0, '#2980b9', 'สร้าง/แก้ไข\nข้อมูลผู้ใช้'),
            (5.8, 3.8, 2.2, 1.0, '#27ae60', 'Validate\nฝ่าย/สำนัก/ตำแหน่ง'),
            (5.8, 2.0, 2.0, 1.0, '#e74c3c', 'แจ้งเตือน\nข้อมูลไม่ถูกต้อง'),
            (8.8, 3.8, 2.0, 1.0, '#8e44ad', 'กำหนด Role\n+ Per-user'),
            (8.8, 2.0, 2.0, 1.0, '#16a085', 'บันทึก\nAudit Log'),
        ],
        [(0,1), (1,2), (2,4), (2,3), (4,5)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "กระบวนการจัดการเจ้าหน้าที่และสิทธิ์การเข้าถึง")

    # 3.6.1.3 Data Validations
    add_subsection_heading(doc, "3.6.1.3", "กฎการตรวจสอบข้อมูล (Data Validations)")
    add_body_after_subsection(doc,
        "กฎการตรวจสอบข้อมูลสำหรับระบบจัดการเจ้าหน้าที่ถูกออกแบบให้ครอบคลุมทุกขั้นตอนการทำงาน "
        "เพื่อให้มั่นใจว่าข้อมูลผู้ใช้มีความถูกต้องและสอดคล้องกับโครงสร้างองค์กรจริง"
    )
    add_numbered_item_subsection(doc, 1, "ฝ่าย (Department) ต้องมีอยู่ใน Master Data ขององค์กร")
    add_numbered_item_subsection(doc, 2, "ตำแหน่ง (Position) ต้องเป็นตำแหน่งที่ถูกต้องสำหรับฝ่ายที่เลือก")
    add_numbered_item_subsection(doc, 3, "ชื่อผู้ใช้ (Username) ต้องไม่ซ้ำกันในระบบ (Unique Constraint)")
    add_numbered_item_subsection(doc, 4, "รหัสผ่านต้องมีความยาวไม่น้อยกว่า 8 ตัวอักษร ประกอบด้วยตัวพิมพ์ใหญ่ ตัวพิมพ์เล็ก ตัวเลข และอักขระพิเศษ")
    add_numbered_item_subsection(doc, 5, "สิทธิ์รายบุคคล (Per-user Permission) ต้อง Override สิทธิ์ของ Role ได้โดยไม่กระทบผู้ใช้คนอื่นใน Role เดียวกัน")

    # 3.6.1.4 Acceptance Tests
    add_subsection_heading(doc, "3.6.1.4", "เกณฑ์การทดสอบและยอมรับ (Acceptance Tests)")
    add_body_after_subsection(doc,
        "เกณฑ์การทดสอบต่อไปนี้ใช้ยืนยันว่าระบบจัดการเจ้าหน้าที่ทำงานได้ตรงตามข้อกำหนด TOR 4.5 "
        "ดังแสดงในตารางที่ 3.6-2"
    )
    tn = next_table_num()
    add_table_caption(doc, tn, "เกณฑ์การทดสอบระบบจัดการเจ้าหน้าที่")
    add_table(doc,
        ["รหัส", "รายการทดสอบ", "ผลลัพธ์ที่คาดหวัง"],
        [
            ["AT-5.1", "สร้างผู้ใช้โดยเลือกฝ่ายที่ไม่มีใน Master Data", "ระบบแสดงข้อผิดพลาด ไม่สามารถบันทึกได้"],
            ["AT-5.2", "กำหนดสิทธิ์รายบุคคลให้ผู้ใช้", "ผู้ใช้เข้าถึงได้เฉพาะเมนูตามสิทธิ์ที่กำหนด"],
            ["AT-5.3", "เปลี่ยนแปลงสิทธิ์ผู้ใช้แล้วตรวจสอบ Audit Trail", "บันทึก Audit แสดงผู้เปลี่ยน วันเวลา และรายละเอียดการเปลี่ยนแปลง"],
            ["AT-5.4", "Zoo Admin สร้างผู้ใช้ข้ามสวนสัตว์", "ระบบปฏิเสธ แสดงข้อความว่าไม่มีสิทธิ์"],
        ],
        col_widths=[1.8, 6.0, 6.0]
    )

    # 3.6.1.5 Gap Analysis
    add_subsection_heading(doc, "3.6.1.5", "การวิเคราะห์ช่องว่าง As-Is vs To-Be (Gap Analysis)")
    add_body_after_subsection(doc,
        "จากการวิเคราะห์ระบบจัดการเจ้าหน้าที่เดิมพบช่องว่างสำคัญ 4 ประการ "
        "ที่จำเป็นต้องปรับปรุงเพื่อให้ระบบรองรับการกำกับดูแลและความปลอดภัยตามมาตรฐาน "
        "ตารางที่ {tn} แสดงการเปรียบเทียบ As-Is และ To-Be "
        "และแผนภาพที่ {dn} แสดงสถาปัตยกรรมเปรียบเทียบก่อนและหลังปรับปรุง".format(
            tn="3.6-" + str(table_counter + 1), dn="3.6-" + str(diagram_counter + 1))
    )
    tn = next_table_num()
    add_table_caption(doc, tn, "Gap Analysis ระบบจัดการเจ้าหน้าที่: As-Is vs To-Be")
    add_table(doc,
        ["ด้าน", "As-Is (ระบบเดิม)", "Gap (ช่องว่าง)", "To-Be (เป้าหมาย)"],
        [
            ["โครงสร้างองค์กร", "ไม่ Validate ฝ่าย/สำนัก/ตำแหน่ง", "ข้อมูลผู้ใช้อาจไม่ตรงกับโครงสร้างจริง", "Validate ทุกระดับจาก Master Data"],
            ["สิทธิ์การเข้าถึง", "กำหนดสิทธิ์ตาม Role เท่านั้น", "ไม่สามารถปรับสิทธิ์รายบุคคลได้", "RBAC + Per-user Permission Override"],
            ["Audit Trail", "ไม่บันทึกประวัติการเปลี่ยนแปลงสิทธิ์", "ไม่สามารถตรวจสอบย้อนหลังได้", "บันทึกทุกการเปลี่ยนแปลงพร้อมผู้ดำเนินการ"],
            ["การกระจายอำนาจ", "Super Admin จัดการทุกสวนสัตว์", "คอขวดที่ส่วนกลาง", "Zoo Admin จัดการภายในสวนสัตว์ของตน"],
        ],
        col_widths=[2.5, 3.5, 3.5, 4.0]
    )
    # Gap diagram
    fig, ax = plt.subplots(1, 1, figsize=(10, 4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 4.5); ax.axis('off')
    ax.text(5, 4.2, 'Gap Analysis: ระบบจัดการเจ้าหน้าที่ — As-Is vs To-Be', ha='center', fontsize=12, fontweight='bold')
    # As-Is side
    ax.add_patch(FancyBboxPatch((0.2, 1.0), 4.0, 2.5, boxstyle="round,pad=0.1", facecolor='#e74c3c', alpha=0.15, edgecolor='#e74c3c'))
    ax.text(2.2, 3.2, 'As-Is', ha='center', fontsize=11, fontweight='bold', color='#e74c3c')
    for i, t in enumerate(['Admin สร้าง User (ไม่ Validate)', 'Role-based เท่านั้น', 'ไม่มี Audit Trail', 'จัดการจากส่วนกลางเท่านั้น']):
        ax.text(2.2, 2.6 - i*0.45, t, ha='center', fontsize=9)
    # Arrow
    ax.annotate('', xy=(5.6, 2.2), xytext=(4.4, 2.2), arrowprops=dict(arrowstyle='->', color='#2c3e50', lw=2.5))
    ax.text(5.0, 2.6, 'ปรับปรุง', ha='center', fontsize=9, fontweight='bold', color='#2c3e50')
    # To-Be side
    ax.add_patch(FancyBboxPatch((5.8, 1.0), 4.0, 2.5, boxstyle="round,pad=0.1", facecolor='#27ae60', alpha=0.15, edgecolor='#27ae60'))
    ax.text(7.8, 3.2, 'To-Be', ha='center', fontsize=11, fontweight='bold', color='#27ae60')
    for i, t in enumerate(['Validate ฝ่าย/สำนัก/ตำแหน่ง', 'RBAC + Per-user Permission', 'Full Audit Trail', 'Zoo Admin กระจายอำนาจ']):
        ax.text(7.8, 2.6 - i*0.45, t, ha='center', fontsize=9)
    fig.tight_layout()
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "เปรียบเทียบ As-Is vs To-Be ระบบจัดการเจ้าหน้าที่")

    # 3.6.1.6 Compatibility
    add_subsection_heading(doc, "3.6.1.6", "แผนความเข้ากันได้กับระบบเดิม (Compatibility Plan)")
    add_body_after_subsection(doc,
        "การปรับปรุงระบบจัดการเจ้าหน้าที่ใช้แนวทาง Additive Schema Changes โดยเพิ่มคอลัมน์สิทธิ์รายบุคคล "
        "(individual_permissions) และคอลัมน์โครงสร้างองค์กร (org_structure_id) เข้าไปในตารางผู้ใช้เดิม "
        "โดยไม่แก้ไขหรือลบคอลัมน์ใดในตารางเดิม ผู้ใช้ที่มีอยู่แล้วยังคงเข้าถึงระบบได้ตามสิทธิ์เดิม "
        "และจะได้รับการ Migrate สิทธิ์เป็นรูปแบบใหม่ในขั้นตอน Data Migration"
    )


# ══════════════════════════════════════════════
# MODULE 2: Animal Registry
# ══════════════════════════════════════════════

def build_module_animal_registry(doc):
    add_section_heading(doc, "3.6.2", "ระบบทะเบียนสัตว์ (Animal Registry) — TOR 4.6")

    add_body_after_section(doc,
        "ระบบทะเบียนสัตว์เป็นศูนย์กลางข้อมูลหลัก (Single Source of Truth) ของระบบ Thai Zoo ARK ทั้งหมด "
        "โดยทุกโมดูลอ้างอิงข้อมูลสัตว์ผ่าน Animal ID จากทะเบียนนี้ "
        "ปัญหาสำคัญของระบบเดิมคือผู้ใช้ทั่วไปสามารถลบข้อมูลสัตว์ได้โดยไม่มีกลไกควบคุม "
        "ซึ่งเป็นความเสี่ยงสูงต่อความสมบูรณ์ของฐานข้อมูล การปรับปรุงครั้งนี้จึงนำ Controlled Deletion มาใช้ "
        "โดยเฉพาะ Super Admin เท่านั้นที่สามารถลบข้อมูลได้ และทุกการลบต้องอ้างอิงใบคำร้อง (Change Request) "
        "พร้อมบันทึกเหตุผล วันเวลา และผู้ดำเนินการในรูปแบบ Soft Delete และ Audit Trail "
        "หัวข้อนี้ประกอบด้วย User Stories (หัวข้อ 3.6.2.1) การออกแบบหน้าจอ (หัวข้อ 3.6.2.2) "
        "กฎการตรวจสอบข้อมูล (หัวข้อ 3.6.2.3) และเกณฑ์การทดสอบ (หัวข้อ 3.6.2.4) "
        "ตารางที่ 3.6-3 แสดง User Stories ตารางที่ 3.6-4 แสดงเกณฑ์การทดสอบ "
        "และแผนภาพที่ 3.6-2 แสดงกระบวนการ Controlled Deletion"
    )

    add_subsection_heading(doc, "3.6.2.1", "User Stories")
    add_body_after_subsection(doc,
        "User Stories ของระบบทะเบียนสัตว์มุ่งเน้นการควบคุมการลบข้อมูลและการรักษาความสมบูรณ์ของข้อมูลอ้างอิง "
        "ดังแสดงในตารางที่ 3.6-3"
    )
    tn = next_table_num()
    add_table_caption(doc, tn, "User Stories ระบบทะเบียนสัตว์")
    add_table(doc,
        ["รหัส", "บทบาท", "ความต้องการ", "เป้าหมาย", "TOR"],
        [
            ["US-6.1", "Super Admin", "ต้องการลบข้อมูลสัตว์โดยอ้างอิงเลขที่ใบคำร้อง", "การลบมีหลักฐานตรวจสอบได้", "4.6"],
            ["US-6.2", "ผู้ใช้ทั่วไป", "ไม่สามารถลบข้อมูลสัตว์ได้", "ป้องกันข้อมูลสูญหาย", "4.6"],
            ["US-6.3", "ผู้ตรวจสอบ", "ต้องการดูประวัติการลบข้อมูลพร้อมเหตุผลและหลักฐาน", "Audit Trail ครบถ้วน", "4.6"],
        ],
        col_widths=[1.8, 2.5, 5.0, 4.0, 1.2]
    )

    add_subsection_heading(doc, "3.6.2.2", "การออกแบบหน้าจอและกระบวนการทำงาน")
    add_body_after_subsection(doc,
        "กระบวนการ Controlled Deletion ออกแบบให้มีขั้นตอนการตรวจสอบหลายชั้น "
        "เริ่มจากการส่งคำร้องผ่านระบบ Change Request ตรวจสอบสิทธิ์ Super Admin "
        "ยืนยันเหตุผลและเลขที่คำร้อง จากนั้นดำเนินการ Soft Delete พร้อมบันทึก Audit Log "
        "แผนภาพที่ 3.6-2 แสดงกระบวนการนี้โดยละเอียด"
    )
    fig = create_flow_diagram(
        "กระบวนการ Controlled Deletion ทะเบียนสัตว์ (TOR 4.6)",
        [
            (0.0, 2.5, 1.8, 1.0, '#2c3e50', 'ผู้ใช้ทั่วไป'),
            (2.2, 3.5, 2.0, 1.0, '#e74c3c', 'ปุ่มลบ\nถูกซ่อน'),
            (2.2, 1.5, 1.8, 1.0, '#2980b9', 'Super Admin'),
            (4.5, 1.5, 2.0, 1.0, '#8e44ad', 'ระบุเลขที่\nใบคำร้อง'),
            (7.0, 2.5, 2.0, 1.0, '#27ae60', 'Soft Delete\n+ เหตุผล'),
            (9.5, 2.5, 1.5, 1.0, '#16a085', 'Audit\nLog'),
        ],
        [(0,1), (2,3), (3,4), (4,5)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "กระบวนการ Controlled Deletion ทะเบียนสัตว์")

    add_subsection_heading(doc, "3.6.2.3", "กฎการตรวจสอบข้อมูล (Data Validations)")
    add_body_after_subsection(doc,
        "กฎการตรวจสอบข้อมูลสำหรับระบบทะเบียนสัตว์ออกแบบให้ป้องกันการสูญเสียข้อมูลโดยไม่ได้รับอนุญาต"
    )
    add_numbered_item_subsection(doc, 1, "การลบต้องมีสิทธิ์ Super Admin เท่านั้น")
    add_numbered_item_subsection(doc, 2, "ต้องระบุเลขที่ใบคำร้อง (CR Reference Number) จากโมดูล Change Request")
    add_numbered_item_subsection(doc, 3, "ต้องระบุเหตุผลในการลบ (Reason Text) ความยาวไม่น้อยกว่า 20 ตัวอักษร")
    add_numbered_item_subsection(doc, 4, "ใช้ Soft Delete โดยเปลี่ยนสถานะเป็น deleted ข้อมูลยังคงอยู่ในฐานข้อมูล")
    add_numbered_item_subsection(doc, 5, "บันทึก Audit Log: ผู้ลบ วันเวลา เลขที่ CR เหตุผล และข้อมูลก่อนลบ (Snapshot)")

    add_subsection_heading(doc, "3.6.2.4", "เกณฑ์การทดสอบและยอมรับ (Acceptance Tests)")
    tn = next_table_num()
    add_table_caption(doc, tn, "เกณฑ์การทดสอบระบบทะเบียนสัตว์")
    add_table(doc,
        ["รหัส", "รายการทดสอบ", "ผลลัพธ์ที่คาดหวัง"],
        [
            ["AT-6.1", "ผู้ใช้ทั่วไปพยายามลบข้อมูลสัตว์", "ปุ่มลบไม่แสดง หรือปฏิเสธการทำงาน"],
            ["AT-6.2", "Super Admin ลบข้อมูลพร้อม CR#", "Audit Log บันทึก CR# เหตุผล และ Timestamp"],
            ["AT-6.3", "ลบข้อมูลแล้วค้นหาในระบบ", "ข้อมูลไม่แสดงในหน้าจอปกติ แต่ยังอยู่ในฐานข้อมูล"],
        ],
        col_widths=[1.8, 6.0, 6.0]
    )

    # Gap Analysis
    add_subsection_heading(doc, "3.6.2.5", "การวิเคราะห์ช่องว่าง As-Is vs To-Be (Gap Analysis)")
    add_body_after_subsection(doc,
        "ระบบทะเบียนสัตว์เดิมมีความสามารถในการจัดการข้อมูลพื้นฐาน แต่มีช่องว่างสำคัญด้านการควบคุมการลบข้อมูล "
        "และการตรวจสอบย้อนหลัง ซึ่งเป็นความเสี่ยงสูงเนื่องจากทะเบียนสัตว์เป็น Single Source of Truth "
        "ตารางที่ {tn} แสดงช่องว่างและแนวทางแก้ไข "
        "แผนภาพที่ {dn} แสดงกระบวนการเปรียบเทียบก่อนและหลังปรับปรุง".format(
            tn="3.6-" + str(table_counter + 1), dn="3.6-" + str(diagram_counter + 1))
    )
    tn = next_table_num()
    add_table_caption(doc, tn, "Gap Analysis ระบบทะเบียนสัตว์: As-Is vs To-Be")
    add_table(doc,
        ["ด้าน", "As-Is (ระบบเดิม)", "Gap (ช่องว่าง)", "To-Be (เป้าหมาย)"],
        [
            ["การลบข้อมูล", "ผู้ใช้ทั่วไปลบได้โดยตรง", "ไม่มีกลไกควบคุม เสี่ยงต่อข้อมูลสูญหาย", "เฉพาะ Super Admin + อ้างอิง CR"],
            ["ประเภทการลบ", "Hard Delete (ลบถาวร)", "ข้อมูลหายไปจากระบบตลอดกาล", "Soft Delete (เปลี่ยนสถานะ)"],
            ["เหตุผลการลบ", "ไม่ต้องระบุเหตุผล", "ไม่มีหลักฐานว่าลบทำไม", "ต้องระบุเหตุผล + เลขที่คำร้อง"],
            ["Audit Trail", "ไม่บันทึกการลบ", "ตรวจสอบย้อนหลังไม่ได้", "บันทึกครบ: ผู้ลบ วันเวลา เหตุผล CR# Snapshot"],
        ],
        col_widths=[2.5, 3.5, 3.5, 4.0]
    )
    fig, ax = plt.subplots(1, 1, figsize=(10, 4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 4.5); ax.axis('off')
    ax.text(5, 4.2, 'Gap Analysis: ระบบทะเบียนสัตว์ — As-Is vs To-Be', ha='center', fontsize=12, fontweight='bold')
    ax.add_patch(FancyBboxPatch((0.2, 1.0), 4.0, 2.5, boxstyle="round,pad=0.1", facecolor='#e74c3c', alpha=0.15, edgecolor='#e74c3c'))
    ax.text(2.2, 3.2, 'As-Is', ha='center', fontsize=11, fontweight='bold', color='#e74c3c')
    for i, t in enumerate(['User ทั่วไปลบได้ (Hard Delete)', 'ไม่มีเหตุผลการลบ', 'ไม่มี Audit Trail', 'ไม่มี CR Reference']):
        ax.text(2.2, 2.6 - i*0.45, t, ha='center', fontsize=9)
    ax.annotate('', xy=(5.6, 2.2), xytext=(4.4, 2.2), arrowprops=dict(arrowstyle='->', color='#2c3e50', lw=2.5))
    ax.text(5.0, 2.6, 'ปรับปรุง', ha='center', fontsize=9, fontweight='bold', color='#2c3e50')
    ax.add_patch(FancyBboxPatch((5.8, 1.0), 4.0, 2.5, boxstyle="round,pad=0.1", facecolor='#27ae60', alpha=0.15, edgecolor='#27ae60'))
    ax.text(7.8, 3.2, 'To-Be', ha='center', fontsize=11, fontweight='bold', color='#27ae60')
    for i, t in enumerate(['Controlled Deletion (Super Admin)', 'Soft Delete + Reason Text', 'Full Audit Trail + Snapshot', 'CR Reference บังคับ']):
        ax.text(7.8, 2.6 - i*0.45, t, ha='center', fontsize=9)
    fig.tight_layout()
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "เปรียบเทียบ As-Is vs To-Be ระบบทะเบียนสัตว์")


# ══════════════════════════════════════════════
# MODULE 3: Nutrition Management
# ══════════════════════════════════════════════

def build_module_nutrition(doc):
    add_section_heading(doc, "3.6.3", "ระบบจัดการโภชนาการสัตว์ (Nutrition Management) — TOR 4.7")

    add_body_after_section(doc,
        "ระบบจัดการโภชนาการเป็นหนึ่งในโมดูลปฏิบัติการที่ผู้ดูแลสัตว์และนักโภชนาการใช้งานมากที่สุดในแต่ละวัน "
        "ปัญหาของระบบเดิมคือ UX การแก้ไขตารางอาหารยังไม่สะดวก และยังไม่มีกลไกบังคับให้เลือกอาหาร "
        "เฉพาะรายการที่อนุญาตตาม Diet Card ของสัตว์แต่ละชนิด "
        "การปรับปรุงจึงมุ่งเน้น Diet-Card-Driven Nutrition Entry ที่ให้เลือกได้เฉพาะรายการอาหารใน Diet Card "
        "พร้อม Inline Editing แบบ Click-to-Edit และ Validation เรื่องหน่วย ปริมาณ ตามมาตรฐาน "
        "นอกจากนี้ยังเชื่อมข้อมูลกับคลังอาหารเพื่อตรวจสอบสต็อกก่อนจัดเตรียมอาหาร "
        "หัวข้อนี้ประกอบด้วย User Stories (หัวข้อ 3.6.3.1) การออกแบบหน้าจอ (หัวข้อ 3.6.3.2) "
        "กฎการตรวจสอบข้อมูล (หัวข้อ 3.6.3.3) และเกณฑ์การทดสอบ (หัวข้อ 3.6.3.4) "
        "ตารางที่ 3.6-5 แสดง User Stories และแผนภาพที่ 3.6-3 แสดงกระบวนการบันทึกข้อมูลโภชนาการ"
    )

    add_subsection_heading(doc, "3.6.3.1", "User Stories")
    tn = next_table_num()
    add_table_caption(doc, tn, "User Stories ระบบจัดการโภชนาการสัตว์")
    add_table(doc,
        ["รหัส", "บทบาท", "ความต้องการ", "เป้าหมาย", "TOR"],
        [
            ["US-7.1", "นักโภชนาการ", "ต้องการ Inline Editing บนตารางอาหาร", "แก้ไขข้อมูลได้รวดเร็ว", "4.7"],
            ["US-7.2", "ระบบ", "จำกัดประเภทอาหารเฉพาะที่ระบุใน Diet Card", "การให้อาหารถูกต้องตามมาตรฐาน", "4.7"],
            ["US-7.3", "นักโภชนาการ", "ต้องการเห็นคำเตือนเมื่อปริมาณเกิน/ต่ำกว่า Diet Card", "ป้องกันการให้อาหารผิดพลาด", "4.7"],
        ],
        col_widths=[1.8, 2.5, 5.5, 3.5, 1.2]
    )

    add_subsection_heading(doc, "3.6.3.2", "การออกแบบหน้าจอและกระบวนการทำงาน")
    add_body_after_subsection(doc,
        "หน้าจอระบบโภชนาการออกแบบเป็นตาราง Inline Editing ที่ผู้ใช้คลิกเซลล์เพื่อแก้ไข "
        "ระบบโหลด Diet Card Rules อัตโนมัติเมื่อเลือกสัตว์ และแสดงรายการอาหารที่อนุญาตใน Dropdown "
        "หากเลือกอาหารนอก Diet Card ระบบจะปฏิเสธพร้อมแสดงคำเตือน "
        "แผนภาพที่ 3.6-3 แสดงกระบวนการทำงานตั้งแต่การเลือกสัตว์จนถึงการบันทึกข้อมูลโภชนาการ"
    )
    fig = create_flow_diagram(
        "กระบวนการบันทึกข้อมูลโภชนาการแบบ Diet-Card-Driven (TOR 4.7)",
        [
            (0.0, 2.5, 2.0, 1.0, '#2c3e50', 'ผู้ดูแลสัตว์\n/ นักโภชนาการ'),
            (2.5, 2.5, 2.0, 1.0, '#2980b9', 'เลือกสัตว์\nโหลด Diet Card'),
            (5.0, 3.5, 2.0, 1.0, '#27ae60', 'อาหาร\nอนุญาต?'),
            (5.0, 1.5, 2.0, 1.0, '#e74c3c', 'ปฏิเสธ\n+ คำเตือน'),
            (7.5, 3.5, 2.0, 1.0, '#8e44ad', 'บันทึก\nFeeding Entry'),
            (7.5, 1.5, 2.0, 1.0, '#16a085', 'อัปเดตสต็อก\nคลังอาหาร'),
        ],
        [(0,1), (1,2), (2,4), (2,3), (4,5)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "กระบวนการบันทึกข้อมูลโภชนาการแบบ Diet-Card-Driven")

    add_subsection_heading(doc, "3.6.3.3", "กฎการตรวจสอบข้อมูล (Data Validations)")
    add_numbered_item_subsection(doc, 1, "ประเภทอาหารต้องมีอยู่ใน Diet Card ของสัตว์ชนิดนั้น")
    add_numbered_item_subsection(doc, 2, "แจ้งเตือน (Warning) เมื่อปริมาณเกินหรือต่ำกว่าช่วงที่ Diet Card กำหนด")
    add_numbered_item_subsection(doc, 3, "หน่วยวัด (Unit of Measure) ต้องตรงกับหน่วยที่กำหนดใน Diet Card")
    add_numbered_item_subsection(doc, 4, "วันที่บันทึกต้องไม่เป็นวันในอนาคต")

    add_subsection_heading(doc, "3.6.3.4", "เกณฑ์การทดสอบและยอมรับ (Acceptance Tests)")
    tn = next_table_num()
    add_table_caption(doc, tn, "เกณฑ์การทดสอบระบบจัดการโภชนาการสัตว์")
    add_table(doc,
        ["รหัส", "รายการทดสอบ", "ผลลัพธ์ที่คาดหวัง"],
        [
            ["AT-7.1", "เลือกอาหารที่ไม่อยู่ใน Diet Card", "ระบบปฏิเสธพร้อมแสดงรายการที่อนุญาต"],
            ["AT-7.2", "ใช้ Inline Editing แก้ไขปริมาณอาหาร", "เซลล์เปลี่ยนเป็น Editable และบันทึกได้ทันที"],
            ["AT-7.3", "ป้อนปริมาณเกินช่วง Diet Card", "แสดงคำเตือนสีเหลืองแต่อนุญาตให้บันทึกพร้อมเหตุผล"],
        ],
        col_widths=[1.8, 6.0, 6.0]
    )

    # Gap Analysis
    add_subsection_heading(doc, "3.6.3.5", "การวิเคราะห์ช่องว่าง As-Is vs To-Be (Gap Analysis)")
    tn = next_table_num()
    add_table_caption(doc, tn, "Gap Analysis ระบบจัดการโภชนาการ: As-Is vs To-Be")
    add_table(doc,
        ["ด้าน", "As-Is (ระบบเดิม)", "Gap (ช่องว่าง)", "To-Be (เป้าหมาย)"],
        [
            ["UX การแก้ไข", "แก้ไขตารางยาก ต้องเปิดฟอร์มใหม่", "ทำงานช้า ผู้ใช้ไม่สะดวก", "Inline Editing แบบ Click-to-Edit"],
            ["Diet Card", "ไม่มี Rule บังคับเลือกอาหาร", "ให้อาหารผิดประเภทได้", "Diet-Card-Driven เลือกได้เฉพาะที่อนุญาต"],
            ["Validation", "ไม่ตรวจสอบปริมาณ/หน่วย", "ปริมาณผิดพลาดไม่มีการเตือน", "Validate หน่วย ปริมาณ ช่วงที่กำหนด"],
            ["เชื่อมคลังอาหาร", "ไม่เชื่อมกับระบบคลังอาหาร", "ไม่ทราบสต็อกก่อนจัดเตรียม", "เชื่อมข้อมูลสต็อกแบบ Real-Time"],
        ],
        col_widths=[2.5, 3.5, 3.5, 4.0]
    )
    fig, ax = plt.subplots(1, 1, figsize=(10, 4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 4.5); ax.axis('off')
    ax.text(5, 4.2, 'Gap Analysis: ระบบจัดการโภชนาการ — As-Is vs To-Be', ha='center', fontsize=12, fontweight='bold')
    ax.add_patch(FancyBboxPatch((0.2, 1.0), 4.0, 2.5, boxstyle="round,pad=0.1", facecolor='#e74c3c', alpha=0.15, edgecolor='#e74c3c'))
    ax.text(2.2, 3.2, 'As-Is', ha='center', fontsize=11, fontweight='bold', color='#e74c3c')
    for i, t in enumerate(['UX ตารางแก้ไขยาก', 'ไม่มี Diet Card Rule', 'ไม่ Validate ปริมาณ/หน่วย', 'ไม่เชื่อมคลังอาหาร']):
        ax.text(2.2, 2.6 - i*0.45, t, ha='center', fontsize=9)
    ax.annotate('', xy=(5.6, 2.2), xytext=(4.4, 2.2), arrowprops=dict(arrowstyle='->', color='#2c3e50', lw=2.5))
    ax.text(5.0, 2.6, 'ปรับปรุง', ha='center', fontsize=9, fontweight='bold', color='#2c3e50')
    ax.add_patch(FancyBboxPatch((5.8, 1.0), 4.0, 2.5, boxstyle="round,pad=0.1", facecolor='#27ae60', alpha=0.15, edgecolor='#27ae60'))
    ax.text(7.8, 3.2, 'To-Be', ha='center', fontsize=11, fontweight='bold', color='#27ae60')
    for i, t in enumerate(['Inline Editing (Click-to-Edit)', 'Diet-Card-Driven Selection', 'Validate หน่วย+ปริมาณ+ช่วง', 'เชื่อมคลังอาหาร Real-Time']):
        ax.text(7.8, 2.6 - i*0.45, t, ha='center', fontsize=9)
    fig.tight_layout()
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "เปรียบเทียบ As-Is vs To-Be ระบบจัดการโภชนาการ")


# ══════════════════════════════════════════════
# MODULE 4: Pharma / Medical Supplies
# ══════════════════════════════════════════════

def build_module_pharma(doc):
    add_section_heading(doc, "3.6.4", "ระบบบริหารการเก็บและใช้ยา/เวชภัณฑ์ (Pharmacy & Medical Supplies) — TOR 4.8, 4.12, 4.13")

    add_body_after_section(doc,
        "ระบบบริหารยาและเวชภัณฑ์ครอบคลุมข้อกำหนด TOR 3 ข้อ ได้แก่ ข้อ 4.8 การจัดการยาเพิ่มเติม "
        "ข้อ 4.12 Responsive UI และข้อ 4.13 Barcode สำหรับยาและเวชภัณฑ์ "
        "ระบบเดิมรองรับการรับเข้า เบิก และคืนยาอยู่แล้ว แต่ยังขาดการเก็บข้อมูลแหล่งที่มาของยา (บริจาค/จัดซื้อ) "
        "ไม่มี Responsive UI สำหรับหน้างาน และไม่มีการใช้ Barcode อย่างเป็นระบบ "
        "การปรับปรุงจึงยกระดับเป็น Medication Inventory + Traceability ที่บันทึก Source, Lot/Batch, Expiry, "
        "Storage Location รองรับ Barcode ทั้งรับเข้า เบิก คืน และพิมพ์ฉลาก "
        "พร้อมเชื่อมข้อมูลไปยังเวชระเบียนและระบบคิดค่ารักษา "
        "ตารางที่ 3.6-7 แสดง User Stories ตารางที่ 3.6-8 แสดงเกณฑ์การทดสอบ "
        "และแผนภาพที่ 3.6-4 แสดงกระบวนการ Barcode Workflow สำหรับยาและเวชภัณฑ์"
    )

    add_subsection_heading(doc, "3.6.4.1", "User Stories")
    tn = next_table_num()
    add_table_caption(doc, tn, "User Stories ระบบบริหารยาและเวชภัณฑ์")
    add_table(doc,
        ["รหัส", "บทบาท", "ความต้องการ", "เป้าหมาย", "TOR"],
        [
            ["US-8.1", "เจ้าหน้าที่คลัง", "ต้องการบันทึกแหล่งที่มา (บริจาค/จัดซื้อ)", "คิดค่ารักษาถูกต้อง", "4.8"],
            ["US-8.2", "เจ้าหน้าที่คลัง", "ต้องการสแกน Barcode เพื่อรับเข้า/เบิก/คืน", "ทำงานรวดเร็วและแม่นยำ", "4.13"],
            ["US-8.3", "เจ้าหน้าที่คลัง", "ต้องการพิมพ์ฉลาก Barcode", "ติดตามยาได้ทุกขั้นตอน", "4.13"],
            ["US-8.4", "สัตวแพทย์", "ต้องการใช้ Responsive UI บนหน้างาน", "ทำงานได้ทุกอุปกรณ์", "4.12"],
        ],
        col_widths=[1.8, 2.5, 5.0, 3.5, 1.2]
    )

    add_subsection_heading(doc, "3.6.4.2", "การออกแบบหน้าจอและกระบวนการทำงาน")
    add_body_after_subsection(doc,
        "หน้าจอระบบยาและเวชภัณฑ์ออกแบบเป็น 4 หน้าจอหลัก ได้แก่ หน้ารับเข้า (Receive) ที่รองรับการสแกนหรือป้อนข้อมูลด้วยมือ "
        "หน้าเบิกยา (Issue) ด้วย Barcode Scanner หน้าคืนยา (Return) และหน้าสต็อกคงเหลือ (Stock Remaining) "
        "ทุกหน้าจอรองรับ Responsive Design สำหรับ Desktop Tablet และ Mobile "
        "แผนภาพที่ 3.6-4 แสดง Barcode Workflow ตั้งแต่การรับเข้า สร้าง Barcode พิมพ์ฉลาก จนถึงการเบิกจ่าย"
    )
    fig = create_flow_diagram(
        "กระบวนการ Barcode Workflow สำหรับยาและเวชภัณฑ์ (TOR 4.8, 4.12, 4.13)",
        [
            (0.0, 3.0, 2.0, 1.0, '#2c3e50', 'เจ้าหน้าที่คลัง'),
            (2.5, 3.8, 2.0, 1.0, '#2980b9', 'รับเข้ายา\n+ บันทึก Source'),
            (2.5, 2.0, 2.0, 1.0, '#27ae60', 'สร้าง/สแกน\nBarcode'),
            (5.2, 3.8, 2.0, 1.0, '#8e44ad', 'พิมพ์ฉลาก\nBarcode'),
            (5.2, 2.0, 2.0, 1.0, '#e67e22', 'Drug\nInventory DB'),
            (7.8, 3.0, 2.0, 1.0, '#e74c3c', 'สัตวแพทย์\nเบิก/คืน'),
            (10.0, 3.0, 1.0, 1.0, '#16a085', 'Billing'),
        ],
        [(0,1), (0,2), (1,3), (2,4), (5,2), (5,6)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "กระบวนการ Barcode Workflow สำหรับยาและเวชภัณฑ์")

    add_subsection_heading(doc, "3.6.4.3", "กฎการตรวจสอบข้อมูล (Data Validations)")
    add_numbered_item_subsection(doc, 1, "แหล่งที่มา (Source) ต้องระบุ: บริจาค หรือ จัดซื้อ (ห้ามเว้นว่าง)")
    add_numbered_item_subsection(doc, 2, "Lot/Batch Number ต้องไม่ซ้ำกันในระบบ")
    add_numbered_item_subsection(doc, 3, "วันหมดอายุ (Expiry Date) ต้องเป็นวันในอนาคต ณ เวลารับเข้า")
    add_numbered_item_subsection(doc, 4, "Barcode ต้องเป็นรูปแบบมาตรฐาน (EAN-13 หรือ Code 128)")
    add_numbered_item_subsection(doc, 5, "จำนวนเบิกต้องไม่เกินสต็อกคงเหลือ")

    add_subsection_heading(doc, "3.6.4.4", "เกณฑ์การทดสอบและยอมรับ (Acceptance Tests)")
    tn = next_table_num()
    add_table_caption(doc, tn, "เกณฑ์การทดสอบระบบบริหารยาและเวชภัณฑ์")
    add_table(doc,
        ["รหัส", "รายการทดสอบ", "ผลลัพธ์ที่คาดหวัง"],
        [
            ["AT-8.1", "รับเข้ายาโดยไม่ระบุแหล่งที่มา", "ระบบปฏิเสธ แสดงข้อความให้ระบุ Source"],
            ["AT-8.2", "สแกน Barcode เบิกยา", "แสดงข้อมูลยาถูกต้อง ลดสต็อกอัตโนมัติ"],
            ["AT-8.3", "พิมพ์ฉลาก Barcode", "ฉลากมี Barcode, ชื่อยา, Lot, Expiry ครบถ้วน"],
            ["AT-8.4", "ใช้งานบน Tablet", "หน้าจอ Responsive ปุ่มใช้งานได้สะดวก"],
        ],
        col_widths=[1.8, 6.0, 6.0]
    )

    # Gap Analysis
    add_subsection_heading(doc, "3.6.4.5", "การวิเคราะห์ช่องว่าง As-Is vs To-Be (Gap Analysis)")
    tn = next_table_num()
    add_table_caption(doc, tn, "Gap Analysis ระบบยา/เวชภัณฑ์: As-Is vs To-Be")
    add_table(doc,
        ["ด้าน", "As-Is (ระบบเดิม)", "Gap (ช่องว่าง)", "To-Be (เป้าหมาย)"],
        [
            ["แหล่งที่มายา", "ไม่แยกบริจาค/จัดซื้อ", "คิดค่ารักษาไม่ถูกต้อง", "บันทึก Source: บริจาค/จัดซื้อ ทุกรายการ"],
            ["Barcode", "ไม่มีระบบ Barcode", "ค้นหายาด้วยมือ ช้าและผิดพลาด", "Barcode สแกนรับ/เบิก/คืน + พิมพ์ฉลาก"],
            ["Responsive", "Desktop Only", "ใช้งานหน้าคลังไม่สะดวก", "Responsive: Desktop/Tablet/Mobile"],
            ["Lot Tracking", "ไม่บันทึก Lot/Batch", "ไม่สามารถเรียกคืนยาตาม Lot ได้", "Lot/Batch + Expiry + Storage Location"],
        ],
        col_widths=[2.5, 3.5, 3.5, 4.0]
    )
    fig, ax = plt.subplots(1, 1, figsize=(10, 4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 4.5); ax.axis('off')
    ax.text(5, 4.2, 'Gap Analysis: ระบบยา/เวชภัณฑ์ — As-Is vs To-Be', ha='center', fontsize=12, fontweight='bold')
    ax.add_patch(FancyBboxPatch((0.2, 1.0), 4.0, 2.5, boxstyle="round,pad=0.1", facecolor='#e74c3c', alpha=0.15, edgecolor='#e74c3c'))
    ax.text(2.2, 3.2, 'As-Is', ha='center', fontsize=11, fontweight='bold', color='#e74c3c')
    for i, t in enumerate(['ไม่แยก Source (บริจาค/จัดซื้อ)', 'ไม่มี Barcode', 'Desktop Only', 'ไม่มี Lot/Batch Tracking']):
        ax.text(2.2, 2.6 - i*0.45, t, ha='center', fontsize=9)
    ax.annotate('', xy=(5.6, 2.2), xytext=(4.4, 2.2), arrowprops=dict(arrowstyle='->', color='#2c3e50', lw=2.5))
    ax.text(5.0, 2.6, 'ปรับปรุง', ha='center', fontsize=9, fontweight='bold', color='#2c3e50')
    ax.add_patch(FancyBboxPatch((5.8, 1.0), 4.0, 2.5, boxstyle="round,pad=0.1", facecolor='#27ae60', alpha=0.15, edgecolor='#27ae60'))
    ax.text(7.8, 3.2, 'To-Be', ha='center', fontsize=11, fontweight='bold', color='#27ae60')
    for i, t in enumerate(['Source Tracking ทุกรายการ', 'Barcode Scan + Print Label', 'Responsive (Desktop/Tablet/Mobile)', 'Lot/Batch + Expiry + Location']):
        ax.text(7.8, 2.6 - i*0.45, t, ha='center', fontsize=9)
    fig.tight_layout()
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "เปรียบเทียบ As-Is vs To-Be ระบบยา/เวชภัณฑ์")


# ══════════════════════════════════════════════
# MODULE 5: Food Warehouse
# ══════════════════════════════════════════════

def build_module_food_warehouse(doc):
    add_section_heading(doc, "3.6.5", "ระบบคลังอาหารสัตว์ (Food Warehouse) — TOR 4.9, 4.10")

    add_body_after_section(doc,
        "ระบบคลังอาหารสัตว์ครอบคลุมข้อกำหนด TOR ข้อ 4.9 (หน้าจอ Touch Screen + Responsive) "
        "และข้อ 4.10 (Barcode สำหรับคลังอาหาร) โดยเป็นโมดูลที่ต้องใช้งานจริงในพื้นที่คลังสินค้า "
        "ผ่าน Touch Screen ขนาด 21.5 นิ้ว ที่จัดหาตาม TOR ข้อ 4.28-4.34 "
        "ระบบเดิมรองรับการรับเข้า เบิก คืน และตรวจสอบสต็อกอยู่แล้ว แต่ยังไม่เหมาะกับ Touch Screen และ Barcode "
        "การปรับปรุงจึงออกแบบเป็น Touch-Enabled Warehouse Workflow ที่มีปุ่มขนาดใหญ่ (ไม่น้อยกว่า 44x44 พิกเซล) "
        "รองรับ 5 หน้าจอหลักตาม TOR ได้แก่ รับเข้าวัตถุดิบ (4.9.1) เบิกวัตถุดิบ (4.9.2) คืนวัตถุดิบ (4.9.3) "
        "Stock Card (4.9.4) และรายการสต็อกคงเหลือ (4.9.5) "
        "พร้อมระบบ Barcode ที่ครอบคลุมการสแกนรับเข้า (4.10.1) สร้างและพิมพ์ฉลาก (4.10.2) และหน้าจัดเก็บข้อมูล (4.10.3) "
        "ตารางที่ 3.6-9 แสดง User Stories และแผนภาพที่ 3.6-5 แสดง Touch Workflow ของคลังอาหาร"
    )

    add_subsection_heading(doc, "3.6.5.1", "User Stories")
    tn = next_table_num()
    add_table_caption(doc, tn, "User Stories ระบบคลังอาหารสัตว์")
    add_table(doc,
        ["รหัส", "บทบาท", "ความต้องการ", "เป้าหมาย", "TOR"],
        [
            ["US-9.1", "เจ้าหน้าที่คลัง", "ต้องการหน้าจอ Touch-Friendly ที่คลังสินค้า", "ใช้งานได้สะดวกด้วย Touch Screen", "4.9"],
            ["US-9.2", "เจ้าหน้าที่คลัง", "ต้องการสแกน Barcode เพื่อรับเข้าวัตถุดิบ", "รับเข้ารวดเร็วและแม่นยำ", "4.10"],
            ["US-9.3", "เจ้าหน้าที่คลัง", "ต้องการพิมพ์ฉลาก Barcode สำหรับสินค้าใหม่", "ติดตามสินค้าได้ทุกรายการ", "4.10"],
            ["US-9.4", "เจ้าหน้าที่คลัง", "ต้องการดู Stock Card และสต็อกคงเหลือ", "บริหารสต็อกได้มีประสิทธิภาพ", "4.9"],
        ],
        col_widths=[1.8, 2.5, 5.0, 3.5, 1.2]
    )

    add_subsection_heading(doc, "3.6.5.2", "การออกแบบหน้าจอและกระบวนการทำงาน")
    add_body_after_subsection(doc,
        "หน้าจอคลังอาหารออกแบบแบบ Touch-First โดยมีปุ่มขนาดใหญ่สำหรับ 5 ฟังก์ชันหลัก "
        "เมนูหลักเป็นแบบ Grid Layout ที่แตะเลือกได้ง่าย รองรับทั้ง Touch Screen 21.5 นิ้ว "
        "Tablet และ Desktop แผนภาพที่ 3.6-5 แสดงกระบวนการทำงานของคลังอาหาร"
    )
    fig = create_flow_diagram(
        "กระบวนการ Touch Workflow คลังอาหารสัตว์ (TOR 4.9, 4.10)",
        [
            (0.0, 2.5, 2.0, 1.0, '#2c3e50', 'เจ้าหน้าที่คลัง\nTouch Screen'),
            (2.5, 4.0, 1.8, 0.8, '#2980b9', 'รับเข้า\n(4.9.1)'),
            (2.5, 3.0, 1.8, 0.8, '#27ae60', 'เบิก\n(4.9.2)'),
            (2.5, 2.0, 1.8, 0.8, '#e67e22', 'คืน\n(4.9.3)'),
            (2.5, 1.0, 1.8, 0.8, '#8e44ad', 'Stock Card\n(4.9.4)'),
            (5.0, 3.0, 2.0, 1.0, '#e74c3c', 'สแกน\nBarcode'),
            (7.5, 3.0, 2.0, 1.0, '#16a085', 'Food\nInventory DB'),
            (7.5, 1.0, 2.0, 1.0, '#d4ac0d', 'พิมพ์ฉลาก\nBarcode'),
        ],
        [(0,1), (0,2), (0,3), (0,4), (1,5), (2,5), (3,5), (5,6), (6,7)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "กระบวนการ Touch Workflow คลังอาหารสัตว์")

    add_subsection_heading(doc, "3.6.5.3", "กฎการตรวจสอบข้อมูล (Data Validations)")
    add_numbered_item_subsection(doc, 1, "รหัสสินค้า (Product Code) ต้องมีอยู่ใน Product Master")
    add_numbered_item_subsection(doc, 2, "จำนวนเบิกต้องไม่เกินสต็อกคงเหลือ (On-Hand Inventory)")
    add_numbered_item_subsection(doc, 3, "หน่วยวัดต้องตรงกับหน่วยที่กำหนดใน Product Master")
    add_numbered_item_subsection(doc, 4, "Barcode ต้องสแกนได้และตรงกับรายการในฐานข้อมูล")
    add_numbered_item_subsection(doc, 5, "Storage Location ต้องระบุสำหรับทุกรายการรับเข้า")

    add_subsection_heading(doc, "3.6.5.4", "เกณฑ์การทดสอบและยอมรับ (Acceptance Tests)")
    tn = next_table_num()
    add_table_caption(doc, tn, "เกณฑ์การทดสอบระบบคลังอาหารสัตว์")
    add_table(doc,
        ["รหัส", "รายการทดสอบ", "ผลลัพธ์ที่คาดหวัง"],
        [
            ["AT-9.1", "ใช้งานบน Touch Screen 21.5 นิ้ว", "ปุ่มขนาดใหญ่ สัมผัสได้ง่าย ไม่กดผิด"],
            ["AT-9.2", "สแกน Barcode รับเข้าอาหาร", "ข้อมูลสินค้าแสดงถูกต้อง สต็อกเพิ่มอัตโนมัติ"],
            ["AT-9.3", "พิมพ์ฉลาก Barcode ผ่าน Thermal Printer", "ฉลากอ่านได้ด้วย Scanner"],
            ["AT-9.4", "ดู Stock Card แสดงความเคลื่อนไหว", "แสดงรายการรับ-เบิก-คืน พร้อมยอดคงเหลือ"],
        ],
        col_widths=[1.8, 6.0, 6.0]
    )

    # Gap Analysis
    add_subsection_heading(doc, "3.6.5.5", "การวิเคราะห์ช่องว่าง As-Is vs To-Be (Gap Analysis)")
    tn = next_table_num()
    add_table_caption(doc, tn, "Gap Analysis ระบบคลังอาหาร: As-Is vs To-Be")
    add_table(doc,
        ["ด้าน", "As-Is (ระบบเดิม)", "Gap (ช่องว่าง)", "To-Be (เป้าหมาย)"],
        [
            ["Touch Screen", "ไม่รองรับ Touch Screen", "ใช้งานจริงหน้าคลังไม่ได้", "Touch-First UI ปุ่มใหญ่ 44x44px+"],
            ["Barcode", "ไม่มีระบบ Barcode", "รับเข้าด้วยมือ ช้าและผิดพลาด", "Scan รับเข้า + สร้าง + พิมพ์ฉลาก"],
            ["Responsive", "Desktop Only", "ไม่เหมาะกับหน้างานจริง", "Touch 21.5\"/Tablet/Desktop"],
            ["Product Master", "ไม่มีระบบ Product Master", "ข้อมูลสินค้าไม่เป็นมาตรฐาน", "Product Master + UoM + Storage Location"],
        ],
        col_widths=[2.5, 3.5, 3.5, 4.0]
    )
    fig, ax = plt.subplots(1, 1, figsize=(10, 4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 4.5); ax.axis('off')
    ax.text(5, 4.2, 'Gap Analysis: ระบบคลังอาหาร — As-Is vs To-Be', ha='center', fontsize=12, fontweight='bold')
    ax.add_patch(FancyBboxPatch((0.2, 1.0), 4.0, 2.5, boxstyle="round,pad=0.1", facecolor='#e74c3c', alpha=0.15, edgecolor='#e74c3c'))
    ax.text(2.2, 3.2, 'As-Is', ha='center', fontsize=11, fontweight='bold', color='#e74c3c')
    for i, t in enumerate(['ไม่รองรับ Touch Screen', 'ไม่มี Barcode', 'Desktop Only', 'ไม่มี Product Master']):
        ax.text(2.2, 2.6 - i*0.45, t, ha='center', fontsize=9)
    ax.annotate('', xy=(5.6, 2.2), xytext=(4.4, 2.2), arrowprops=dict(arrowstyle='->', color='#2c3e50', lw=2.5))
    ax.text(5.0, 2.6, 'ปรับปรุง', ha='center', fontsize=9, fontweight='bold', color='#2c3e50')
    ax.add_patch(FancyBboxPatch((5.8, 1.0), 4.0, 2.5, boxstyle="round,pad=0.1", facecolor='#27ae60', alpha=0.15, edgecolor='#27ae60'))
    ax.text(7.8, 3.2, 'To-Be', ha='center', fontsize=11, fontweight='bold', color='#27ae60')
    for i, t in enumerate(['Touch-First UI (44x44px+)', 'Barcode Scan + Print Label', 'Touch 21.5\"/Tablet/Desktop', 'Product Master + UoM + Location']):
        ax.text(7.8, 2.6 - i*0.45, t, ha='center', fontsize=9)
    fig.tight_layout()
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "เปรียบเทียบ As-Is vs To-Be ระบบคลังอาหาร")


# ══════════════════════════════════════════════
# MODULE 6: Zookeeper Module
# ══════════════════════════════════════════════

def build_module_keeper(doc):
    add_section_heading(doc, "3.6.6", "ระบบผู้ดูแลสัตว์ (Zookeeper Mobile Module) — TOR 4.11")

    add_body_after_section(doc,
        "ระบบผู้ดูแลสัตว์เป็นโมดูลที่ออกแบบสำหรับการใช้งานภาคสนามผ่านอุปกรณ์เคลื่อนที่ "
        "โดยผู้ดูแลสัตว์ซึ่งเป็นกลุ่มผู้ใช้จำนวนมากที่สุด (ประมาณ 100-200 คนทั่วทุกสวนสัตว์) "
        "ต้องบันทึกข้อมูลประจำวันในพื้นที่เลี้ยงสัตว์จริง ระบบเดิมยังไม่รองรับ Mobile Device "
        "และยังขาดประสบการณ์ใช้งานแบบ Quick Actions สำหรับหน้างาน "
        "การปรับปรุงจึงออกแบบเป็น Mobile-First Keeper Operations ที่แสดงรายการสัตว์ที่รับผิดชอบ "
        "รองรับ 9 ฟังก์ชันตาม TOR 4.11.1-4.11.9 ได้แก่ รายชื่อสัตว์ รายการรับอาหาร สภาพแวดล้อม "
        "ส่งเสริมคุณภาพชีวิต แจ้งเกิด แจ้งป่วย แจ้งตั้งครรภ์/วางไข่ แจ้งตาย และแจ้งสูญหาย "
        "ตารางที่ 3.6-11 แสดง User Stories ตารางที่ 3.6-12 แสดงฟังก์ชัน 9 รายการ "
        "และแผนภาพที่ 3.6-6 แสดงกระบวนการทำงานของผู้ดูแลสัตว์ผ่าน Mobile"
    )

    add_subsection_heading(doc, "3.6.6.1", "User Stories")
    tn = next_table_num()
    add_table_caption(doc, tn, "User Stories ระบบผู้ดูแลสัตว์")
    add_table(doc,
        ["รหัส", "บทบาท", "ความต้องการ", "เป้าหมาย", "TOR"],
        [
            ["US-11.1", "ผู้ดูแลสัตว์", "ต้องการเข้าถึงรายชื่อสัตว์ที่ดูแลผ่าน Mobile", "ทำงานภาคสนามได้สะดวก", "4.11"],
            ["US-11.2", "ผู้ดูแลสัตว์", "ต้องการแจ้งเกิด/ป่วย/ตายผ่าน Mobile ได้ทันที", "สัตวแพทย์ตอบสนองได้รวดเร็ว", "4.11"],
            ["US-11.3", "ผู้ดูแลสัตว์", "ต้องการบันทึกสภาพแวดล้อมและกิจกรรม Enrichment", "ข้อมูลสวัสดิภาพครบถ้วน", "4.11"],
        ],
        col_widths=[1.8, 2.5, 5.5, 3.5, 1.2]
    )

    add_subsection_heading(doc, "3.6.6.2", "ฟังก์ชัน 9 รายการตาม TOR 4.11")
    add_body_after_subsection(doc,
        "ระบบผู้ดูแลสัตว์รองรับฟังก์ชัน 9 รายการตามข้อกำหนด TOR 4.11.1-4.11.9 "
        "โดยทุกฟังก์ชันออกแบบให้ใช้งานได้ผ่าน Smartphone และ Tablet "
        "ดังแสดงในตารางที่ 3.6-12"
    )
    tn = next_table_num()
    add_table_caption(doc, tn, "ฟังก์ชัน 9 รายการของระบบผู้ดูแลสัตว์")
    add_table(doc,
        ["ลำดับ", "ฟังก์ชัน", "รายละเอียด", "TOR"],
        [
            ["1", "รายชื่อสัตว์ที่ดูแล", "แสดงรายการสัตว์ที่ผู้ดูแลรับผิดชอบ พร้อมรูปภาพและสถานะ", "4.11.1"],
            ["2", "รายการรับอาหาร", "บันทึกการรับอาหารจากคลัง ตรวจสอบกับ Diet Card", "4.11.2"],
            ["3", "สภาพแวดล้อม", "บันทึกอุณหภูมิ ความชื้น ความสะอาดของพื้นที่เลี้ยง", "4.11.3"],
            ["4", "ส่งเสริมคุณภาพชีวิตสัตว์", "บันทึกกิจกรรม Enrichment และการตอบสนองของสัตว์", "4.11.4"],
            ["5", "การแจ้งเกิด", "แจ้งเหตุการณ์เกิดใหม่ พร้อมข้อมูลพ่อ-แม่ เพศ น้ำหนัก", "4.11.5"],
            ["6", "การแจ้งป่วย", "แจ้งอาการป่วย ส่งแจ้งเตือนถึงสัตวแพทย์อัตโนมัติ", "4.11.6"],
            ["7", "การแจ้งตั้งครรภ์/วางไข่", "บันทึกข้อมูลการตั้งครรภ์หรือวางไข่ พร้อมกำหนดคลอด", "4.11.7"],
            ["8", "การแจ้งตาย", "แจ้งเหตุการณ์ตาย อัปเดตสถานะในทะเบียนสัตว์อัตโนมัติ", "4.11.8"],
            ["9", "การแจ้งสูญหาย", "แจ้งสัตว์สูญหาย ส่งแจ้งเตือนถึงผู้บริหาร", "4.11.9"],
        ],
        col_widths=[1.0, 3.5, 6.5, 1.2]
    )

    add_subsection_heading(doc, "3.6.6.3", "การออกแบบหน้าจอและกระบวนการทำงาน")
    fig = create_flow_diagram(
        "กระบวนการทำงานผู้ดูแลสัตว์ผ่าน Mobile (TOR 4.11)",
        [
            (0.0, 2.5, 2.0, 1.0, '#2c3e50', 'ผู้ดูแลสัตว์\nMobile'),
            (2.5, 4.0, 2.0, 0.8, '#2980b9', 'สัตว์ที่ดูแล\n(My Animals)'),
            (2.5, 3.0, 2.0, 0.8, '#27ae60', 'รับอาหาร\n(Feed)'),
            (2.5, 2.0, 2.0, 0.8, '#e67e22', 'สภาพแวดล้อม\n(Environment)'),
            (2.5, 1.0, 2.0, 0.8, '#8e44ad', 'Enrichment'),
            (5.5, 3.5, 2.0, 1.0, '#e74c3c', 'แจ้งเหตุการณ์\nเกิด/ป่วย/ตาย/สูญหาย'),
            (8.0, 3.5, 2.0, 1.0, '#16a085', 'แจ้งเตือน\nสัตวแพทย์/ผู้บริหาร'),
            (8.0, 1.5, 2.0, 1.0, '#d4ac0d', 'Animal Care\nDB'),
        ],
        [(0,1), (0,2), (0,3), (0,4), (1,5), (5,6), (2,7), (3,7), (4,7)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "กระบวนการทำงานผู้ดูแลสัตว์ผ่าน Mobile")

    add_subsection_heading(doc, "3.6.6.4", "เกณฑ์การทดสอบและยอมรับ (Acceptance Tests)")
    tn = next_table_num()
    add_table_caption(doc, tn, "เกณฑ์การทดสอบระบบผู้ดูแลสัตว์")
    add_table(doc,
        ["รหัส", "รายการทดสอบ", "ผลลัพธ์ที่คาดหวัง"],
        [
            ["AT-11.1", "เข้าใช้งานผ่าน Smartphone (Android/iOS)", "หน้าจอ Responsive ใช้งานได้ครบทุกฟังก์ชัน"],
            ["AT-11.2", "แจ้งสัตว์ป่วยผ่าน Mobile", "สัตวแพทย์ได้รับการแจ้งเตือนภายใน 1 นาที"],
            ["AT-11.3", "แจ้งสัตว์ตาย", "สถานะในทะเบียนสัตว์อัปเดตเป็น 'ตาย' อัตโนมัติ"],
            ["AT-11.4", "บันทึกสภาพแวดล้อมและ Enrichment", "ข้อมูลบันทึกลงฐานข้อมูลพร้อม Timestamp"],
        ],
        col_widths=[1.8, 6.0, 6.0]
    )

    # Gap Analysis
    add_subsection_heading(doc, "3.6.6.5", "การวิเคราะห์ช่องว่าง As-Is vs To-Be (Gap Analysis)")
    tn = next_table_num()
    add_table_caption(doc, tn, "Gap Analysis ระบบผู้ดูแลสัตว์: As-Is vs To-Be")
    add_table(doc,
        ["ด้าน", "As-Is (ระบบเดิม)", "Gap (ช่องว่าง)", "To-Be (เป้าหมาย)"],
        [
            ["Mobile", "Desktop Only", "ผู้ดูแลไม่สามารถบันทึกข้อมูลในพื้นที่ได้", "Mobile-First Responsive (Smartphone/Tablet)"],
            ["Quick Actions", "ไม่มีฟังก์ชันแจ้งเหตุด่วน", "แจ้งป่วย/ตายต้องเดินมาคอมพิวเตอร์", "9 ฟังก์ชัน Quick Report ผ่าน Mobile"],
            ["Notification", "ไม่มีการแจ้งเตือน", "สัตวแพทย์ไม่ทราบเหตุการณ์ทันเวลา", "Push Notification ถึงสัตวแพทย์/ผู้บริหาร"],
            ["Enrichment", "ไม่มีระบบบันทึก Enrichment", "ขาดข้อมูลสวัสดิภาพสัตว์", "บันทึกกิจกรรมและการตอบสนอง"],
            ["สภาพแวดล้อม", "ไม่มีระบบบันทึกสภาพแวดล้อม", "ขาดข้อมูลอุณหภูมิ/ความชื้น", "บันทึกอุณหภูมิ ความชื้น ความสะอาด"],
        ],
        col_widths=[2.5, 3.5, 3.5, 4.0]
    )
    fig, ax = plt.subplots(1, 1, figsize=(10, 4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 4.5); ax.axis('off')
    ax.text(5, 4.2, 'Gap Analysis: ระบบผู้ดูแลสัตว์ — As-Is vs To-Be', ha='center', fontsize=12, fontweight='bold')
    ax.add_patch(FancyBboxPatch((0.2, 0.8), 4.0, 2.8, boxstyle="round,pad=0.1", facecolor='#e74c3c', alpha=0.15, edgecolor='#e74c3c'))
    ax.text(2.2, 3.3, 'As-Is', ha='center', fontsize=11, fontweight='bold', color='#e74c3c')
    for i, t in enumerate(['Desktop Only', 'ไม่มีแจ้งเหตุด่วน', 'ไม่มี Notification', 'ไม่มี Enrichment/Environment']):
        ax.text(2.2, 2.7 - i*0.45, t, ha='center', fontsize=9)
    ax.annotate('', xy=(5.6, 2.2), xytext=(4.4, 2.2), arrowprops=dict(arrowstyle='->', color='#2c3e50', lw=2.5))
    ax.text(5.0, 2.6, 'ปรับปรุง', ha='center', fontsize=9, fontweight='bold', color='#2c3e50')
    ax.add_patch(FancyBboxPatch((5.8, 0.8), 4.0, 2.8, boxstyle="round,pad=0.1", facecolor='#27ae60', alpha=0.15, edgecolor='#27ae60'))
    ax.text(7.8, 3.3, 'To-Be', ha='center', fontsize=11, fontweight='bold', color='#27ae60')
    for i, t in enumerate(['Mobile-First (9 ฟังก์ชัน)', 'Quick Report ทันที', 'Push Notification อัตโนมัติ', 'Enrichment + Environment']):
        ax.text(7.8, 2.7 - i*0.45, t, ha='center', fontsize=9)
    fig.tight_layout()
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "เปรียบเทียบ As-Is vs To-Be ระบบผู้ดูแลสัตว์")


# ══════════════════════════════════════════════
# MODULE 7: Studbook Management
# ══════════════════════════════════════════════

def build_module_studbook(doc):
    add_section_heading(doc, "3.6.7", "ระบบ Studbook Management (ทะเบียนสายพันธุ์) — TOR 4.14")

    add_body_after_section(doc,
        "ระบบ Studbook Management เป็นโมดูลใหม่ที่พัฒนาขึ้นทั้งหมดเพื่อรองรับการจัดทำทะเบียนสายพันธุ์สัตว์ "
        "ซึ่งเป็นหนึ่งในวัตถุประสงค์หลักของโครงการตาม TOR ข้อ 4.14 "
        "Studbook เป็นเครื่องมือสำคัญในการบริหารประชากรสัตว์ในกรงเลี้ยง (Captive Population Management) "
        "ที่ใช้ทั่วโลกโดยสอดคล้องกับมาตรฐาน Species360/ZIMS "
        "ระบบจะดึงข้อมูลจากทะเบียนสัตว์ทุกสถานะ และจัดเก็บข้อมูลมาตรฐาน 10 ฟิลด์ตาม TOR 4.14.3 "
        "ได้แก่ Studbook No, Animal ID, เพศ, อายุ, วันเกิด, วันตาย, หมายเลขพ่อ, หมายเลขแม่, "
        "แหล่งที่มา และสถานที่ปัจจุบัน โดยเรียงลำดับตามวันเกิด (TOR 4.14.1) "
        "และครอบคลุมสัตว์ทุกสถานะ: เกิด ตาย รับเข้า ส่งออก สูญหาย บริจาค จัดซื้อ (TOR 4.14.2) "
        "ตารางที่ 3.6-14 แสดง User Stories ตารางที่ 3.6-15 แสดงโครงสร้างข้อมูล "
        "และแผนภาพที่ 3.6-7 แสดงกระบวนการสร้าง Studbook"
    )

    add_subsection_heading(doc, "3.6.7.1", "User Stories")
    tn = next_table_num()
    add_table_caption(doc, tn, "User Stories ระบบ Studbook Management")
    add_table(doc,
        ["รหัส", "บทบาท", "ความต้องการ", "เป้าหมาย", "TOR"],
        [
            ["US-14.1", "เจ้าหน้าที่ Studbook", "ต้องการ Auto-Numbering ตาม DOB", "ลำดับมาตรฐาน", "4.14.1"],
            ["US-14.2", "เจ้าหน้าที่ Studbook", "ต้องการบันทึกสัตว์ทุกสถานะ", "Studbook ครบถ้วน", "4.14.2"],
            ["US-14.3", "นักวิจัย", "ต้องการข้อมูลพ่อ-แม่สำหรับ Genetic Analysis", "วิเคราะห์พันธุกรรมได้", "4.14.3"],
        ],
        col_widths=[1.8, 3.0, 5.0, 3.5, 1.2]
    )

    add_subsection_heading(doc, "3.6.7.2", "โครงสร้างข้อมูล 10 ฟิลด์มาตรฐาน (TOR 4.14.3)")
    tn = next_table_num()
    add_table_caption(doc, tn, "โครงสร้างข้อมูล Studbook 10 ฟิลด์มาตรฐาน")
    add_table(doc,
        ["ลำดับ", "ฟิลด์", "คำอธิบาย", "ประเภท", "Validation"],
        [
            ["1", "Studbook No", "ลำดับประจำ Studbook (Auto-generate)", "Integer", "Unique, Auto-increment by DOB"],
            ["2", "Animal ID", "รหัสประจำตัวสัตว์ (FK → Animal Registry)", "String", "FK, Not Null"],
            ["3", "เพศ (Sex)", "เพศของสัตว์", "Enum", "M/F/Unknown"],
            ["4", "อายุ (Age)", "อายุ ณ ปัจจุบัน (คำนวณจาก DOB)", "Computed", "Auto-calculated"],
            ["5", "วันเกิด (DOB)", "วันเดือนปีเกิด", "Date", "ไม่เป็นอนาคต"],
            ["6", "วันตาย (DOD)", "วันเดือนปีตาย (ถ้ามี)", "Date", "Nullable, ≥ DOB"],
            ["7", "หมายเลขพ่อ (Sire ID)", "Studbook No ของพ่อ", "FK", "FK → Studbook, Nullable"],
            ["8", "หมายเลขแม่ (Dam ID)", "Studbook No ของแม่", "FK", "FK → Studbook, Nullable"],
            ["9", "แหล่งที่มา (Origin)", "แหล่งที่มาของสัตว์", "String", "เกิด/รับเข้า/บริจาค/จัดซื้อ"],
            ["10", "สถานที่ปัจจุบัน", "สวนสัตว์ที่สัตว์อยู่ปัจจุบัน", "FK", "FK → Zoo Master"],
        ],
        col_widths=[1.0, 2.5, 4.0, 2.0, 4.0]
    )

    add_subsection_heading(doc, "3.6.7.3", "การออกแบบหน้าจอและกระบวนการทำงาน")
    fig = create_flow_diagram(
        "กระบวนการสร้างและจัดการ Studbook (TOR 4.14)",
        [
            (0.0, 2.5, 2.0, 1.0, '#c0392b', 'Animal\nRegistry'),
            (0.0, 1.0, 2.0, 1.0, '#8e44ad', 'Taxonomy\nDB'),
            (2.8, 2.5, 2.2, 1.0, '#2980b9', 'Studbook\nBuilder'),
            (5.5, 2.5, 2.0, 1.0, '#27ae60', 'Auto-Generate\nStudbook No.'),
            (8.0, 3.5, 2.0, 0.8, '#e67e22', 'Genetic\nAnalysis'),
            (8.0, 2.5, 2.0, 0.8, '#16a085', 'Pedigree\nDiagram'),
            (8.0, 1.5, 2.0, 0.8, '#d4ac0d', 'Population\nDashboard'),
        ],
        [(0,2), (1,2), (2,3), (3,4), (3,5), (3,6)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "กระบวนการสร้างและจัดการ Studbook")


# ══════════════════════════════════════════════
# MODULE 8: Genetic Analysis
# ══════════════════════════════════════════════

def build_module_genetic(doc):
    add_section_heading(doc, "3.6.8", "โมดูลวิเคราะห์พันธุกรรม (Genetic Analysis Module) — TOR 4.15")

    add_body_after_section(doc,
        "โมดูลวิเคราะห์พันธุกรรมเป็นหัวใจสำคัญของความสามารถใหม่ด้าน Conservation Data Platform "
        "ที่คำนวณความสัมพันธ์ทางพันธุกรรมอัตโนมัติตาม TOR ข้อ 4.15 ครอบคลุม 5 ความสามารถหลัก "
        "ได้แก่ คำนวณค่า Inbreeding Coefficient (F) อัตโนมัติเมื่อมีการอัปเดตข้อมูลพ่อ-แม่ (4.15.1) "
        "วิเคราะห์บรรพบุรุษร่วมหลายเส้นทาง (Multi-Path Common Ancestor Analysis) (4.15.2) "
        "ประเมินความลึกและความสมบูรณ์ของข้อมูล Pedigree (Pedigree Completeness) (4.15.3) "
        "ทำนายค่า F ของลูกจากคู่ผสมที่เลือก (Expected F Prediction) (4.15.4) "
        "และแสดงผลการเปรียบเทียบคู่ผสมที่เข้าใจง่าย (4.15.5) "
        "ตารางที่ 3.6-16 แสดง User Stories ตารางที่ 3.6-17 แสดง 5 ความสามารถหลัก "
        "และแผนภาพที่ 3.6-8 แสดงสถาปัตยกรรมการคำนวณ Genetic Analysis"
    )

    add_subsection_heading(doc, "3.6.8.1", "User Stories")
    tn = next_table_num()
    add_table_caption(doc, tn, "User Stories โมดูลวิเคราะห์พันธุกรรม")
    add_table(doc,
        ["รหัส", "บทบาท", "ความต้องการ", "เป้าหมาย", "TOR"],
        [
            ["US-15.1", "นักวิจัย", "ต้องการเห็นค่า F คำนวณอัตโนมัติเมื่ออัปเดตพ่อ-แม่", "ข้อมูลวิเคราะห์ทันสมัยเสมอ", "4.15.1"],
            ["US-15.2", "นักวิจัย", "ต้องการทำนายค่า F ของลูกจากคู่ผสมที่เลือก", "ตัดสินใจจัดคู่ผสมด้วยข้อมูล", "4.15.4"],
            ["US-15.3", "นักวิจัย", "ต้องการเปรียบเทียบคู่ผสมหลายคู่พร้อมกัน", "เลือกคู่ผสมที่ดีที่สุด", "4.15.5"],
        ],
        col_widths=[1.8, 2.0, 5.5, 3.5, 1.2]
    )

    add_subsection_heading(doc, "3.6.8.2", "5 ความสามารถหลักตาม TOR 4.15")
    tn = next_table_num()
    add_table_caption(doc, tn, "ความสามารถหลัก 5 ประการของโมดูล Genetic Analysis")
    add_table(doc,
        ["ลำดับ", "ความสามารถ", "รายละเอียด", "TOR", "อัลกอริทึม"],
        [
            ["1", "Auto-Calculate F", "คำนวณค่า F อัตโนมัติเมื่ออัปเดต Sire/Dam", "4.15.1", "Wright's Path Coefficient"],
            ["2", "Multi-Path Ancestor", "วิเคราะห์เส้นทางบรรพบุรุษร่วมทุกเส้นทาง", "4.15.2", "Graph Traversal (DFS/BFS)"],
            ["3", "Pedigree Completeness", "ประเมินความลึกและความสมบูรณ์ของข้อมูล", "4.15.3", "Completeness Index (CI)"],
            ["4", "Expected F Prediction", "ทำนายค่า F ของลูกจากคู่ที่เลือก", "4.15.4", "Simulated Offspring F"],
            ["5", "Pair Comparison", "เปรียบเทียบคู่ผสมหลายคู่แบบเห็นภาพชัด", "4.15.5", "Side-by-Side + Color Coding"],
        ],
        col_widths=[1.0, 3.0, 4.5, 1.0, 4.0]
    )

    add_subsection_heading(doc, "3.6.8.3", "สถาปัตยกรรมการคำนวณ")
    fig = create_flow_diagram(
        "สถาปัตยกรรม Genetic Analysis Engine (TOR 4.15)",
        [
            (0.0, 2.5, 2.0, 1.0, '#c0392b', 'Studbook\nDB'),
            (2.5, 2.5, 2.0, 1.0, '#2980b9', 'Pedigree\nResolver'),
            (5.0, 3.5, 2.2, 1.0, '#27ae60', 'Ancestor Path\nAnalysis'),
            (5.0, 1.5, 2.2, 1.0, '#8e44ad', 'Compute F\n& Relatedness'),
            (7.8, 3.5, 2.2, 1.0, '#e67e22', 'Pairing\nSimulator'),
            (7.8, 1.5, 2.2, 1.0, '#16a085', 'Genetic\nMetrics DB'),
        ],
        [(0,1), (1,2), (1,3), (2,3), (3,4), (3,5), (4,5)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "สถาปัตยกรรม Genetic Analysis Engine")


# ══════════════════════════════════════════════
# MODULE 9: Pedigree Diagram
# ══════════════════════════════════════════════

def build_module_pedigree(doc):
    add_section_heading(doc, "3.6.9", "ระบบ Pedigree Diagram (ผังสายโลหิต) — TOR 4.16")

    add_body_after_section(doc,
        "ระบบ Pedigree Diagram แสดงผังสายโลหิตสัตว์แบบ Interactive ตาม TOR ข้อ 4.16 "
        "ครอบคลุม 6 ข้อกำหนด ได้แก่ แสดงความสัมพันธ์พ่อ-แม่ (4.16.1) เชื่อมกับ Taxonomy (4.16.2) "
        "แสดง ID, ชื่อ, เพศ, อายุ, สถานะ, ค่า F (4.16.3) แสดงค่า Expected F บนเส้นเชื่อม (4.16.4) "
        "สี/สัญลักษณ์เตือนเมื่อ F เกินเกณฑ์ (4.16.5) และส่งออก PDF/Excel/Image (4.16.6) "
        "ระบบใช้ D3.js หรือ GoJS ในการแสดงผล Interactive Tree ที่ซูม, แพน, คลิกเพื่อดูรายละเอียดได้ "
        "ตารางที่ 3.6-18 แสดงข้อกำหนด 6 รายการ และแผนภาพที่ 3.6-9 แสดงสถาปัตยกรรม Pedigree"
    )

    add_subsection_heading(doc, "3.6.9.1", "ข้อกำหนด 6 รายการตาม TOR 4.16")
    tn = next_table_num()
    add_table_caption(doc, tn, "ข้อกำหนด Pedigree Diagram 6 รายการ")
    add_table(doc,
        ["ลำดับ", "ข้อกำหนด", "รายละเอียดการออกแบบ", "TOR"],
        [
            ["1", "แสดงความสัมพันธ์พ่อ-แม่", "Tree Layout แบบ Top-Down, Sire ซ้าย Dam ขวา", "4.16.1"],
            ["2", "เชื่อมกับ Taxonomy", "แสดงชื่อวิทยาศาสตร์และชื่อสามัญจาก Taxonomy DB", "4.16.2"],
            ["3", "แสดงข้อมูล 6 รายการต่อ Node", "Animal ID, ชื่อ, เพศ, อายุ, สถานะ, ค่า F", "4.16.3"],
            ["4", "แสดง Expected F บนเส้นเชื่อม", "ตัวเลข Expected F แสดงบนเส้นระหว่างพ่อ-แม่", "4.16.4"],
            ["5", "สี/สัญลักษณ์เตือน F เกินเกณฑ์", "Node สีแดงเมื่อ F > 0.0625 (เกณฑ์ตั้งค่าได้)", "4.16.5"],
            ["6", "ส่งออก PDF/Excel/Image", "ปุ่ม Export ทั้ง 3 รูปแบบ พร้อม High-DPI", "4.16.6"],
        ],
        col_widths=[1.0, 3.5, 6.5, 1.2]
    )

    add_subsection_heading(doc, "3.6.9.2", "สถาปัตยกรรมและเทคโนโลยี")
    fig = create_flow_diagram(
        "สถาปัตยกรรม Pedigree Diagram (TOR 4.16)",
        [
            (0.0, 2.5, 2.0, 1.0, '#c0392b', 'Studbook\nDB'),
            (0.0, 1.0, 2.0, 1.0, '#2980b9', 'Genetic\nMetrics DB'),
            (0.0, 3.8, 2.0, 0.8, '#8e44ad', 'Taxonomy\nDB'),
            (3.0, 2.5, 2.5, 1.0, '#27ae60', 'Pedigree Graph\nEngine (D3.js)'),
            (6.5, 3.5, 2.0, 0.8, '#e67e22', 'Threshold\nColoring'),
            (6.5, 2.5, 2.0, 0.8, '#16a085', 'Expected F\non Lines'),
            (6.5, 1.5, 2.0, 0.8, '#d4ac0d', 'PDF/Excel/\nImage Export'),
        ],
        [(0,3), (1,3), (2,3), (3,4), (3,5), (3,6)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "สถาปัตยกรรม Pedigree Diagram")


# ══════════════════════════════════════════════
# MODULE 10: Reports & Dashboard
# ══════════════════════════════════════════════

def build_module_reports_dashboard(doc):
    add_section_heading(doc, "3.6.10", "ระบบรายงานและ Dashboard (Reports & Dashboard) — TOR 4.17")

    add_body_after_section(doc,
        "ระบบรายงานและ Dashboard เป็นโมดูลใหม่ตาม TOR ข้อ 4.17 ที่จัดทำรายงานด้านประชากรสัตว์และพันธุกรรม "
        "ครอบคลุม 4 รายงานหลักและ 1 ฟังก์ชันส่งออก ได้แก่ Inbreeding Coefficient Report (4.17.1) "
        "ที่แสดงค่า F รายตัว รายคู่ และค่าเฉลี่ยประชากร Age Pyramid Report (4.17.2) "
        "ที่แสดงโครงสร้างอายุแยกตามเพศ พร้อมตั้งค่าช่วงอายุได้ตามชนิดสัตว์ "
        "Birth Seasonality Report (4.17.3) ที่แสดงแนวโน้มการเกิดรายเดือน/รายฤดูกาล "
        "Population Structure Report (4.17.4) ที่แสดงโครงสร้างประชากรแยกตามเพศ กลุ่มอายุ และสถานะ "
        "และทุกรายงานส่งออกเป็น PDF ได้เป็นอย่างน้อย (4.17.5) "
        "ตารางที่ 3.6-19 แสดงรายงาน 4 รายการ และแผนภาพที่ 3.6-10 แสดงสถาปัตยกรรม Reporting"
    )

    add_subsection_heading(doc, "3.6.10.1", "รายงาน 4 รายการ + ฟังก์ชันส่งออก")
    tn = next_table_num()
    add_table_caption(doc, tn, "รายงาน 4 รายการและฟังก์ชันส่งออก")
    add_table(doc,
        ["ลำดับ", "รายงาน", "รายละเอียด", "TOR", "Visualization"],
        [
            ["1", "Inbreeding Coefficient", "ค่า F รายตัว รายคู่ ค่าเฉลี่ยประชากร", "4.17.1", "Table + Heat Map"],
            ["2", "Age Pyramid", "โครงสร้างอายุแยกเพศ ตั้งค่าช่วงได้", "4.17.2", "Horizontal Bar Chart"],
            ["3", "Birth Seasonality", "แนวโน้มการเกิดรายเดือน/ฤดูกาล", "4.17.3", "Line/Bar Chart"],
            ["4", "Population Structure", "โครงสร้างตามเพศ กลุ่มอายุ สถานะ", "4.17.4", "Stacked Bar / Pie"],
            ["5", "PDF Export", "ทุกรายงานส่งออก PDF (อย่างน้อย)", "4.17.5", "Puppeteer / jsPDF"],
        ],
        col_widths=[1.0, 3.0, 4.5, 1.0, 4.0]
    )

    add_subsection_heading(doc, "3.6.10.2", "สถาปัตยกรรม Reporting")
    fig = create_flow_diagram(
        "สถาปัตยกรรม Reports & Dashboard (TOR 4.17)",
        [
            (0.0, 3.0, 1.8, 0.8, '#c0392b', 'Animal\nRegistry'),
            (0.0, 2.0, 1.8, 0.8, '#2980b9', 'Studbook'),
            (0.0, 1.0, 1.8, 0.8, '#27ae60', 'Genetic\nMetrics'),
            (2.5, 2.0, 2.5, 1.0, '#8e44ad', 'Reporting\nData Mart'),
            (5.5, 3.5, 2.0, 0.8, '#e67e22', 'Inbreeding\nReport'),
            (5.5, 2.5, 2.0, 0.8, '#16a085', 'Age\nPyramid'),
            (5.5, 1.5, 2.0, 0.8, '#d4ac0d', 'Birth\nSeasonality'),
            (5.5, 0.5, 2.0, 0.8, '#e74c3c', 'Population\nStructure'),
            (8.5, 2.0, 2.0, 1.0, '#2c3e50', 'PDF\nExport'),
        ],
        [(0,3), (1,3), (2,3), (3,4), (3,5), (3,6), (3,7), (4,8), (5,8), (6,8), (7,8)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "สถาปัตยกรรม Reports & Dashboard")


# ══════════════════════════════════════════════
# MODULE 11: Veterinary Billing
# ══════════════════════════════════════════════

def build_module_billing(doc):
    add_section_heading(doc, "3.6.11", "ระบบคิดค่าบริการดูแลรักษาสัตว์ (Veterinary Billing) — TOR 4.18")

    add_body_after_section(doc,
        "ระบบคิดค่าบริการเป็นโมดูลใหม่ที่ซับซ้อนที่สุดตาม TOR ข้อ 4.18 ครอบคลุม 7 ระบบย่อย "
        "รวม 24 ข้อกำหนดย่อย ทำงานภายใต้เวชระเบียนสัตว์โดยเชื่อมกับข้อมูลหัตถการ ยาและเวชภัณฑ์ "
        "อัตราค่าบริการ สิทธิ์สมาชิก และระบบชำระเงิน ระบบย่อย 7 ส่วน ได้แก่ "
        "การบันทึกหัตถการ (4.18.1) การตรวจสอบสิทธิ์ (4.18.2) กฎราคายา (4.18.3) "
        "จัดการอัตราค่าบริการ (4.18.4) คำนวณค่าใช้จ่าย (4.18.5) บันทึกการชำระเงิน (4.18.6) "
        "และ Audit Trail + รายงาน (4.18.7) "
        "ตารางที่ 3.6-20 แสดง 7 ระบบย่อยและข้อกำหนด "
        "ตารางที่ 3.6-21 แสดง Medicine Pricing Logic "
        "และแผนภาพที่ 3.6-11 แสดง Billing Workflow ตั้งแต่บันทึกหัตถการจนถึงชำระเงิน"
    )

    add_subsection_heading(doc, "3.6.11.1", "7 ระบบย่อยและข้อกำหนด 24 รายการ")
    tn = next_table_num()
    add_table_caption(doc, tn, "ระบบย่อย 7 ส่วนของ Veterinary Billing")
    add_table(doc,
        ["ระบบย่อย", "รายละเอียด", "ข้อกำหนดย่อย", "TOR"],
        [
            ["4.18.1 บันทึกหัตถการ", "เลือกรหัสหัตถการมาตรฐาน บันทึกจำนวน หน่วย วันที่ เชื่อม Rate Card", "3 ข้อ", "4.18.1.1-1.3"],
            ["4.18.2 ตรวจสอบสิทธิ์", "ตรวจสอบสมาชิกภายใน/ภายนอก ภายใน→ไม่คิดค่า ภายนอก→คิดตามเงื่อนไข", "3 ข้อ", "4.18.2.1-2.3"],
            ["4.18.3 กฎราคายา", "ยาบริจาค→ไม่คิดค่ายา ยาจัดซื้อ→คิดค่ายา", "2 ข้อ", "4.18.3.1-3.2"],
            ["4.18.4 อัตราค่าบริการ", "กำหนดรหัส อัตรา หน่วย อัปเดต/ยกเลิก พร้อมประวัติ วันมีผล/สิ้นสุด เชื่อม Billing", "4 ข้อ", "4.18.4.1-4.4"],
            ["4.18.5 คำนวณค่าใช้จ่าย", "แยกหัตถการ/ยา ต่อครั้ง/ต่อระยะรักษา ยกเว้น/ส่วนลด+เหตุผล ออกใบแจ้ง", "4 ข้อ", "4.18.5.1-5.4"],
            ["4.18.6 บันทึกชำระเงิน", "แสดงใบแจ้ง + สถานะชำระ", "1 ข้อ", "4.18.6.1"],
            ["4.18.7 Audit + รายงาน", "Audit Trail ทุกขั้นตอน รายงานค่ารักษา/หัตถการ/ชำระ ส่งออก PDF/Excel", "3 ข้อ", "4.18.7.1-7.3"],
        ],
        col_widths=[3.0, 5.5, 1.8, 2.5]
    )

    add_subsection_heading(doc, "3.6.11.2", "ตรรกะการคิดราคายา (Medicine Pricing Logic)")
    tn = next_table_num()
    add_table_caption(doc, tn, "ตรรกะการคิดราคายา")
    add_table(doc,
        ["แหล่งที่มายา", "สิทธิ์สัตว์", "คิดค่าหัตถการ", "คิดค่ายา", "TOR"],
        [
            ["จัดซื้อ", "ภายนอก", "คิด (ตาม Rate Card)", "คิด (ตามราคาจัดซื้อ)", "4.18.2+4.18.3"],
            ["จัดซื้อ", "ภายใน", "ไม่คิด", "ไม่คิด", "4.18.2.2"],
            ["บริจาค", "ภายนอก", "คิด (ตาม Rate Card)", "ไม่คิด", "4.18.3.1"],
            ["บริจาค", "ภายใน", "ไม่คิด", "ไม่คิด", "4.18.2.2+4.18.3.1"],
        ],
        col_widths=[2.5, 2.5, 3.0, 3.0, 2.5]
    )

    add_subsection_heading(doc, "3.6.11.3", "กระบวนการทำงาน Billing Workflow")
    fig = create_flow_diagram(
        "กระบวนการ Billing Workflow (TOR 4.18)",
        [
            (0.0, 2.5, 2.0, 1.0, '#2c3e50', 'เวชระเบียน\n(Medical Record)'),
            (2.5, 3.5, 2.0, 1.0, '#2980b9', 'บันทึก\nหัตถการ'),
            (2.5, 1.5, 2.0, 1.0, '#27ae60', 'ตรวจสอบ\nสิทธิ์'),
            (5.2, 3.5, 2.0, 1.0, '#8e44ad', 'Rate Card\nEngine'),
            (5.2, 1.5, 2.0, 1.0, '#e67e22', 'Drug\nPricing'),
            (7.8, 2.5, 2.0, 1.0, '#e74c3c', 'คำนวณ\nค่าใช้จ่าย'),
            (10.0, 3.0, 1.0, 0.8, '#16a085', 'Invoice'),
            (10.0, 1.8, 1.0, 0.8, '#d4ac0d', 'Audit'),
        ],
        [(0,1), (0,2), (1,3), (2,3), (2,4), (3,5), (4,5), (5,6), (5,7)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "กระบวนการ Billing Workflow")

    add_subsection_heading(doc, "3.6.11.4", "เกณฑ์การทดสอบและยอมรับ (Acceptance Tests)")
    tn = next_table_num()
    add_table_caption(doc, tn, "เกณฑ์การทดสอบระบบ Veterinary Billing")
    add_table(doc,
        ["รหัส", "รายการทดสอบ", "ผลลัพธ์ที่คาดหวัง"],
        [
            ["AT-18.1", "บันทึกหัตถการ → คำนวณค่าใช้จ่าย", "ค่าหัตถการคำนวณถูกต้องตาม Rate Card"],
            ["AT-18.2", "สัตว์ภายใน → ไม่คิดค่าบริการ", "ยอดรวม = 0 บาท"],
            ["AT-18.3", "ยาบริจาค → ไม่คิดค่ายา", "คิดเฉพาะค่าหัตถการ ไม่คิดค่ายา"],
            ["AT-18.4", "ออก Invoice → ส่งออก PDF", "Invoice มีรายละเอียดครบ ส่งออก PDF ได้"],
            ["AT-18.5", "ตรวจสอบ Audit Trail", "บันทึกทุกขั้นตอนพร้อมผู้ดำเนินการและ Timestamp"],
        ],
        col_widths=[1.8, 6.0, 6.0]
    )


# ══════════════════════════════════════════════
# MODULE 12: Change Request
# ══════════════════════════════════════════════

def build_module_change_request(doc):
    add_section_heading(doc, "3.6.12", "ระบบคำร้องขอเปลี่ยนแปลงแก้ไข (Change Request) — TOR 4.19")

    add_body_after_section(doc,
        "ระบบคำร้องขอเปลี่ยนแปลงเป็นโมดูลใหม่ตาม TOR ข้อ 4.19 ทำหน้าที่เป็น Governance Workflow "
        "สำหรับการขอแก้ไขข้อมูลหรือปรับปรุงรายการสำคัญในระบบ ครอบคลุม 3 ข้อกำหนด "
        "ได้แก่ จัดการหมวดหมู่คำร้อง (4.19.1) บันทึกคำร้องพร้อมแนบ PDF (4.19.2) "
        "และติดตามสถานะคำร้อง: รอดำเนินการ → กำลังดำเนินการ → เสร็จสิ้น → ปฏิเสธ (4.19.3) "
        "ระบบนี้เชื่อมโยงโดยตรงกับระบบทะเบียนสัตว์ (หัวข้อ 3.6.2) สำหรับกระบวนการ Controlled Deletion "
        "ตารางที่ 3.6-23 แสดง User Stories และแผนภาพที่ 3.6-12 แสดงกระบวนการ Change Request"
    )

    add_subsection_heading(doc, "3.6.12.1", "User Stories")
    tn = next_table_num()
    add_table_caption(doc, tn, "User Stories ระบบ Change Request")
    add_table(doc,
        ["รหัส", "บทบาท", "ความต้องการ", "เป้าหมาย", "TOR"],
        [
            ["US-19.1", "ผู้ร้องขอ", "ต้องการส่งคำร้องพร้อมแนบเอกสาร PDF", "มีหลักฐานประกอบคำร้อง", "4.19.2"],
            ["US-19.2", "ผู้อนุมัติ", "ต้องการอนุมัติ/ปฏิเสธคำร้องพร้อมเหตุผล", "Governance ครบถ้วน", "4.19.3"],
            ["US-19.3", "ผู้ร้องขอ", "ต้องการติดตามสถานะคำร้องแบบ Real-Time", "ทราบความคืบหน้า", "4.19.3"],
        ],
        col_widths=[1.8, 2.5, 5.5, 3.0, 1.2]
    )

    add_subsection_heading(doc, "3.6.12.2", "กระบวนการทำงาน")
    fig = create_flow_diagram(
        "กระบวนการ Change Request Workflow (TOR 4.19)",
        [
            (0.0, 2.5, 2.0, 1.0, '#2c3e50', 'ผู้ร้องขอ'),
            (2.5, 2.5, 2.0, 1.0, '#2980b9', 'สร้างคำร้อง\n+ หมวดหมู่'),
            (5.0, 2.5, 2.0, 1.0, '#8e44ad', 'แนบ PDF\n+ ส่ง'),
            (7.5, 3.2, 2.0, 0.8, '#27ae60', 'อนุมัติ'),
            (7.5, 1.8, 2.0, 0.8, '#e74c3c', 'ปฏิเสธ'),
            (10.0, 2.5, 1.0, 1.0, '#16a085', 'CR\nDB'),
        ],
        [(0,1), (1,2), (2,3), (2,4), (3,5), (4,5)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "กระบวนการ Change Request Workflow")

    add_subsection_heading(doc, "3.6.12.3", "กฎการตรวจสอบข้อมูลและเกณฑ์ทดสอบ")
    add_numbered_item_subsection(doc, 1, "หมวดหมู่คำร้องต้องเลือกจาก Master Data (ไม่อนุญาต Free Text)")
    add_numbered_item_subsection(doc, 2, "เอกสารแนบต้องเป็นไฟล์ PDF ขนาดไม่เกิน 10 MB")
    add_numbered_item_subsection(doc, 3, "การเปลี่ยนสถานะต้องบันทึกผู้ดำเนินการ วันเวลา และเหตุผล")
    add_numbered_item_subsection(doc, 4, "คำร้องที่อนุมัติแล้วสามารถอ้างอิงใน Controlled Deletion ของทะเบียนสัตว์ได้")


# ══════════════════════════════════════════════
# MODULE 13: Manual CMS + Video
# ══════════════════════════════════════════════

def build_module_manual_cms(doc):
    add_section_heading(doc, "3.6.13", "ระบบจัดเก็บเนื้อหาคู่มือการใช้งาน (Manual CMS + Video) — TOR 4.20")

    add_body_after_section(doc,
        "ระบบจัดเก็บเนื้อหาคู่มือเป็นโมดูลใหม่ตาม TOR ข้อ 4.20 ที่เป็นศูนย์รวมคู่มือการใช้งาน "
        "รองรับ 3 ข้อกำหนด ได้แก่ อัปโหลดไฟล์ PDF (4.20.1) แนบลิงก์และ Embed วิดีโอ (4.20.2) "
        "และจัดทำวิดีโอคู่มือ 5 โมดูลที่ TOR กำหนด (4.20.3) ได้แก่ ระบบคลังอาหาร ระบบผู้ดูแลสัตว์ "
        "ระบบเวชระเบียน ระบบยาและเวชภัณฑ์ และระบบโภชนาการ "
        "ตารางที่ 3.6-24 แสดงวิดีโอคู่มือ 5 รายการ และแผนภาพที่ 3.6-13 แสดงสถาปัตยกรรม Manual CMS"
    )

    add_subsection_heading(doc, "3.6.13.1", "วิดีโอคู่มือ 5 รายการตาม TOR 4.20.3")
    tn = next_table_num()
    add_table_caption(doc, tn, "วิดีโอคู่มือ 5 โมดูลตามข้อกำหนด TOR 4.20.3")
    add_table(doc,
        ["ลำดับ", "โมดูล", "เนื้อหาวิดีโอ", "ความยาวประมาณ", "รูปแบบ"],
        [
            ["1", "คลังอาหารสัตว์", "Touch Screen: รับเข้า เบิก คืน สแกน Barcode พิมพ์ฉลาก", "10-15 นาที", "MP4 + Embed"],
            ["2", "ผู้ดูแลสัตว์", "Mobile: 9 ฟังก์ชัน แจ้งเกิด/ป่วย/ตาย บันทึกสิ่งแวดล้อม", "10-15 นาที", "MP4 + Embed"],
            ["3", "เวชระเบียน", "บันทึกการตรวจ สั่งยา ดูประวัติ คิดค่ารักษา", "10-15 นาที", "MP4 + Embed"],
            ["4", "ยา/เวชภัณฑ์", "รับเข้า เบิก คืน Barcode บริจาค/จัดซื้อ", "8-12 นาที", "MP4 + Embed"],
            ["5", "โภชนาการ", "Diet Card Inline Editing ตรวจสอบปริมาณ", "8-12 นาที", "MP4 + Embed"],
        ],
        col_widths=[1.0, 2.5, 5.5, 2.0, 2.5]
    )

    add_subsection_heading(doc, "3.6.13.2", "สถาปัตยกรรม Manual CMS")
    fig = create_flow_diagram(
        "สถาปัตยกรรม Manual CMS (TOR 4.20)",
        [
            (0.0, 2.5, 2.0, 1.0, '#2c3e50', 'Admin\n/ Trainer'),
            (2.5, 3.2, 2.0, 0.8, '#2980b9', 'Upload PDF'),
            (2.5, 2.0, 2.0, 0.8, '#27ae60', 'Add Video\nLink'),
            (5.5, 2.5, 2.5, 1.0, '#8e44ad', 'Knowledge\nBase DB'),
            (8.5, 3.2, 2.0, 0.8, '#e67e22', 'ค้นหา\nตามโมดูล'),
            (8.5, 2.0, 2.0, 0.8, '#16a085', 'ดู PDF\n/ วิดีโอ'),
        ],
        [(0,1), (0,2), (1,3), (2,3), (3,4), (3,5)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "สถาปัตยกรรม Manual CMS")


# ══════════════════════════════════════════════
# MODULE 14: System Log
# ══════════════════════════════════════════════

def build_module_system_log(doc):
    add_section_heading(doc, "3.6.14", "ระบบบันทึกเหตุการณ์ System Log — TOR 4.21")

    add_body_after_section(doc,
        "ระบบบันทึกเหตุการณ์เป็นโมดูลใหม่ตาม TOR ข้อ 4.21 ที่ทำหน้าที่เป็น Central Event Logger "
        "สำหรับตรวจสอบย้อนหลังและควบคุมความมั่นคงปลอดภัย ครอบคลุม 5 ข้อกำหนด ได้แก่ "
        "บันทึก Login/Logout (4.21.1) บันทึก User Actions: Create, Update เป็นอย่างน้อย (4.21.2) "
        "บันทึก Timestamp (4.21.3) บันทึก Username/User ID (4.21.4) "
        "และกำหนด Retention Period ได้ไม่น้อยกว่า 90 วัน (4.21.5) "
        "นอกจากนี้ ระบบยังเสริมการบันทึก IP Address และ User Agent เพื่อเพิ่มศักยภาพด้านความปลอดภัย "
        "ข้อมูล Log จัดเก็บในตารางแยก (Partitioned Table) เพื่อประสิทธิภาพในการค้นหาและจัดการ Retention "
        "ตารางที่ 3.6-25 แสดงข้อกำหนด 5 รายการ ตารางที่ 3.6-26 แสดงโครงสร้างข้อมูล System Log "
        "และแผนภาพที่ 3.6-14 แสดงสถาปัตยกรรม System Log"
    )

    add_subsection_heading(doc, "3.6.14.1", "ข้อกำหนด 5 รายการตาม TOR 4.21")
    tn = next_table_num()
    add_table_caption(doc, tn, "ข้อกำหนด System Log 5 รายการ")
    add_table(doc,
        ["ลำดับ", "ข้อกำหนด", "รายละเอียดการออกแบบ", "TOR"],
        [
            ["1", "Login/Logout", "บันทึกทุกเหตุการณ์เข้า-ออกระบบ พร้อม IP และ User Agent", "4.21.1"],
            ["2", "User Actions", "Create, Update, Delete (Soft) เป็นอย่างน้อย + Module Name", "4.21.2"],
            ["3", "Timestamp", "ISO 8601 format พร้อม Timezone (Asia/Bangkok)", "4.21.3"],
            ["4", "Username/User ID", "บันทึกทั้ง Username และ User ID ทุกเหตุการณ์", "4.21.4"],
            ["5", "Configurable Retention", "ตั้งค่าได้ ค่าเริ่มต้น 90 วัน รองรับปรับเป็น 180/365 วัน", "4.21.5"],
        ],
        col_widths=[1.0, 3.0, 6.5, 1.2]
    )

    add_subsection_heading(doc, "3.6.14.2", "โครงสร้างข้อมูล System Log")
    tn = next_table_num()
    add_table_caption(doc, tn, "โครงสร้างข้อมูลตาราง System Log")
    add_table(doc,
        ["ฟิลด์", "ประเภท", "คำอธิบาย", "ตัวอย่าง"],
        [
            ["id", "BIGINT (PK)", "รหัสเหตุการณ์ Auto-increment", "100001"],
            ["timestamp", "DATETIME", "วันเวลาเหตุการณ์ (ISO 8601)", "2569-03-15T10:30:00+07:00"],
            ["user_id", "INT (FK)", "รหัสผู้ใช้", "42"],
            ["username", "VARCHAR(100)", "ชื่อผู้ใช้", "somchai.k"],
            ["action", "ENUM", "ประเภทเหตุการณ์", "LOGIN/CREATE/UPDATE/DELETE"],
            ["module", "VARCHAR(50)", "โมดูลที่เกี่ยวข้อง", "animal_registry"],
            ["detail", "JSON", "รายละเอียดเหตุการณ์", '{"animal_id": "A001", "field": "name"}'],
            ["ip_address", "VARCHAR(45)", "IP Address ผู้ใช้", "192.168.1.100"],
            ["user_agent", "VARCHAR(255)", "Browser/Device ผู้ใช้", "Chrome/120 Windows"],
        ],
        col_widths=[2.0, 2.5, 3.5, 5.0]
    )

    add_subsection_heading(doc, "3.6.14.3", "สถาปัตยกรรม System Log")
    fig = create_flow_diagram(
        "สถาปัตยกรรม System Log (TOR 4.21)",
        [
            (0.0, 2.5, 2.0, 1.0, '#2c3e50', 'All\nModules'),
            (2.5, 2.5, 2.0, 1.0, '#2980b9', 'Event\nCollector'),
            (5.0, 2.5, 2.0, 1.0, '#27ae60', 'Normalize\nEvent'),
            (7.5, 3.5, 2.0, 0.8, '#8e44ad', 'Search\n& Filter'),
            (7.5, 2.5, 2.0, 0.8, '#e67e22', 'System Log\nStore'),
            (7.5, 1.5, 2.0, 0.8, '#e74c3c', 'Retention\nPolicy ≥90d'),
        ],
        [(0,1), (1,2), (2,4), (4,3), (4,5)]
    )
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "สถาปัตยกรรม System Log")

    add_subsection_heading(doc, "3.6.14.4", "เกณฑ์การทดสอบและยอมรับ (Acceptance Tests)")
    tn = next_table_num()
    add_table_caption(doc, tn, "เกณฑ์การทดสอบระบบ System Log")
    add_table(doc,
        ["รหัส", "รายการทดสอบ", "ผลลัพธ์ที่คาดหวัง"],
        [
            ["AT-21.1", "Login เข้าระบบแล้วตรวจสอบ Log", "พบบันทึก LOGIN พร้อม Timestamp, User ID, IP"],
            ["AT-21.2", "สร้างข้อมูลสัตว์แล้วตรวจสอบ Log", "พบบันทึก CREATE พร้อม Module = animal_registry"],
            ["AT-21.3", "ตั้งค่า Retention เป็น 90 วัน", "Log เก่ากว่า 90 วันถูกลบอัตโนมัติ"],
            ["AT-21.4", "ค้นหา Log ตาม User/Module/Date Range", "แสดงผลลัพธ์ถูกต้องภายใน 2 วินาที"],
        ],
        col_widths=[1.8, 6.0, 6.0]
    )


# ══════════════════════════════════════════════
# SECTION 3.6.15: Best Practices & Benchmarks
# ══════════════════════════════════════════════

def build_module_benchmarks(doc):
    add_section_heading(doc, "3.6.15", "แนวปฏิบัติที่ดีและมาตรฐานอ้างอิงระดับสากล (Best Practices & Benchmarks)")

    add_body_after_section(doc,
        "หัวข้อนี้นำเสนอแนวปฏิบัติที่ดีระดับสากลและ Benchmarks สำหรับโมดูลใหม่ 8 โมดูล "
        "ที่พัฒนาขึ้นในโครงการนี้ เพื่อให้มั่นใจว่าการออกแบบและพัฒนาระบบ Thai Zoo ARK "
        "สอดคล้องกับมาตรฐานที่ยอมรับในระดับนานาชาติ โดยอ้างอิงจาก 4 แหล่งหลัก "
        "ส่วนแรก (หัวข้อ 3.6.15.1) นำเสนอมาตรฐาน Species360/ZIMS และ EAZA Population Management "
        "สำหรับระบบ Studbook, Genetic Analysis และ Pedigree "
        "ส่วนที่สอง (หัวข้อ 3.6.15.2) นำเสนอแนวปฏิบัติ Veterinary Practice Management "
        "สำหรับระบบ Billing รวมถึงหลักการ Charge Capture Integrity และ Revenue Leakage Prevention "
        "ส่วนที่สาม (หัวข้อ 3.6.15.3) นำเสนอมาตรฐาน OWASP Security Logging "
        "สำหรับระบบ System Log, Change Request และ Manual CMS "
        "และส่วนที่สี่ (หัวข้อ 3.6.15.4) สรุปตาราง Benchmark Mapping "
        "ที่แสดงการเชื่อมโยงระหว่างมาตรฐานสากลกับการออกแบบในแต่ละโมดูล "
        "ตารางที่ {t1} ถึง {t4} แสดงรายละเอียดทั้ง 4 ส่วน "
        "และแผนภาพที่ {d1} และ {d2} แสดงกรอบแนวคิดสำคัญ".format(
            t1="3.6-" + str(table_counter + 1), t4="3.6-" + str(table_counter + 4),
            d1="3.6-" + str(diagram_counter + 1), d2="3.6-" + str(diagram_counter + 2))
    )

    # ── 3.6.15.1 Species360/ZIMS & EAZA ──
    add_subsection_heading(doc, "3.6.15.1", "มาตรฐาน Species360/ZIMS และ EAZA Population Management")
    add_body_after_subsection(doc,
        "Species360 เป็นองค์กรไม่แสวงกำไรที่รวมชุมชนสวนสัตว์และพิพิธภัณฑ์สัตว์น้ำกว่า 1,300 แห่งทั่วโลก "
        "โดยพัฒนาระบบ ZIMS (Zoological Information Management System) "
        "ซึ่งเป็นมาตรฐานสากลสำหรับการจัดการข้อมูลสัตว์ในกรงเลี้ยง "
        "ในขณะที่ EAZA (European Association of Zoos and Aquaria) จัดทำ Population Management Manual "
        "ที่กำหนดมาตรฐานและขั้นตอนสำหรับการบริหารจัดการประชากรสัตว์ "
        "แนวปฏิบัติเหล่านี้เป็นพื้นฐานสำคัญในการออกแบบ Studbook, Genetic Analysis และ Pedigree "
        "ของระบบ Thai Zoo ARK ดังแสดงในตารางที่ {tn}".format(tn="3.6-" + str(table_counter + 1))
    )
    tn = next_table_num()
    add_table_caption(doc, tn, "แนวปฏิบัติ Species360/ZIMS & EAZA สำหรับ Conservation Modules")
    add_table(doc,
        ["มาตรฐาน/แนวปฏิบัติ", "หลักการสำคัญ", "การนำมาใช้ใน Thai Zoo ARK"],
        [
            ["ZIMS Studbook Taxonomy", "Taxonomy ใน Studbook ต้องตรงกับ Husbandry Records เพื่อเชื่อมโยงข้อมูลได้", "Studbook เชื่อมกับ Taxonomy DB + Animal Registry ผ่าน species_id"],
            ["EAZA Population Management Manual", "ใช้ Wright's Path Coefficient สำหรับคำนวณ F-Coefficient และ Mean Kinship (MK)", "Genetic Analysis ใช้ Wright's Path Coefficient + Multi-Path Ancestor"],
            ["Species360 Best Practices in Record Keeping", "บันทึกสัตว์ทุกสถานะ (Born, Died, Transfer, Missing) พร้อมวันที่แม่นยำ", "Studbook ครอบคลุม 7 สถานะตาม TOR 4.14.2"],
            ["EAZA Pedigree Completeness Index", "ประเมินความสมบูรณ์ของ Pedigree ใช้ Completeness Index (CI) สำหรับประเมินความน่าเชื่อถือ", "Genetic Analysis คำนวณ CI ทุก Pedigree (TOR 4.15.3)"],
            ["AZA Institutional Records Keeping", "รหัสสัตว์ต้อง Unique และ Standardized สอดคล้องกับ ZIMS", "Animal ID เป็น Single Source of Truth ทุกโมดูลอ้างอิง"],
            ["Retriever & Pointer Software Approach", "แยก Inbreeding Evaluation กับ Mating Optimization เป็น 2 ขั้นตอน", "Genetic Analysis แยก Compute F กับ Pairing Simulator"],
        ],
        col_widths=[3.5, 5.0, 5.0]
    )

    # Population management framework diagram
    fig, ax = plt.subplots(1, 1, figsize=(10, 5.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis('off')
    ax.text(5, 5.7, 'กรอบแนวคิด Population Management (อ้างอิง EAZA/Species360)', ha='center', fontsize=12, fontweight='bold')

    layers = [
        (0.5, 4.2, 9.0, 1.0, '#1a5276', 'Species360/ZIMS Standards\nTaxonomy | Global ID | Data Exchange Protocol'),
        (0.5, 2.8, 4.0, 1.0, '#2980b9', 'Studbook Management\nAuto-Numbering | DOB Ordering\n7 Statuses | 10 Fields'),
        (5.0, 2.8, 4.5, 1.0, '#27ae60', 'Genetic Analysis Engine\nWright F-Coefficient | Multi-Path\nCI Index | Pairing Simulator'),
        (0.5, 1.4, 4.0, 1.0, '#8e44ad', 'Pedigree Visualization\nD3.js Interactive Tree\nThreshold Coloring | Export'),
        (5.0, 1.4, 4.5, 1.0, '#e67e22', 'Population Reports & Dashboard\nInbreeding | Age Pyramid\nBirth Seasonality | Structure'),
        (0.5, 0.2, 9.0, 0.8, '#16a085', 'Animal Registry — Single Source of Truth (Animal ID)'),
    ]
    for x, y, w, h, color, label in layers:
        box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08", facecolor=color, alpha=0.9)
        ax.add_patch(box)
        ax.text(x + w/2, y + h/2, label, ha='center', va='center', fontsize=8, color='white', fontweight='bold')

    for sy, ey in [(4.2, 3.8), (2.8, 2.4), (2.8, 2.4)]:
        ax.annotate('', xy=(5, ey), xytext=(5, sy), arrowprops=dict(arrowstyle='->', color='#333', lw=1.2))
    fig.tight_layout()
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "กรอบแนวคิด Population Management อ้างอิงมาตรฐาน EAZA/Species360")

    # ── 3.6.15.2 Veterinary Billing Best Practices ──
    add_subsection_heading(doc, "3.6.15.2", "แนวปฏิบัติ Veterinary Practice Management สำหรับระบบ Billing")
    add_body_after_subsection(doc,
        "การออกแบบระบบ Billing สำหรับสัตวแพทย์อ้างอิงแนวปฏิบัติจากอุตสาหกรรม Veterinary Practice Management "
        "ซึ่งงานวิจัยจาก AAHA (American Animal Hospital Association) พบว่าโรงพยาบาลสัตว์สูญเสียรายได้สูงถึง 17% "
        "จากบริการวินิจฉัยเพียงอย่างเดียว เนื่องจากการบันทึกค่าใช้จ่ายด้วยมือและความผิดพลาดจากมนุษย์ "
        "แนวปฏิบัติสำคัญ 5 ประการที่นำมาใช้ในระบบ Thai Zoo ARK ดังแสดงในตารางที่ {tn}".format(
            tn="3.6-" + str(table_counter + 1))
    )
    tn = next_table_num()
    add_table_caption(doc, tn, "แนวปฏิบัติ Veterinary Billing สำหรับ Thai Zoo ARK")
    add_table(doc,
        ["หลักการ", "แนวปฏิบัติสากล", "การนำมาใช้ใน Thai Zoo ARK"],
        [
            ["Charge Capture Integrity", "ทุกหัตถการและยาที่ใช้ต้องถูกบันทึกเป็นค่าใช้จ่ายอัตโนมัติ ไม่พึ่งพาการป้อนด้วยมือ", "Auto-Link หัตถการ/ยา กับ Rate Card (TOR 4.18.1.3, 4.18.4.4)"],
            ["Revenue Leakage Prevention", "ลดการสูญเสียรายได้จากค่าบริการที่ไม่ได้บันทึก โดยเฉลี่ยอุตสาหกรรมสูญเสีย 5-17%", "คำนวณค่าใช้จ่ายอัตโนมัติ แยกหัตถการ/ยา (TOR 4.18.5.1)"],
            ["Eligibility-Based Billing", "ตรวจสอบสิทธิ์ก่อนคิดค่า: สมาชิกภายในไม่คิด ภายนอกคิดตามเงื่อนไข", "ระบบตรวจสอบสิทธิ์อัตโนมัติ (TOR 4.18.2)"],
            ["Audit Trail ทุกขั้นตอน", "ทุกการเปลี่ยนแปลง Rate Card/Invoice ต้องมี Audit Trail สำหรับตรวจสอบ", "Full Audit Trail ตั้งแต่บันทึกถึงชำระ (TOR 4.18.7.1)"],
            ["Digital Invoice + Export", "ออก Invoice ดิจิทัลพร้อมส่งออก PDF/Excel ลดกระดาษและข้อผิดพลาด", "Generate Invoice + Export PDF/Excel (TOR 4.18.5.4, 4.18.7.3)"],
        ],
        col_widths=[3.0, 5.5, 5.0]
    )

    # ── 3.6.15.3 OWASP Security Logging ──
    add_subsection_heading(doc, "3.6.15.3", "มาตรฐาน OWASP Security Logging and Monitoring")
    add_body_after_subsection(doc,
        "การออกแบบระบบ System Log อ้างอิงมาตรฐาน OWASP Top 10 ข้อ A09:2021 "
        "เรื่อง Security Logging and Monitoring Failures ซึ่งระบุว่าการบันทึกและตรวจสอบ Log "
        "ที่ไม่เพียงพอเป็นหนึ่งในช่องโหว่สำคัญที่ทำให้องค์กรไม่สามารถตรวจจับการโจมตีได้ทันเวลา "
        "จากรายงาน IBM Cost of a Data Breach 2023 องค์กรที่ตรวจจับและควบคุมเหตุการณ์ภายใน 200 วัน "
        "ประหยัดค่าใช้จ่ายเฉลี่ย 1.02 ล้านดอลลาร์สหรัฐ เทียบกับองค์กรที่ใช้เวลานานกว่า "
        "หลักการ 7 ประการตาม OWASP ที่นำมาใช้แสดงในตารางที่ {tn}".format(
            tn="3.6-" + str(table_counter + 1))
    )
    tn = next_table_num()
    add_table_caption(doc, tn, "หลักการ OWASP Security Logging ที่นำมาใช้ใน Thai Zoo ARK")
    add_table(doc,
        ["หลักการ OWASP", "คำอธิบาย", "การนำมาใช้ใน Thai Zoo ARK"],
        [
            ["บันทึก Auditable Events", "Login, Failed Login, High-Value Transactions ต้องบันทึกทุกรายการ", "System Log บันทึก Login/Logout + ทุก CRUD Action (TOR 4.21)"],
            ["Log Integrity Protection", "ป้องกันการแก้ไข Log โดยไม่ได้รับอนุญาต", "Append-Only Log Store + Partitioned Table"],
            ["Centralized Logging", "Log จากทุกโมดูลต้องรวมศูนย์ ไม่กระจาย", "Event Collector รวมศูนย์จากทุกโมดูล"],
            ["Configurable Retention", "กำหนด Retention Period ตามนโยบาย (90-365 วัน)", "Retention ตั้งค่าได้ ค่าเริ่มต้น 90 วัน (TOR 4.21.5)"],
            ["Alerting on Anomalies", "แจ้งเตือนเมื่อพบพฤติกรรมผิดปกติ", "Alert Rules สำหรับ Failed Login, Mass Delete, Unauthorized Access"],
            ["Structured Log Format", "Log ในรูปแบบ Structured (JSON) เพื่อค้นหาและวิเคราะห์ได้ง่าย", "JSON Format + ISO 8601 Timestamp + Module Tag"],
            ["Sensitive Data Handling", "ไม่บันทึกข้อมูลส่วนบุคคล/รหัสผ่านใน Log", "Mask PII + ไม่บันทึก Password/Token ใน Log"],
        ],
        col_widths=[3.0, 5.0, 5.5]
    )

    # OWASP Logging Architecture Diagram
    fig, ax = plt.subplots(1, 1, figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5.5); ax.axis('off')
    ax.text(5, 5.2, 'สถาปัตยกรรม Security Logging ตามมาตรฐาน OWASP', ha='center', fontsize=12, fontweight='bold')

    # Sources
    sources = [
        (0.2, 3.5, 1.6, 0.8, '#2c3e50', 'Staff Mgmt'),
        (0.2, 2.5, 1.6, 0.8, '#2980b9', 'Animal Reg'),
        (0.2, 1.5, 1.6, 0.8, '#27ae60', 'Billing'),
        (0.2, 0.5, 1.6, 0.8, '#8e44ad', 'All Other\nModules'),
    ]
    for x, y, w, h, c, l in sources:
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05", facecolor=c, alpha=0.85))
        ax.text(x+w/2, y+h/2, l, ha='center', va='center', fontsize=7, color='white', fontweight='bold')
        ax.annotate('', xy=(2.3, y+h/2), xytext=(1.8, y+h/2), arrowprops=dict(arrowstyle='->', color='#444', lw=1))

    # Pipeline
    pipeline = [
        (2.5, 2.0, 1.8, 1.5, '#e74c3c', 'Event\nCollector'),
        (4.8, 2.0, 1.8, 1.5, '#e67e22', 'Normalize\n+ Enrich\n(JSON)'),
        (7.0, 2.0, 1.5, 1.5, '#16a085', 'Log\nStore'),
    ]
    for x, y, w, h, c, l in pipeline:
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08", facecolor=c, alpha=0.9))
        ax.text(x+w/2, y+h/2, l, ha='center', va='center', fontsize=8, color='white', fontweight='bold')

    ax.annotate('', xy=(4.8, 2.75), xytext=(4.3, 2.75), arrowprops=dict(arrowstyle='->', color='#444', lw=1.5))
    ax.annotate('', xy=(7.0, 2.75), xytext=(6.6, 2.75), arrowprops=dict(arrowstyle='->', color='#444', lw=1.5))

    # Outputs
    outputs = [
        (8.8, 3.5, 1.0, 0.8, '#1a5276', 'Search\n& Filter'),
        (8.8, 2.5, 1.0, 0.8, '#c0392b', 'Alert\nRules'),
        (8.8, 1.5, 1.0, 0.8, '#7f8c8d', 'Retention\nPolicy'),
        (8.8, 0.5, 1.0, 0.8, '#d4ac0d', 'Audit\nReview'),
    ]
    for x, y, w, h, c, l in outputs:
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05", facecolor=c, alpha=0.85))
        ax.text(x+w/2, y+h/2, l, ha='center', va='center', fontsize=7, color='white', fontweight='bold')
        ax.annotate('', xy=(8.8, y+h/2), xytext=(8.5, 2.75), arrowprops=dict(arrowstyle='->', color='#444', lw=1))

    fig.tight_layout()
    add_image_from_fig(doc, fig)
    dn = next_diagram_num()
    add_figure_caption(doc, dn, "สถาปัตยกรรม Security Logging ตามมาตรฐาน OWASP")

    # ── 3.6.15.4 Benchmark Mapping Summary ──
    add_subsection_heading(doc, "3.6.15.4", "สรุป Benchmark Mapping ระหว่างมาตรฐานสากลกับโมดูลใหม่")
    add_body_after_subsection(doc,
        "ตารางที่ {tn} สรุปการเชื่อมโยงระหว่างมาตรฐานและแนวปฏิบัติสากลที่อ้างอิง "
        "กับการออกแบบและพัฒนาในแต่ละโมดูลใหม่ของระบบ Thai Zoo ARK "
        "เพื่อยืนยันว่าทุกโมดูลมีพื้นฐานจากมาตรฐานที่ได้รับการยอมรับในระดับสากล".format(
            tn="3.6-" + str(table_counter + 1))
    )
    tn = next_table_num()
    add_table_caption(doc, tn, "Benchmark Mapping: มาตรฐานสากล กับ โมดูลใหม่ Thai Zoo ARK")
    add_table(doc,
        ["โมดูลใหม่", "มาตรฐาน/Benchmark อ้างอิง", "หลักการสำคัญที่นำมาใช้", "TOR"],
        [
            ["Studbook", "Species360 ZIMS Studbook, AZA Records Keeping", "Taxonomy Linking, Auto-Numbering, 7 Statuses", "4.14"],
            ["Genetic Analysis", "Wright's Path Coefficient, EAZA Pop. Mgmt Manual, Retriever & Pointer Software", "F-Coefficient Calculation, Multi-Path Ancestor, CI Index", "4.15"],
            ["Pedigree Diagram", "EAZA Pedigree Analysis for Group Species, ZIMS Visualization", "Interactive Tree, Threshold Coloring, PDF/Excel/Image Export", "4.16"],
            ["Reports & Dashboard", "EAZA Population Management Manual, Species360 Analytics", "Inbreeding Report, Age Pyramid, Birth Seasonality, Population Structure", "4.17"],
            ["Veterinary Billing", "AAHA Revenue Standards, Digitail Secure Payments, PIMS Best Practices", "Charge Capture Integrity, Eligibility Check, Full Audit Trail", "4.18"],
            ["Change Request", "ITIL Change Management, ISO 20000", "Category Management, Approval Workflow, Status Tracking", "4.19"],
            ["Manual CMS", "Knowledge Management ISO 30401, WCAG 2.2 Accessibility", "PDF/Video Upload, Module-Based Search, Video Manual 5 Modules", "4.20"],
            ["System Log", "OWASP A09:2021, NIST SP 800-92, CIS Log Management Benchmark", "Centralized Logging, JSON Format, Retention Policy, Alert Rules", "4.21"],
        ],
        col_widths=[2.0, 4.0, 4.5, 1.0]
    )


if __name__ == "__main__":
    build_document()
