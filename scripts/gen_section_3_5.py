#!/usr/bin/env python3
"""Generate Section 3.5 UX/UI: Responsive + Touch + Barcode-first
for Thai Zoo ARK technical proposal.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from io import BytesIO

# ── Thai Font Registration ──
_FONT_DIR = os.path.expanduser("~/Library/Fonts/")
for _fname in [
    "THSarabunNew.ttf", "THSarabunNew Bold.ttf",
    "THSarabunNew Italic.ttf", "THSarabunNew BoldItalic.ttf",
    "THSarabun.ttf", "THSarabun Bold.ttf",
]:
    _path = os.path.join(_FONT_DIR, _fname)
    if os.path.exists(_path):
        fm.fontManager.addfont(_path)
plt.rcParams['font.family'] = 'TH Sarabun New'
plt.rcParams['font.sans-serif'] = ['TH Sarabun New', 'TH SarabunPSK', 'DejaVu Sans']

# ── Constants ──
FONT_TH    = "TH Sarabun New"
BODY_SIZE  = Pt(16)
HEAD_SIZE  = Pt(16)
CAP_SIZE   = Pt(14)
HEADER_BG  = "D6E3F8"
OUT_DIR    = os.path.expanduser("~/Library/CloudStorage/OneDrive-Personal/Personal/ZTL")
OUT_FILE   = os.path.join(OUT_DIR, "Section_3_5_UX_UI_Responsive_Touch_Barcode.docx")

# ════════════════════════════════════════════════════════════
#  LOW-LEVEL HELPERS
# ════════════════════════════════════════════════════════════

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _ensure_cs(run, font_name, size, bold):
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:cs="{font_name}" '
            f'w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        )
        rPr.insert(0, rF)
    else:
        for a in ('cs', 'ascii', 'hAnsi'):
            rF.set(qn(f'w:{a}'), font_name)
    sz = rPr.find(qn('w:szCs'))
    v = str(int(size.pt * 2))
    if sz is None:
        rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="{v}"/>'))
    else:
        sz.set(qn('w:val'), v)
    if bold and rPr.find(qn('w:bCs')) is None:
        rPr.append(parse_xml(f'<w:bCs {nsdecls("w")}/>'))


def sfont(run, font_name=FONT_TH, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    _ensure_cs(run, font_name, size, bold)
    if color:
        run.font.color.rgb = color


def _heading(doc, num, title, tab, left, hang, bef, aft):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="{bef}" w:after="{aft}"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="{left}" w:hanging="{hang}"/>'))
    pPr.append(parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="{tab}"/></w:tabs>'))
    run = p.add_run(f"{num}\t{title}")
    sfont(run, bold=True, size=HEAD_SIZE)
    return p


def h_chapter(doc, n, t):    return _heading(doc, n, t, 567,  567,  567, 240, 120)
def h_section(doc, n, t):    return _heading(doc, n, t, 1418, 1418, 851, 200,  80)
def h_subsection(doc, n, t): return _heading(doc, n, t, 2268, 2268, 850, 160,  80)


def _body(doc, text, fi):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="120" w:line="240" w:lineRule="auto"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="{fi}"/>'))
    run = p.add_run(text)
    sfont(run)
    return p


def body_ch(doc, t):  return _body(doc, t,  567)
def body_sec(doc, t): return _body(doc, t, 1418)
def body_sub(doc, t): return _body(doc, t, 2268)


def nitem_ch(doc, n, t):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="60" w:line="240" w:lineRule="auto"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="927" w:hanging="360"/>'))
    run = p.add_run(f"{n}) {t}")
    sfont(run)
    return p


def nitem_sec(doc, n, t):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="60" w:line="240" w:lineRule="auto"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="1778" w:hanging="360"/>'))
    run = p.add_run(f"{n}) {t}")
    sfont(run)
    return p


def nitem_sub(doc, n, t):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="60" w:line="240" w:lineRule="auto"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="2628" w:hanging="360"/>'))
    run = p.add_run(f"{n}) {t}")
    sfont(run)
    return p


def tbl_cap(doc, t):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="240" w:after="120"/>'))
    run = p.add_run(t)
    sfont(run, bold=True)


def fig_cap(doc, t):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="240" w:after="120"/>'))
    run = p.add_run(t)
    sfont(run, bold=True, italic=True, size=CAP_SIZE)


def tbl(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # header row
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, HEADER_BG)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.paragraphs[0].add_run(h)
        sfont(run, bold=True)
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            run = table.rows[ri + 1].cells[ci].paragraphs[0].add_run(str(val))
            sfont(run)
    return table


def add_img(doc, fig, w=5.8):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Inches(w))
    plt.close(fig)


# ════════════════════════════════════════════════════════════
#  DIAGRAM GENERATORS
# ════════════════════════════════════════════════════════════

def _box(ax, x, y, w, h, color, label, fs=9, text_color='white'):
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08",
                 facecolor=color, alpha=0.90, edgecolor='#222', linewidth=1.2))
    ax.text(x + w / 2, y + h / 2, label, ha='center', va='center',
            fontsize=fs, color=text_color, fontweight='bold')


def _arrow(ax, x1, y1, x2, y2, label='', c='#333', lw=1.5, cs='arc3,rad=0'):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=c, lw=lw,
                                connectionstyle=cs))
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(mx, my + 0.12, label, ha='center', fontsize=7.5, color=c)


# ── Diagram 1: Responsive Breakpoints ──────────────────────
def diagram_responsive_breakpoints():
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6); ax.axis('off')
    ax.text(5.5, 5.7, 'ระบบออกแบบ Responsive Design — Breakpoints และ Layout Grid',
            ha='center', fontsize=13, fontweight='bold')

    # Desktop
    ax.add_patch(FancyBboxPatch((0.3, 1.0), 3.2, 4.0, boxstyle="round,pad=0.1",
                 facecolor='#EBF5FB', edgecolor='#2E86C1', linewidth=2.0))
    ax.text(1.9, 4.7, 'Desktop (> 1,024px)', ha='center', fontsize=10,
            color='#2E86C1', fontweight='bold')
    # Desktop layout: sidebar + main
    ax.add_patch(plt.Rectangle((0.5, 1.2), 0.7, 3.2, facecolor='#AED6F1', alpha=0.8))
    ax.text(0.85, 2.8, 'Side\nbar', ha='center', fontsize=8, color='#1B4F72')
    ax.add_patch(plt.Rectangle((1.3, 1.2), 2.0, 3.2, facecolor='#D6EAF8', alpha=0.8))
    ax.text(2.3, 2.8, 'Main\nContent\nArea', ha='center', fontsize=8, color='#1B4F72')
    # grid lines inside main
    for yy in [1.8, 2.4, 3.0, 3.6, 4.2]:
        ax.plot([1.3, 3.3], [yy, yy], color='#AED6F1', lw=0.6, alpha=0.7)

    # Tablet
    ax.add_patch(FancyBboxPatch((3.9, 1.2), 2.4, 3.6, boxstyle="round,pad=0.1",
                 facecolor='#EAFAF1', edgecolor='#1E8449', linewidth=2.0))
    ax.text(5.1, 4.5, 'Tablet (768–1,024px)', ha='center', fontsize=10,
            color='#1E8449', fontweight='bold')
    # Tablet: 2-col grid
    ax.add_patch(plt.Rectangle((4.0, 1.4), 1.1, 2.8, facecolor='#A9DFBF', alpha=0.8))
    ax.text(4.55, 2.8, 'Col 1', ha='center', fontsize=8, color='#145A32')
    ax.add_patch(plt.Rectangle((5.2, 1.4), 1.0, 2.8, facecolor='#D5F5E3', alpha=0.8))
    ax.text(5.7, 2.8, 'Col 2', ha='center', fontsize=8, color='#145A32')

    # Mobile
    ax.add_patch(FancyBboxPatch((6.9, 1.5), 1.6, 3.0, boxstyle="round,pad=0.1",
                 facecolor='#FDEDEC', edgecolor='#C0392B', linewidth=2.0))
    ax.text(7.7, 4.2, 'Mobile (< 768px)', ha='center', fontsize=10,
            color='#C0392B', fontweight='bold')
    # Mobile: single column
    ax.add_patch(plt.Rectangle((7.0, 1.7), 1.4, 2.6, facecolor='#F9EBEA', alpha=0.8))
    ax.text(7.7, 3.0, 'Full\nWidth\nCol', ha='center', fontsize=8, color='#922B21')

    # Touch Screen Kiosk
    ax.add_patch(FancyBboxPatch((9.0, 1.0), 1.8, 4.0, boxstyle="round,pad=0.1",
                 facecolor='#F4ECF7', edgecolor='#7D3C98', linewidth=2.0))
    ax.text(9.9, 4.7, 'Touch Kiosk', ha='center', fontsize=10,
            color='#7D3C98', fontweight='bold')
    # Kiosk: big buttons
    for j, (label, c) in enumerate([('รับเข้า', '#8E44AD'), ('เบิก', '#9B59B6'), ('คืน', '#A569BD'), ('สต็อก', '#BB8FCE')]):
        ax.add_patch(FancyBboxPatch((9.1, 1.2 + j * 0.8), 1.5, 0.65,
                     boxstyle="round,pad=0.06", facecolor=c, alpha=0.85, edgecolor='white', linewidth=1))
        ax.text(9.85, 1.52 + j * 0.8, label, ha='center', va='center',
                fontsize=9, color='white', fontweight='bold')

    # Labels below
    ax.text(1.9, 0.7, 'Full navigation\n+ sidebar', ha='center', fontsize=8.5, color='#2E86C1')
    ax.text(5.1, 0.7, 'Collapsed nav\n2-col grid', ha='center', fontsize=8.5, color='#1E8449')
    ax.text(7.7, 0.7, 'Bottom nav\nStacked single col', ha='center', fontsize=8.5, color='#C0392B')
    ax.text(9.9, 0.7, 'Large buttons\nTouch-first', ha='center', fontsize=8.5, color='#7D3C98')

    # Touch target note
    ax.text(5.5, 0.15, '* ทุก interactive element ต้องมีขนาด Touch Target ≥ 44 × 44 px (WCAG 2.2)',
            ha='center', fontsize=8, color='#555', style='italic')

    fig.tight_layout()
    return fig


# ── Diagram 2: Touch UI Wireframe – Food Warehouse ─────────
def diagram_touch_wireframe_food():
    fig, axes = plt.subplots(1, 2, figsize=(12, 7))
    fig.suptitle('หน้าจอ Touch Screen — ระบบคลังอาหารสัตว์ (Food Warehouse)\n'
                 'ซ้าย: เมนูหลัก | ขวา: หน้าจอรับเข้าวัตถุดิบ',
                 fontsize=11, fontweight='bold')

    # ── Left: Main Menu ──
    ax = axes[0]
    ax.set_xlim(0, 6); ax.set_ylim(0, 9); ax.axis('off')
    # Phone frame
    ax.add_patch(FancyBboxPatch((0.3, 0.3), 5.4, 8.4, boxstyle="round,pad=0.2",
                 facecolor='#F8F9FA', edgecolor='#495057', linewidth=2.5))
    # Header
    ax.add_patch(plt.Rectangle((0.3, 7.6), 5.4, 1.1, facecolor='#1A5276'))
    ax.text(3.0, 8.15, '🐘  Thai Zoo ARK — คลังอาหารสัตว์', ha='center', fontsize=10,
            color='white', fontweight='bold')

    # Big action buttons
    btns = [
        (0.5, 5.5, 2.2, 1.8, '#2E86C1', '📥\nรับเข้า\nวัตถุดิบ'),
        (3.2, 5.5, 2.2, 1.8, '#1E8449', '📤\nเบิก\nวัตถุดิบ'),
        (0.5, 3.3, 2.2, 1.8, '#7D3C98', '↩\nคืน\nวัตถุดิบ'),
        (3.2, 3.3, 2.2, 1.8, '#B7950B', '📋\nStock\nCard'),
        (0.5, 1.1, 2.2, 1.8, '#117A65', '📊\nสต็อก\nคงเหลือ'),
        (3.2, 1.1, 2.2, 1.8, '#922B21', '🖨\nพิมพ์\nบาร์โค้ด'),
    ]
    for bx, by, bw, bh, bc, bl in btns:
        ax.add_patch(FancyBboxPatch((bx, by), bw, bh, boxstyle="round,pad=0.12",
                     facecolor=bc, alpha=0.92, edgecolor='white', linewidth=1.5))
        ax.text(bx + bw / 2, by + bh / 2, bl, ha='center', va='center',
                fontsize=10, color='white', fontweight='bold')

    ax.text(3.0, 0.5, 'ขนาดปุ่ม ≥ 44×44px | Touch-optimized', ha='center',
            fontsize=7.5, color='#666', style='italic')

    # ── Right: Inbound screen ──
    ax2 = axes[1]
    ax2.set_xlim(0, 6); ax2.set_ylim(0, 9); ax2.axis('off')
    ax2.add_patch(FancyBboxPatch((0.3, 0.3), 5.4, 8.4, boxstyle="round,pad=0.2",
                  facecolor='#F8F9FA', edgecolor='#495057', linewidth=2.5))
    # Header
    ax2.add_patch(plt.Rectangle((0.3, 7.6), 5.4, 1.1, facecolor='#2E86C1'))
    ax2.text(3.0, 8.15, '← รับเข้าวัตถุดิบ (Inbound)', ha='center', fontsize=10,
             color='white', fontweight='bold')

    # Barcode scan area
    ax2.add_patch(FancyBboxPatch((0.6, 6.2), 4.7, 1.2, boxstyle="round,pad=0.08",
                  facecolor='#EBF5FB', edgecolor='#2E86C1', linewidth=1.5, linestyle='--'))
    ax2.text(3.0, 6.82, '[ แสกนบาร์โค้ด หรือ กรอก Lot No. ]', ha='center',
             fontsize=9, color='#2E86C1', style='italic')
    ax2.text(3.0, 6.4, '📷  กด เพื่อเปิดกล้องสแกน', ha='center', fontsize=8.5, color='#555')

    # Form fields
    fields = [
        (0.6, 5.1, 4.7, 0.85, '#ECF0F1', 'ชื่อวัตถุดิบ: [ข้าวโพด — อัตโนมัติจากบาร์โค้ด]'),
        (0.6, 4.0, 4.7, 0.85, '#ECF0F1', 'จำนวนรับเข้า (กก.): [____]'),
        (0.6, 2.9, 4.7, 0.85, '#ECF0F1', 'Lot / Batch: [____]'),
        (0.6, 1.8, 4.7, 0.85, '#ECF0F1', 'วันหมดอายุ: [📅 เลือกวันที่]'),
    ]
    for fx, fy, fw, fh, fc, fl in fields:
        ax2.add_patch(FancyBboxPatch((fx, fy), fw, fh, boxstyle="round,pad=0.06",
                      facecolor=fc, edgecolor='#BDC3C7', linewidth=1.0))
        ax2.text(fx + 0.15, fy + fh / 2, fl, va='center', fontsize=8.5, color='#2C3E50')

    # Confirm button
    ax2.add_patch(FancyBboxPatch((1.0, 0.5), 3.8, 0.95, boxstyle="round,pad=0.1",
                  facecolor='#27AE60', edgecolor='white', linewidth=1.5))
    ax2.text(2.9, 0.97, '✔  ยืนยันรับเข้า', ha='center', fontsize=11,
             color='white', fontweight='bold')

    fig.tight_layout()
    return fig


# ── Diagram 3: Barcode Workflow Sequence ───────────────────
def diagram_barcode_sequence():
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12); ax.set_ylim(0, 8); ax.axis('off')
    ax.text(6.0, 7.7, 'ลำดับขั้นตอนการทำงาน Barcode Workflow\n'
            '(ระบบคลังอาหารสัตว์และยา/เวชภัณฑ์)',
            ha='center', fontsize=12, fontweight='bold')

    # Swimlane headers
    lanes = [
        (0.0, 2.5,  '#D6EAF8', 'เจ้าหน้าที่\nคลัง'),
        (2.5, 2.5,  '#D5F5E3', 'หน้าจอ\nTouch UI'),
        (5.0, 2.5,  '#FDEBD0', 'API Server\n(Backend)'),
        (7.5, 2.5,  '#F9EBEA', 'ฐานข้อมูล\nคลัง'),
        (10.0, 2.0, '#F4ECF7', 'เครื่อง\nพิมพ์'),
    ]
    for lx, lw, lc, lt in lanes:
        ax.add_patch(plt.Rectangle((lx, 0), lw, 7.5, facecolor=lc, alpha=0.35, edgecolor='#BDC3C7'))
        ax.text(lx + lw / 2, 7.2, lt, ha='center', va='center', fontsize=9,
                fontweight='bold', color='#2C3E50')

    # Lifeline dashes
    lx_centers = [1.25, 3.75, 6.25, 8.75, 11.0]
    for cx in lx_centers:
        ax.plot([cx, cx], [0.3, 6.8], color='#888', lw=1.0, linestyle='--', alpha=0.6)

    # Steps (y from top)
    steps = [
        # (from_cx, to_cx, y, label, color)
        (1.25, 3.75, 6.4,  '1. สแกนบาร์โค้ด / กรอก Lot No.',       '#2E86C1'),
        (3.75, 6.25, 5.8,  '2. ส่ง Barcode / ID → Lookup API',      '#1E8449'),
        (6.25, 8.75, 5.2,  '3. Query ข้อมูลสินค้าจาก DB',            '#7D3C98'),
        (8.75, 6.25, 4.6,  '4. คืนชื่อ, หน่วย, Batch, Expiry',       '#7D3C98'),
        (6.25, 3.75, 4.0,  '5. แสดงรายละเอียดสินค้าบนหน้าจอ',       '#2E86C1'),
        (3.75, 1.25, 3.4,  '6. เจ้าหน้าที่ยืนยัน / ระบุจำนวน',      '#C0392B'),
        (1.25, 3.75, 2.8,  '7. กดปุ่มยืนยัน (รับเข้า/เบิก/คืน)',    '#C0392B'),
        (3.75, 6.25, 2.2,  '8. POST /inventory/transaction',         '#1E8449'),
        (6.25, 8.75, 1.6,  '9. บันทึก Stock Ledger + Audit Log',     '#7D3C98'),
        (6.25, 11.0, 1.0,  '10. ส่งคำสั่งพิมพ์ Label (ถ้ารับเข้าใหม่)', '#E67E22'),
    ]

    for fx, tx, y, lbl, c in steps:
        # Draw arrow
        ax.annotate('', xy=(tx, y), xytext=(fx, y),
                    arrowprops=dict(arrowstyle='->', color=c, lw=1.8,
                                    connectionstyle='arc3,rad=0'))
        # Activity box on source
        ax.add_patch(FancyBboxPatch((fx - 0.25, y - 0.18), 0.5, 0.36,
                     boxstyle="round,pad=0.04", facecolor=c, alpha=0.7, edgecolor='none'))
        # Label
        mid_x = (fx + tx) / 2
        offset = 0.2 if tx > fx else -0.2
        ax.text(mid_x, y + 0.22, lbl, ha='center', fontsize=7.5, color=c, fontweight='bold')

    # Return path marker
    ax.text(0.05, 0.2, '* เส้นทึบ = ส่งข้อมูล   เส้นกลับ = รับข้อมูล/response',
            fontsize=7.5, color='#666', style='italic')

    fig.tight_layout()
    return fig


# ── Diagram 4: Mobile Keeper Interface Wireframe ───────────
def diagram_mobile_keeper():
    fig, axes = plt.subplots(1, 3, figsize=(13, 8))
    fig.suptitle('Wireframe หน้าจอมือถือ — โมดูลผู้ดูแลสัตว์ (Zookeeper Mobile Module)',
                 fontsize=11, fontweight='bold')

    # ── Screen A: My Animals ──
    ax = axes[0]
    ax.set_xlim(0, 5); ax.set_ylim(0, 10); ax.axis('off')
    ax.text(2.5, 9.7, 'หน้ารายชื่อสัตว์ที่ดูแล', ha='center', fontsize=10, fontweight='bold', color='#1A5276')

    # Phone outline
    ax.add_patch(FancyBboxPatch((0.2, 0.2), 4.6, 9.2, boxstyle="round,pad=0.15",
                 facecolor='#F8F9FA', edgecolor='#34495E', linewidth=2.5))
    # Status bar
    ax.add_patch(plt.Rectangle((0.2, 8.9), 4.6, 0.5, facecolor='#1A5276'))
    ax.text(2.5, 9.15, '09:41  🔋 97%  📶', ha='center', fontsize=8, color='white')

    # Nav header
    ax.add_patch(plt.Rectangle((0.2, 8.2), 4.6, 0.65, facecolor='#2E86C1'))
    ax.text(2.5, 8.52, '🐾 สัตว์ของฉัน (12 ตัว)', ha='center', fontsize=9,
            color='white', fontweight='bold')

    # Search bar
    ax.add_patch(FancyBboxPatch((0.4, 7.6), 4.2, 0.45, boxstyle="round,pad=0.05",
                 facecolor='#EBF5FB', edgecolor='#AED6F1', linewidth=1))
    ax.text(0.7, 7.82, '🔍  ค้นหา...', fontsize=8, color='#AAA')

    # Animal cards
    animals = [
        ('🐘 ช้างไทย "มะลิ"', 'Female | 12 ปี', '#D5F5E3', '#1E8449'),
        ('🦒 ยีราฟ "ลองเน็ค"', 'Male | 7 ปี', '#D6EAF8', '#2E86C1'),
        ('🐯 เสือ "ราชา"', 'Male | 5 ปี | ⚠ ป่วย', '#FDEDEC', '#C0392B'),
        ('🦁 สิงโต "ซาฟารี"', 'Female | 9 ปี', '#FEF9E7', '#B7950B'),
    ]
    for i, (name, detail, bg, tc) in enumerate(animals):
        yy = 6.8 - i * 1.35
        ax.add_patch(FancyBboxPatch((0.4, yy), 4.2, 1.1, boxstyle="round,pad=0.07",
                     facecolor=bg, edgecolor=tc, linewidth=1.2))
        ax.text(0.7, yy + 0.75, name, fontsize=9, color=tc, fontweight='bold')
        ax.text(0.7, yy + 0.35, detail, fontsize=8, color='#555')
        ax.text(4.35, yy + 0.55, '›', fontsize=14, color=tc, ha='center', va='center')

    # Bottom nav
    ax.add_patch(plt.Rectangle((0.2, 0.2), 4.6, 0.75, facecolor='#ECF0F1'))
    for ni, (icon, label) in enumerate([('🐾','สัตว์ของฉัน'),('📋','บันทึก'),('🔔','แจ้งเตือน'),('👤','โปรไฟล์')]):
        nx = 0.8 + ni * 1.05
        ax.text(nx, 0.62, icon, ha='center', fontsize=10)
        ax.text(nx, 0.32, label, ha='center', fontsize=6.5, color='#555')
    ax.plot([0.2, 4.8], [0.95, 0.95], color='#CCC', lw=0.8)

    # ── Screen B: Quick Actions ──
    ax2 = axes[1]
    ax2.set_xlim(0, 5); ax2.set_ylim(0, 10); ax2.axis('off')
    ax2.text(2.5, 9.7, 'หน้าบันทึกรายวัน', ha='center', fontsize=10, fontweight='bold', color='#1A5276')

    ax2.add_patch(FancyBboxPatch((0.2, 0.2), 4.6, 9.2, boxstyle="round,pad=0.15",
                  facecolor='#F8F9FA', edgecolor='#34495E', linewidth=2.5))
    ax2.add_patch(plt.Rectangle((0.2, 8.9), 4.6, 0.5, facecolor='#1A5276'))
    ax2.text(2.5, 9.15, '09:41  🔋 97%  📶', ha='center', fontsize=8, color='white')
    ax2.add_patch(plt.Rectangle((0.2, 8.2), 4.6, 0.65, facecolor='#2E86C1'))
    ax2.text(2.5, 8.52, '← 🐘 มะลิ — การดูแลวันนี้', ha='center', fontsize=8.5,
             color='white', fontweight='bold')

    actions = [
        ('🌾', 'รายการรับอาหาร',             '#2ECC71', '#1A5276'),
        ('🌡', 'ตรวจสภาพแวดล้อม',           '#3498DB', '#1A5276'),
        ('🎭', 'ส่งเสริมคุณภาพชีวิต',        '#9B59B6', '#F9F9F9'),
        ('🍼', 'แจ้งเกิด',                   '#E67E22', '#F9F9F9'),
        ('🤒', 'แจ้งป่วย',                   '#E74C3C', '#F9F9F9'),
        ('🤰', 'แจ้งตั้งครรภ์/วางไข่',       '#F39C12', '#1A5276'),
        ('💔', 'แจ้งตาย',                    '#95A5A6', '#F9F9F9'),
        ('❓', 'แจ้งสูญหาย',                  '#7F8C8D', '#F9F9F9'),
    ]
    for i, (icon, label, bg, tc) in enumerate(actions):
        row = i // 2; col = i % 2
        bx = 0.4 + col * 2.2; by = 6.8 - row * 1.55
        ax2.add_patch(FancyBboxPatch((bx, by), 1.9, 1.25, boxstyle="round,pad=0.1",
                      facecolor=bg, alpha=0.88, edgecolor='#DDD', linewidth=1))
        ax2.text(bx + 0.95, by + 0.85, icon, ha='center', fontsize=14)
        ax2.text(bx + 0.95, by + 0.35, label, ha='center', fontsize=7.5,
                 color=tc, fontweight='bold')

    ax2.add_patch(plt.Rectangle((0.2, 0.2), 4.6, 0.75, facecolor='#ECF0F1'))
    for ni, (icon, label) in enumerate([('🐾','สัตว์ของฉัน'),('📋','บันทึก'),('🔔','แจ้งเตือน'),('👤','โปรไฟล์')]):
        nx = 0.8 + ni * 1.05
        ax2.text(nx, 0.62, icon, ha='center', fontsize=10)
        ax2.text(nx, 0.32, label, ha='center', fontsize=6.5, color='#555')
    ax2.plot([0.2, 4.8], [0.95, 0.95], color='#CCC', lw=0.8)

    # ── Screen C: Report Birth Form ──
    ax3 = axes[2]
    ax3.set_xlim(0, 5); ax3.set_ylim(0, 10); ax3.axis('off')
    ax3.text(2.5, 9.7, 'แบบฟอร์มแจ้งเหตุการณ์', ha='center', fontsize=10, fontweight='bold', color='#1A5276')

    ax3.add_patch(FancyBboxPatch((0.2, 0.2), 4.6, 9.2, boxstyle="round,pad=0.15",
                  facecolor='#F8F9FA', edgecolor='#34495E', linewidth=2.5))
    ax3.add_patch(plt.Rectangle((0.2, 8.9), 4.6, 0.5, facecolor='#1A5276'))
    ax3.text(2.5, 9.15, '09:41  🔋 97%  📶', ha='center', fontsize=8, color='white')
    ax3.add_patch(plt.Rectangle((0.2, 8.2), 4.6, 0.65, facecolor='#27AE60'))
    ax3.text(2.5, 8.52, '← 🍼 แจ้งเกิด — มะลิ', ha='center', fontsize=9,
             color='white', fontweight='bold')

    form_fields = [
        (7.55, 'วันเวลาที่เกิด *', '📅 [เลือกวันที่และเวลา]'),
        (6.65, 'เพศลูก *',         '○ ตัวผู้   ○ ตัวเมีย   ○ ไม่ทราบ'),
        (5.75, 'น้ำหนักแรกเกิด',   '[____] กิโลกรัม'),
        (4.85, 'สถานะ *',          '○ มีชีวิต   ○ ตายแรกเกิด'),
        (3.95, 'หมายเหตุ',         '[กรอกข้อสังเกตเพิ่มเติม...]'),
    ]
    for fy, flabel, fval in form_fields:
        ax3.text(0.6, fy, flabel, fontsize=8.5, color='#2C3E50', fontweight='bold')
        ax3.add_patch(FancyBboxPatch((0.4, fy - 0.7), 4.2, 0.55,
                      boxstyle="round,pad=0.05", facecolor='white',
                      edgecolor='#AED6F1', linewidth=1))
        ax3.text(0.65, fy - 0.42, fval, fontsize=8, color='#888')

    # Photo attach
    ax3.add_patch(FancyBboxPatch((0.4, 2.9), 4.2, 0.7, boxstyle="round,pad=0.05",
                  facecolor='#EBF5FB', edgecolor='#AED6F1', linewidth=1, linestyle='--'))
    ax3.text(2.5, 3.25, '📷 แนบรูปถ่าย (ไม่บังคับ)', ha='center', fontsize=8.5, color='#2E86C1')

    # Buttons
    ax3.add_patch(FancyBboxPatch((0.4, 1.85), 1.8, 0.8, boxstyle="round,pad=0.08",
                  facecolor='#BDC3C7', edgecolor='none'))
    ax3.text(1.3, 2.25, 'ยกเลิก', ha='center', fontsize=9, color='#2C3E50', fontweight='bold')

    ax3.add_patch(FancyBboxPatch((2.8, 1.85), 1.8, 0.8, boxstyle="round,pad=0.08",
                  facecolor='#27AE60', edgecolor='none'))
    ax3.text(3.7, 2.25, '✔ บันทึก', ha='center', fontsize=9, color='white', fontweight='bold')

    ax3.text(2.5, 1.35, '* ระบบจะแจ้ง Notification ไปยังสัตวแพทย์อัตโนมัติ',
             ha='center', fontsize=7.5, color='#E74C3C', style='italic')

    ax3.add_patch(plt.Rectangle((0.2, 0.2), 4.6, 0.75, facecolor='#ECF0F1'))
    for ni, (icon, label) in enumerate([('🐾','สัตว์ของฉัน'),('📋','บันทึก'),('🔔','แจ้งเตือน'),('👤','โปรไฟล์')]):
        nx = 0.8 + ni * 1.05
        ax3.text(nx, 0.62, icon, ha='center', fontsize=10)
        ax3.text(nx, 0.32, label, ha='center', fontsize=6.5, color='#555')
    ax3.plot([0.2, 4.8], [0.95, 0.95], color='#CCC', lw=0.8)

    fig.tight_layout()
    return fig


# ── Diagram 5: Pharma Touch Wireframe ──────────────────────
def diagram_touch_wireframe_pharma():
    fig, axes = plt.subplots(1, 2, figsize=(11, 7))
    fig.suptitle('หน้าจอ Touch Screen — ระบบยาและเวชภัณฑ์ (Pharma Inventory)\n'
                 'ซ้าย: เมนูหลัก | ขวา: หน้าจอเบิกยา (Issue)',
                 fontsize=11, fontweight='bold')

    # Left: Main Menu
    ax = axes[0]
    ax.set_xlim(0, 5.5); ax.set_ylim(0, 9); ax.axis('off')
    ax.add_patch(FancyBboxPatch((0.2, 0.2), 5.1, 8.4, boxstyle="round,pad=0.15",
                 facecolor='#F8F9FA', edgecolor='#495057', linewidth=2.5))
    ax.add_patch(plt.Rectangle((0.2, 7.6), 5.1, 1.0, facecolor='#117A65'))
    ax.text(2.75, 8.1, '💊  Thai Zoo ARK — ยา/เวชภัณฑ์', ha='center', fontsize=9.5,
            color='white', fontweight='bold')

    pharma_btns = [
        (0.4, 5.3, 2.1, 2.0, '#117A65', '📥\nรับยา\nเข้าคลัง'),
        (2.8, 5.3, 2.1, 2.0, '#2E86C1', '💊\nเบิกยา\nสำหรับสัตว์'),
        (0.4, 2.9, 2.1, 2.0, '#7D3C98', '↩\nคืนยา\nเหลือใช้'),
        (2.8, 2.9, 2.1, 2.0, '#B7950B', '📦\nสต็อก\nคงเหลือ'),
        (0.4, 0.6, 2.1, 2.0, '#922B21', '⏰\nแจ้งเตือน\nใกล้หมดอายุ'),
        (2.8, 0.6, 2.1, 2.0, '#1A5276', '📊\nรายงาน\nการใช้ยา'),
    ]
    for bx, by, bw, bh, bc, bl in pharma_btns:
        ax.add_patch(FancyBboxPatch((bx, by), bw, bh, boxstyle="round,pad=0.12",
                     facecolor=bc, alpha=0.92, edgecolor='white', linewidth=1.5))
        ax.text(bx + bw / 2, by + bh / 2, bl, ha='center', va='center',
                fontsize=9.5, color='white', fontweight='bold')

    # Right: Issue drug screen
    ax2 = axes[1]
    ax2.set_xlim(0, 5.5); ax2.set_ylim(0, 9); ax2.axis('off')
    ax2.add_patch(FancyBboxPatch((0.2, 0.2), 5.1, 8.4, boxstyle="round,pad=0.15",
                  facecolor='#F8F9FA', edgecolor='#495057', linewidth=2.5))
    ax2.add_patch(plt.Rectangle((0.2, 7.6), 5.1, 1.0, facecolor='#2E86C1'))
    ax2.text(2.75, 8.1, '← เบิกยาสำหรับสัตว์', ha='center', fontsize=9.5,
             color='white', fontweight='bold')

    # Animal search
    ax2.add_patch(FancyBboxPatch((0.4, 6.75), 4.7, 0.65, boxstyle="round,pad=0.06",
                  facecolor='#EBF5FB', edgecolor='#AED6F1', linewidth=1.2))
    ax2.text(0.75, 7.08, '🔍 ค้นหาสัตว์ (ID/ชื่อ): [____]', fontsize=8.5, color='#2E86C1')

    # Drug barcode scan
    ax2.add_patch(FancyBboxPatch((0.4, 5.7), 4.7, 0.8, boxstyle="round,pad=0.07",
                  facecolor='#D5F5E3', edgecolor='#1E8449', linewidth=1.5, linestyle='--'))
    ax2.text(2.75, 6.12, '📷 สแกนบาร์โค้ดยา', ha='center', fontsize=9,
             color='#1E8449', fontweight='bold')
    ax2.text(2.75, 5.83, '(หรือกรอก Drug Code ด้วยตนเอง)', ha='center', fontsize=7.5,
             color='#555', style='italic')

    # Drug details card
    ax2.add_patch(FancyBboxPatch((0.4, 4.0), 4.7, 1.5, boxstyle="round,pad=0.08",
                  facecolor='#FDEDEC', edgecolor='#E74C3C', linewidth=1.2))
    ax2.text(0.65, 5.28, '💊 Amoxicillin 500mg', fontsize=9.5, color='#C0392B', fontweight='bold')
    ax2.text(0.65, 4.92, 'Lot: L2024-099  |  Exp: 2025-12-31', fontsize=8, color='#555')
    ax2.text(0.65, 4.58, 'คงเหลือ: 240 แคปซูล  |  แหล่งที่มา: จัดซื้อ', fontsize=8, color='#555')
    ax2.text(0.65, 4.22, '⚠ เชื่อมกับเวชระเบียนโดยอัตโนมัติ', fontsize=7.5, color='#E74C3C')

    # Quantity + dosage
    ax2.add_patch(FancyBboxPatch((0.4, 2.9), 4.7, 0.9, boxstyle="round,pad=0.06",
                  facecolor='#ECF0F1', edgecolor='#BDC3C7', linewidth=1))
    ax2.text(0.65, 3.57, 'จำนวนที่เบิก: [  3  ] แคปซูล', fontsize=9, color='#2C3E50')
    ax2.text(0.65, 3.18, 'หมายเหตุ: [__________________________]', fontsize=8.5, color='#2C3E50')

    # Confirm
    ax2.add_patch(FancyBboxPatch((0.8, 1.9), 3.8, 0.8, boxstyle="round,pad=0.1",
                  facecolor='#2E86C1', edgecolor='white', linewidth=1.5))
    ax2.text(2.7, 2.3, '✔  ยืนยันเบิกยา → บันทึก + Audit Log', ha='center',
             fontsize=8.5, color='white', fontweight='bold')

    ax2.text(2.75, 1.35, 'ระบบจะตัดสต็อกและอัปเดตเวชระเบียนโดยอัตโนมัติ',
             ha='center', fontsize=7.5, color='#1E8449', style='italic')

    fig.tight_layout()
    return fig


# ════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width  = 11906  # A4
    section.page_height = 16838
    for attr in ('left_margin', 'right_margin', 'top_margin', 'bottom_margin'):
        setattr(section, attr, Pt(72))  # 1 inch

    # ────────────────────────────────────────────────────────
    #  3.5  Chapter Heading
    # ────────────────────────────────────────────────────────
    h_chapter(doc, '3.5',
              'UX/UI Design: Responsive, Touch-first และ Barcode-first')

    body_ch(doc,
        'การออกแบบประสบการณ์ผู้ใช้ (User Experience) และส่วนต่อประสานกับผู้ใช้ (User Interface) '
        'ของระบบ Thai Zoo ARK ระยะที่ 2 ตั้งอยู่บนหลักการสำคัญ 3 ประการที่สอดคล้องกับบริบทการ'
        'ปฏิบัติงานจริงของเจ้าหน้าที่สวนสัตว์ทั้ง 7 แห่งทั่วประเทศ ได้แก่ (1) Responsive Design '
        'ที่รองรับอุปกรณ์ทุกประเภทตั้งแต่โทรศัพท์มือถือ แท็บเล็ต จนถึงคอมพิวเตอร์ตั้งโต๊ะ '
        '(2) Touch-first Interface ที่ออกแบบสำหรับการใช้งานบน Touch Screen Kiosk โดยเฉพาะในระบบ'
        'คลังอาหารสัตว์และยา/เวชภัณฑ์ และ (3) Barcode-first Workflow ที่บูรณาการการสแกน'
        'บาร์โค้ดเข้าสู่ขั้นตอนการรับเข้า เบิก และคืนสินค้าทุกรายการ หัวข้อนี้ประกอบด้วยเนื้อหา '
        '4 ส่วน ส่วนแรก (หัวข้อ 3.5.1) อธิบายระบบ Responsive Design ครอบคลุม Breakpoints '
        'ที่กำหนด Touch Target Specification และ Browser Support Matrix '
        'ส่วนที่สอง (หัวข้อ 3.5.2) นำเสนอการออกแบบ Touchscreen Flow สำหรับระบบคลังอาหารสัตว์'
        'และยา/เวชภัณฑ์ พร้อม Wireframe หน้าจอจริงที่ผ่านการออกแบบตาม TOR ข้อ 4.9 และ 4.12 '
        'ส่วนที่สาม (หัวข้อ 3.5.3) อธิบาย Barcode Workflow อย่างละเอียดด้วย Sequence Diagram '
        'แสดงลำดับขั้นตอนตั้งแต่การสแกนจนถึงการบันทึกฐานข้อมูลและพิมพ์ Label '
        'และส่วนที่สี่ (หัวข้อ 3.5.4) แสดง Mobile Interface สำหรับโมดูลผู้ดูแลสัตว์ 9 ฟังก์ชัน '
        'พร้อม Wireframe สามหน้าจอหลัก '
        'ตารางที่ 3.5-1 แสดง Responsive Breakpoint Specification '
        'ตารางที่ 3.5-2 แสดง Touch UX Design Standards '
        'ตารางที่ 3.5-3 แสดง Barcode Technology Specification '
        'และตารางที่ 3.5-4 แสดง Mobile Keeper 9 ฟังก์ชันพร้อม TOR Reference '
        'แผนภาพที่ 3.5-1 แสดงโครงสร้าง Responsive Breakpoints '
        'แผนภาพที่ 3.5-2 แสดง Touch Wireframe ระบบคลังอาหาร '
        'แผนภาพที่ 3.5-3 แสดง Barcode Sequence Diagram '
        'แผนภาพที่ 3.5-4 แสดง Mobile Keeper Wireframe '
        'และแผนภาพที่ 3.5-5 แสดง Touch Wireframe ระบบยา/เวชภัณฑ์'
    )

    # ────────────────────────────────────────────────────────
    #  3.5.1  Responsive Design System
    # ────────────────────────────────────────────────────────
    h_section(doc, '3.5.1', 'Responsive Design System')

    body_sec(doc,
        'ระบบ Responsive Design ของ Thai Zoo ARK ออกแบบตามแนวทาง Mobile-first โดยกำหนด '
        'Breakpoints 3 ระดับที่สอดคล้องกับอุปกรณ์ที่เจ้าหน้าที่ใช้งานจริงในสวนสัตว์ '
        'ได้แก่ Mobile (< 768px) สำหรับผู้ดูแลสัตว์ที่ปฏิบัติงานภาคสนาม '
        'Tablet (768–1,024px) สำหรับสัตวแพทย์และเจ้าหน้าที่คลังที่เคลื่อนที่ภายในอาคาร '
        'และ Desktop (> 1,024px) สำหรับผู้บริหารและเจ้าหน้าที่ธุรการที่ทำงานบนโต๊ะ '
        'นอกจากนี้ยังมีการออกแบบพิเศษสำหรับ Touch Screen Kiosk ที่ติดตั้งในพื้นที่คลังอาหารและ'
        'คลังยา/เวชภัณฑ์ ซึ่งต้องการปุ่มขนาดใหญ่และ Layout ที่เหมาะกับการใช้งานในสภาพแวดล้อม'
        'ที่มีความชื้น ฝุ่น หรือผู้ใช้สวมถุงมือ '
        'ตารางที่ 3.5-1 แสดงข้อกำหนด Breakpoints ทั้ง 4 ประเภท '
        'ตารางที่ 3.5-2 แสดงมาตรฐาน Touch UX ที่ใช้ทั่วทั้งระบบ '
        'และแผนภาพที่ 3.5-1 แสดงภาพรวมโครงสร้าง Layout ในแต่ละ Breakpoint'
    )

    h_subsection(doc, '3.5.1.1', 'Breakpoints และ Layout Grid')

    body_sub(doc,
        'การกำหนด Breakpoints ยึดหลัก Content-first โดยเลือกจุดแบ่งจากขนาดที่ Content เริ่ม'
        'ดูแย่ลง ไม่ใช่จากขนาด Device ยี่ห้อใดยี่ห้อหนึ่ง ทำให้ระบบรองรับอุปกรณ์ทุกประเภทได้อย่าง'
        'ยืดหยุ่น CSS Framework ที่ใช้คือ Tailwind CSS หรือ Bootstrap 5 ซึ่งมี Breakpoint System '
        'ในตัว ทำให้ทีมพัฒนาสามารถ implement ได้อย่างสม่ำเสมอทั่วทั้งระบบ'
    )

    tbl_cap(doc, 'ตารางที่ 3.5-1 Responsive Breakpoint Specification และ Layout Grid')
    tbl(doc,
        ['Breakpoint', 'ช่วงความกว้าง', 'Layout', 'Navigation', 'อุปกรณ์เป้าหมาย', 'TOR Ref'],
        [
            ['Mobile', '< 768px', 'Single Column\n(Full Width)', 'Bottom Tab Bar', 'Smartphone (Keeper)', 'TOR 4.1, 4.11'],
            ['Tablet', '768–1,024px', '2-Column Grid\n(Collapsible)', 'Top Navbar\n+ Hamburger', 'Tablet (Vet/Nutrition)', 'TOR 4.1'],
            ['Desktop', '> 1,024px', 'Sidebar +\nMain Content', 'Left Sidebar\nNav', 'PC/Laptop (Admin)', 'TOR 4.1, 4.4'],
            ['Touch Kiosk', 'Fixed 1,024–1,280px', 'Large Button\nGrid (2×3)', 'On-screen\nLarge Buttons', 'Touch Screen (Warehouse)', 'TOR 4.9, 4.12'],
        ]
    )

    h_subsection(doc, '3.5.1.2', 'Touch Target Specification และ Browser Support')

    body_sub(doc,
        'มาตรฐาน Touch Target ที่กำหนดคือขนาดขั้นต่ำ 44 × 44 pixels ตามแนวทาง WCAG 2.2 '
        'Success Criterion 2.5.8 ซึ่งเป็นมาตรฐานสากลสำหรับการออกแบบ Interface ที่รองรับ'
        'ผู้ใช้ทุกกลุ่ม สำหรับ Touch Screen Kiosk ในคลังอาหารและคลังยา ปุ่มหลักต้องมีขนาด'
        'ไม่น้อยกว่า 88 × 88 pixels เพื่อรองรับการใช้งานขณะสวมถุงมือหรือมือเปียก '
        'Browser Support ครอบคลุม Edge, Chrome, Safari เวอร์ชันล่าสุด 2 เวอร์ชัน '
        'ตามที่ TOR ข้อ 4.4 กำหนด'
    )

    tbl_cap(doc, 'ตารางที่ 3.5-2 Touch UX Design Standards และ Browser Support Matrix')
    tbl(doc,
        ['รายการ', 'ข้อกำหนด', 'เหตุผล', 'อ้างอิง'],
        [
            ['Touch Target ขั้นต่ำ (ทั่วไป)', '44 × 44 px', 'WCAG 2.2 SC 2.5.8', 'W3C WCAG'],
            ['Touch Target (Kiosk หลัก)', '88 × 88 px', 'ใช้งานขณะสวมถุงมือ', 'TOR 4.9'],
            ['ระยะห่างระหว่างปุ่ม', '≥ 8 px', 'ป้องกัน Fat-finger Error', 'UX Best Practice'],
            ['Font ขั้นต่ำบน Mobile', '14 pt (TH Sarabun)', 'อ่านง่ายกลางแจ้ง', 'Internal Standard'],
            ['Font ขั้นต่ำบน Kiosk', '18 pt', 'มองเห็นจากระยะ 50 ซม.', 'TOR 4.9'],
            ['Contrast Ratio', '≥ 4.5:1', 'WCAG AA Level', 'WCAG 2.1'],
            ['Browser: Google Chrome', 'Latest 2 versions', 'เบราว์เซอร์หลัก', 'TOR 4.4'],
            ['Browser: Microsoft Edge', 'Latest 2 versions', 'Windows Default', 'TOR 4.4'],
            ['Browser: Safari', 'Latest 2 versions', 'iOS/macOS', 'TOR 4.4'],
            ['Orientation Support', 'Portrait + Landscape', 'Tablet usage', 'TOR 4.1'],
        ]
    )

    # Diagram 1
    add_img(doc, diagram_responsive_breakpoints(), w=5.8)
    fig_cap(doc, 'แผนภาพที่ 3.5-1 โครงสร้าง Responsive Design — Breakpoints และ Layout Grid '
                 'แสดงการจัดวาง UI ในแต่ละ Breakpoint ตั้งแต่ Desktop, Tablet, Mobile จนถึง Touch Kiosk')

    # ────────────────────────────────────────────────────────
    #  3.5.2  Touchscreen Flows
    # ────────────────────────────────────────────────────────
    h_section(doc, '3.5.2', 'Touchscreen Flow — ระบบคลังอาหารสัตว์และยา/เวชภัณฑ์')

    body_sec(doc,
        'การออกแบบ Touchscreen Flow สำหรับระบบคลังอาหารสัตว์ (TOR 4.9) และระบบยา/เวชภัณฑ์ '
        '(TOR 4.12) ตั้งอยู่บนหลักการ Zero-Training-Required กล่าวคือ เจ้าหน้าที่คลังที่ไม่เคย'
        'ใช้ระบบมาก่อนต้องสามารถดำเนินการรับเข้า เบิก และคืนสินค้าได้โดยไม่ต้องอ่านคู่มือ '
        'ผ่านการออกแบบ Affordance ที่ชัดเจน ใช้ไอคอนสากลร่วมกับข้อความภาษาไทย '
        'และใช้สีที่แตกต่างกันชัดเจนสำหรับแต่ละการดำเนินงาน '
        'หัวข้อนี้ประกอบด้วยเนื้อหา 2 ส่วน ส่วนแรก (หัวข้อ 3.5.2.1) อธิบาย Touchscreen Flow '
        'ของระบบคลังอาหารสัตว์พร้อม Wireframe 2 หน้าจอ '
        'และส่วนที่สอง (หัวข้อ 3.5.2.2) อธิบาย Touchscreen Flow ของระบบยา/เวชภัณฑ์ '
        'แผนภาพที่ 3.5-2 แสดง Wireframe ระบบคลังอาหาร '
        'และแผนภาพที่ 3.5-5 แสดง Wireframe ระบบยา/เวชภัณฑ์'
    )

    h_subsection(doc, '3.5.2.1', 'Touchscreen Flow — ระบบคลังอาหารสัตว์ (TOR 4.9)')

    body_sub(doc,
        'ระบบคลังอาหารสัตว์ออกแบบให้ทำงานบน Touch Screen ขนาด 21 นิ้วขึ้นไปที่ติดตั้งใน'
        'พื้นที่คลังสินค้า หน้าจอหลักประกอบด้วยปุ่มขนาดใหญ่ 6 ปุ่มสำหรับ 6 ฟังก์ชันหลัก '
        'ได้แก่ รับเข้าวัตถุดิบ เบิกวัตถุดิบ คืนวัตถุดิบ Stock Card รายการสต็อกคงเหลือ '
        'และพิมพ์บาร์โค้ด แต่ละปุ่มใช้สีเฉพาะเพื่อให้แยกแยะได้ในทันที '
        'เมื่อเจ้าหน้าที่กดปุ่มรับเข้า ระบบจะเปิดกล้องสแกนบาร์โค้ดโดยอัตโนมัติ '
        'หากสแกนสำเร็จ ข้อมูลสินค้าจะปรากฏบนหน้าจอพร้อมให้เจ้าหน้าที่ระบุจำนวนและยืนยัน '
        'ระบบจะบันทึก Stock Ledger และ Audit Log โดยอัตโนมัติ '
        'รองรับ TOR ข้อ 4.9.1 (รับเข้า) 4.9.2 (เบิก) 4.9.3 (คืน) 4.9.4 (Stock Card) '
        'และ 4.9.5 (สต็อกคงเหลือ) ครบถ้วน'
    )

    # Diagram 2
    add_img(doc, diagram_touch_wireframe_food(), w=5.8)
    fig_cap(doc, 'แผนภาพที่ 3.5-2 Wireframe หน้าจอ Touch Screen ระบบคลังอาหารสัตว์ '
                 'ซ้าย: เมนูหลัก 6 ปุ่ม Touch-optimized | ขวา: หน้าจอรับเข้าวัตถุดิบพร้อม Barcode Scan')

    h_subsection(doc, '3.5.2.2', 'Touchscreen Flow — ระบบยา/เวชภัณฑ์ (TOR 4.12)')

    body_sub(doc,
        'ระบบยา/เวชภัณฑ์ใช้หลักการออกแบบเดียวกับระบบคลังอาหาร แต่เพิ่มมิติความปลอดภัยของ'
        'ข้อมูลสุขภาพสัตว์ โดยการเบิกยาทุกครั้งต้องระบุ Animal ID ที่เชื่อมกับทะเบียนสัตว์ '
        'ซึ่งทำให้ระบบสามารถตรวจสอบ Drug Interaction และอัปเดตเวชระเบียนโดยอัตโนมัติ '
        'สำหรับยาที่รับเข้าจากการบริจาค ระบบจะบันทึกแหล่งที่มาแยกจากยาจัดซื้อ '
        'เพื่อความถูกต้องในการคิดค่าบริการรักษาสัตว์ตาม TOR ข้อ 4.18 '
        'หน้าจอเบิกยาประกอบด้วย 4 ขั้นตอนหลัก ได้แก่ ค้นหาสัตว์ สแกนบาร์โค้ดยา '
        'ระบุจำนวนและหมายเหตุ และยืนยันการเบิก ทุกขั้นตอนมี Validation ก่อนบันทึก'
    )

    # Diagram 5
    add_img(doc, diagram_touch_wireframe_pharma(), w=5.8)
    fig_cap(doc, 'แผนภาพที่ 3.5-5 Wireframe หน้าจอ Touch Screen ระบบยา/เวชภัณฑ์ '
                 'ซ้าย: เมนูหลัก | ขวา: หน้าจอเบิกยาพร้อม Barcode Scan และการเชื่อมเวชระเบียน')

    # ────────────────────────────────────────────────────────
    #  3.5.3  Barcode Flows
    # ────────────────────────────────────────────────────────
    h_section(doc, '3.5.3', 'Barcode Workflow — การบูรณาการบาร์โค้ดในกระบวนการทำงาน')

    body_sec(doc,
        'การบูรณาการ Barcode เข้าสู่กระบวนการทำงานของระบบ Thai Zoo ARK ครอบคลุม 2 ระบบหลัก '
        'ได้แก่ ระบบคลังอาหารสัตว์ (TOR 4.10) และระบบยา/เวชภัณฑ์ (TOR 4.13) '
        'เทคโนโลยีบาร์โค้ดที่เลือกใช้คือ Code-128 และ QR Code ซึ่งสแกนได้ทั้งจาก '
        'Barcode Scanner อุปกรณ์ภายนอกและกล้องของอุปกรณ์ผ่าน Library ZXing หรือ QuaggaJS '
        'โดยไม่ต้องติดตั้งซอฟต์แวร์เพิ่มเติม หัวข้อนี้ประกอบด้วยเนื้อหา 2 ส่วน '
        'ส่วนแรก (หัวข้อ 3.5.3.1) อธิบาย Barcode Technology Specification '
        'และส่วนที่สอง (หัวข้อ 3.5.3.2) อธิบาย Barcode Workflow ทุกประเภทด้วย Sequence Diagram '
        'ตารางที่ 3.5-3 แสดง Barcode Technology Specification '
        'และแผนภาพที่ 3.5-3 แสดง Barcode Sequence Diagram ฉบับสมบูรณ์'
    )

    h_subsection(doc, '3.5.3.1', 'Barcode Technology Specification')

    body_sub(doc,
        'ระบบรองรับ Barcode ทั้งแบบ 1D (Code-128) และ 2D (QR Code) ซึ่งเป็นมาตรฐานที่ใช้งาน'
        'กันทั่วไปในอุตสาหกรรมคลังสินค้าและการแพทย์ การสแกนทำได้ผ่าน 2 ช่องทาง คือ '
        'Barcode Scanner USB/Bluetooth ที่เชื่อมต่อกับ Touch Screen Kiosk '
        'และกล้องของอุปกรณ์ผ่าน Web Browser API ทำให้เจ้าหน้าที่สามารถใช้ Smartphone '
        'สแกนบาร์โค้ดได้โดยตรงในภาคสนาม การพิมพ์ Label ใช้ Thermal Printer '
        'ที่รองรับ ZPL (Zebra Programming Language) หรือ ESC/POS ซึ่งเป็นมาตรฐาน'
        'ที่เครื่องพิมพ์ Thermal ทุกยี่ห้อรองรับ'
    )

    tbl_cap(doc, 'ตารางที่ 3.5-3 Barcode Technology Specification')
    tbl(doc,
        ['รายการ', 'ข้อกำหนด', 'ระบบที่ใช้', 'TOR Ref'],
        [
            ['Barcode Type', 'Code-128 (1D)', 'คลังอาหาร, ยา/เวชภัณฑ์', 'TOR 4.10, 4.13'],
            ['QR Code (2D)', 'QR Code v2.0', 'ข้อมูลสัตว์, Label เพิ่มเติม', 'Value-add'],
            ['Scan Method 1', 'USB/BT Barcode Scanner', 'Touch Kiosk (Warehouse)', 'TOR 4.10.1, 4.13.1'],
            ['Scan Method 2', 'Camera API (ZXing/QuaggaJS)', 'Mobile / Tablet', 'TOR 4.10.1'],
            ['Label Printing', 'Thermal Printer (ZPL/ESC-POS)', 'คลังอาหาร, ยา/เวชภัณฑ์', 'TOR 4.10.2, 4.13.2'],
            ['Label Content', 'ชื่อสินค้า + Lot + Exp + Barcode', 'สินค้ารับเข้าใหม่', 'TOR 4.10.2'],
            ['Data on Barcode', 'Internal Item ID (UUID)', 'เชื่อมกับ DB ภายใน', 'Internal'],
            ['Error Handling', 'ไม่พบบาร์โค้ด → Manual Entry', 'Fallback mode', 'UX Requirement'],
        ]
    )

    h_subsection(doc, '3.5.3.2', 'Barcode Sequence Diagram — ลำดับขั้นตอนการทำงาน')

    body_sub(doc,
        'Sequence Diagram ในแผนภาพที่ 3.5-3 แสดงลำดับขั้นตอนการทำงานของ Barcode Workflow '
        'ครบ 10 ขั้นตอน ตั้งแต่เจ้าหน้าที่สแกนบาร์โค้ด (ขั้นที่ 1) จนถึงการพิมพ์ Label '
        'สำหรับสินค้ารับเข้าใหม่ (ขั้นที่ 10) โดยการสื่อสารระหว่าง Touch UI และ API Server '
        'ใช้ HTTP/REST Protocol ผ่าน Secure Connection (HTTPS/SSL) ทุกขั้นตอน '
        'การตอบสนองของระบบต้องไม่เกิน 500ms สำหรับการ Lookup บาร์โค้ด '
        'เพื่อให้การใช้งานหน้างานเป็นไปอย่างลื่นไหล '
        'กรณีที่เครือข่ายช้า ระบบจะ Cache ข้อมูลสินค้าที่ใช้บ่อยไว้ใน Local Storage '
        'เพื่อให้การสแกนยังทำงานได้แม้อินเทอร์เน็ตมีปัญหาชั่วคราว'
    )

    # Diagram 3
    add_img(doc, diagram_barcode_sequence(), w=5.8)
    fig_cap(doc, 'แผนภาพที่ 3.5-3 Barcode Workflow Sequence Diagram — '
                 'แสดง 10 ขั้นตอนการทำงานตั้งแต่สแกนบาร์โค้ดจนถึงบันทึก Stock Ledger '
                 'และพิมพ์ Label ครอบคลุมทั้งระบบคลังอาหารสัตว์และยา/เวชภัณฑ์')

    # ────────────────────────────────────────────────────────
    #  3.5.4  Mobile Keeper Interface
    # ────────────────────────────────────────────────────────
    h_section(doc, '3.5.4', 'Mobile Interface — โมดูลผู้ดูแลสัตว์ (Zookeeper Module)')

    body_sec(doc,
        'โมดูลผู้ดูแลสัตว์ (TOR 4.11) ออกแบบสำหรับการใช้งานบนมือถือในภาคสนามเป็นหลัก '
        'โดยเจ้าหน้าที่ผู้ดูแลสัตว์ (Zookeeper) กว่า 100–200 คนทั่ว 7 สวนสัตว์ต้องสามารถ'
        'บันทึกข้อมูลการดูแลสัตว์ประจำวันและแจ้งเหตุการณ์สำคัญได้ทันทีที่เกิดขึ้น '
        'โดยไม่ต้องรอกลับมายังโต๊ะทำงาน การออกแบบยึดหลัก 3-tap Rule คือ '
        'เจ้าหน้าที่ต้องเข้าถึงฟังก์ชันใดก็ได้ใน 3 การแตะหน้าจอหรือน้อยกว่า '
        'หัวข้อนี้ประกอบด้วยเนื้อหา 2 ส่วน '
        'ส่วนแรก (หัวข้อ 3.5.4.1) อธิบายการออกแบบ 9 ฟังก์ชันหลักพร้อมตารางอ้างอิง TOR '
        'และส่วนที่สอง (หัวข้อ 3.5.4.2) อธิบาย Notification System สำหรับการแจ้งเตือน'
        'ไปยังสัตวแพทย์และผู้เกี่ยวข้องโดยอัตโนมัติ '
        'ตารางที่ 3.5-4 แสดง 9 ฟังก์ชันพร้อม TOR Reference '
        'และแผนภาพที่ 3.5-4 แสดง Wireframe 3 หน้าจอหลักของโมดูลผู้ดูแลสัตว์'
    )

    h_subsection(doc, '3.5.4.1', '9 ฟังก์ชันหลักของโมดูลผู้ดูแลสัตว์')

    body_sub(doc,
        'โมดูลผู้ดูแลสัตว์ประกอบด้วย 9 ฟังก์ชันตาม TOR ข้อ 4.11 ซึ่งครอบคลุมทั้งงาน'
        'ปกติประจำวัน ได้แก่ การรับอาหาร การตรวจสภาพแวดล้อม และการส่งเสริมคุณภาพชีวิตสัตว์ '
        'และงานแจ้งเหตุการณ์สำคัญ 5 ประเภท ได้แก่ เกิด ป่วย ตั้งครรภ์/วางไข่ ตาย และสูญหาย '
        'แต่ละฟังก์ชันออกแบบให้กรอกข้อมูลได้รวดเร็วโดยใช้ปุ่มตัวเลือกสำเร็จรูป (Radio Buttons) '
        'และ Dropdown ที่โหลดข้อมูลจากทะเบียนสัตว์โดยอัตโนมัติ '
        'ลดการพิมพ์ข้อความให้น้อยที่สุด เพื่อความสะดวกในการใช้งานภาคสนาม'
    )

    tbl_cap(doc, 'ตารางที่ 3.5-4 9 ฟังก์ชันโมดูลผู้ดูแลสัตว์ — ข้อกำหนดและ TOR Reference')
    tbl(doc,
        ['#', 'ฟังก์ชัน', 'ข้อมูลที่บันทึก', 'Notification', 'TOR Ref'],
        [
            ['1', 'รายชื่อสัตว์ที่ดูแล', 'แสดงรายการสัตว์ตาม Assignment', '-', '4.11.1'],
            ['2', 'รายการรับอาหาร', 'ประเภทอาหาร, ปริมาณ, เวลา', '-', '4.11.2'],
            ['3', 'ตรวจสภาพแวดล้อม', 'อุณหภูมิ, ความสะอาด, ความปลอดภัย', '-', '4.11.3'],
            ['4', 'ส่งเสริมคุณภาพชีวิต', 'ประเภท Enrichment, การตอบสนองสัตว์', '-', '4.11.4'],
            ['5', 'แจ้งเกิด', 'วันเวลา, เพศ, น้ำหนัก, สถานะ', 'สัตวแพทย์, Registry', '4.11.5'],
            ['6', 'แจ้งป่วย', 'อาการ, ระดับความรุนแรง, ตำแหน่ง', 'สัตวแพทย์ (ทันที)', '4.11.6'],
            ['7', 'แจ้งตั้งครรภ์/วางไข่', 'วันที่สังเกต, ประมาณวันครบกำหนด', 'สัตวแพทย์, Nutritionist', '4.11.7'],
            ['8', 'แจ้งตาย', 'วันเวลา, สาเหตุเบื้องต้น, พยาน', 'สัตวแพทย์, Admin (ทันที)', '4.11.8'],
            ['9', 'แจ้งสูญหาย', 'วันเวลา, ตำแหน่งล่าสุดที่พบ', 'ทีมค้นหา, Admin (ทันที)', '4.11.9'],
        ]
    )

    h_subsection(doc, '3.5.4.2', 'In-app Notification System สำหรับเหตุการณ์สำคัญ')

    body_sub(doc,
        'เมื่อเจ้าหน้าที่บันทึกเหตุการณ์สำคัญประเภท ป่วย ตาย หรือสูญหาย ระบบจะส่ง '
        'Push Notification ไปยังกลุ่มผู้รับที่กำหนดโดยอัตโนมัติภายใน 30 วินาที '
        'โดยใช้ Web Push API ที่รองรับทั้ง Android และ iOS '
        'Notification ประกอบด้วยชื่อสัตว์ ประเภทเหตุการณ์ สวนสัตว์ที่เกิดเหตุ '
        'และ Deep Link ที่คลิกแล้วเปิดหน้าจอที่เกี่ยวข้องโดยตรง '
        'ผู้รับ Notification แบ่งตามประเภทเหตุการณ์ดังแสดงในตารางที่ 3.5-4 '
        'โดย Role ของผู้รับจะถูกจัดการผ่านระบบสิทธิ์ (RBAC) ที่กำหนดไว้ใน TOR ข้อ 4.5'
    )

    # Diagram 4
    add_img(doc, diagram_mobile_keeper(), w=6.0)
    fig_cap(doc, 'แผนภาพที่ 3.5-4 Wireframe หน้าจอ Mobile ระบบผู้ดูแลสัตว์ '
                 'ซ้าย: หน้ารายชื่อสัตว์ที่ดูแล | กลาง: เมนู 8 ฟังก์ชันบันทึกรายวัน | '
                 'ขวา: แบบฟอร์มแจ้งเกิดพร้อมการแจ้งเตือนสัตวแพทย์อัตโนมัติ')

    # Save
    os.makedirs(OUT_DIR, exist_ok=True)
    doc.save(OUT_FILE)
    print(f"✅  Saved: {OUT_FILE}")


if __name__ == '__main__':
    build()
